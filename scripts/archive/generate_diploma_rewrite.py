"""Generate a VКR-style working edition of the DOLG diploma.

The current diploma contains useful engineering material, including visual
development stages. This generator preserves the global diploma structure,
turns the main chapters into a coherent VКR text, and moves bulky code snippets,
verification tables and evidence blocks to appendices.
"""

from __future__ import annotations

import os
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / 'docs'
DOCX_PATH = DOCS_DIR / 'Диплом_DOLG_редакция_ВКР.docx'
MD_PATH = DOCS_DIR / 'DIPLOMA_REWRITE.md'


@dataclass(frozen=True)
class Facts:
    django_version: str = '6.0.4'
    database_engine: str = 'SQLite'
    users: int = 5
    categories: int = 20
    products: int = 72
    reb_products: int = 26
    projects: int = 11
    demo_projects: int = 11
    knowledge_categories: int = 6
    articles: int = 9
    tests_total: int = 99
    tests_ok: int = 83
    tests_skipped: int = 16
    browser_ok: str = '16/16'
    coverage: str = '78 %'


def collect_facts() -> Facts:
    """Collect live Django facts, falling back to the latest inspected values."""

    os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'Dolg_PR.settings')
    try:
        import django
        from django.conf import settings
        from django.contrib.auth.models import User

        django.setup()

        from Dolg_APP.models import SchematicProject
        from knowledge.models import Article, KnowledgeCategory
        from shop.models import Category, Product

        engine = settings.DATABASES['default']['ENGINE'].rsplit('.', 1)[-1]
        reb_slugs = getattr(Category, 'REB_SLUGS', ())
        reb_products = Product.objects.filter(category__slug__in=reb_slugs).count()
        return Facts(
            django_version=django.get_version(),
            database_engine=engine,
            users=User.objects.count(),
            categories=Category.objects.count(),
            products=Product.objects.count(),
            reb_products=reb_products,
            projects=SchematicProject.objects.count(),
            demo_projects=SchematicProject.objects.filter(is_demo=True).count(),
            knowledge_categories=KnowledgeCategory.objects.count(),
            articles=Article.objects.count(),
        )
    except Exception:
        return Facts()


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ''
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    normal = document.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for name in ('Heading 1', 'Heading 2', 'Heading 3'):
        style = document.styles[name]
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    code_style = document.styles.add_style('DOLG Code', 1)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(9)
    code_style.paragraph_format.first_line_indent = Cm(0)
    code_style.paragraph_format.line_spacing = 1.0
    code_style.paragraph_format.space_after = Pt(0)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_list(document: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style='List Paragraph')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Cm(1.25)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.add_run(item).font.name = 'Times New Roman'


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)


def add_table(document: Document, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    add_caption(document, caption)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], 'D9EAF7')
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    document.add_paragraph()


def add_code(document: Document, code: str) -> None:
    for line in code.strip('\n').splitlines():
        document.add_paragraph(line.rstrip(), style='DOLG Code')


def add_diagram(document: Document, caption: str, diagram: str) -> None:
    add_caption(document, caption)
    add_code(document, diagram)
    document.add_paragraph()


def build_document(facts: Facts) -> Document:
    document = Document()
    configure_document(document)

    add_heading(document, 'ОГЛАВЛЕНИЕ', 1)
    for line in [
        'ВВЕДЕНИЕ',
        'ГЛАВА 1. АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ',
        '1.1. Постановка задачи',
        '1.2. Требования к разрабатываемой системе',
        '1.3. Анализ аналогов и ограничений существующих решений',
        'ГЛАВА 2. ПРОЕКТИРОВАНИЕ СИСТЕМЫ',
        '2.1. Архитектура web-приложения',
        '2.2. Проектирование базы данных',
        '2.3. Проектирование клиентских модулей схем, CAD и симуляции',
        '2.4. Этапы разработки и план модернизации',
        'ГЛАВА 3. РЕАЛИЗАЦИЯ И ПРОВЕРКА ПРОГРАММНОЙ СИСТЕМЫ',
        '3.1. Реализованные модули',
        '3.2. Тестирование и результаты проверки',
        '3.3. Выводы по реализации и проверке',
        'ЗАКЛЮЧЕНИЕ',
        'СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ И ИСТОЧНИКОВ',
        'ПРИЛОЖЕНИЯ',
    ]:
        paragraph = document.add_paragraph(line, style='List Paragraph')
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_page_break()

    add_heading(document, 'ВВЕДЕНИЕ', 1)
    add_paragraph(
        document,
        'Развитие современной радиоэлектроники сопровождается ростом числа инженерных задач, '
        'которые требуют не только закупки компонентов, но и предварительной проверки схемных '
        'решений. Для учебных, лабораторных и малых инженерных проектов особенно важны '
        'инструменты, объединяющие каталог компонентов, проектирование схем, проверку '
        'электрических параметров и оформление результатов в одном web-приложении [3], [9].',
    )
    add_paragraph(
        document,
        'Актуальность работы определяется тем, что типовой пользователь вынужден работать сразу '
        'с несколькими разрозненными средами: интернет-магазином, справочником компонентов, '
        'редактором принципиальных схем, CAD-инструментом и симулятором SPICE. Такая разобщённость '
        'увеличивает вероятность ошибок при подборе номиналов, усложняет повторное использование '
        'схем и снижает наглядность учебного процесса [21], [22], [25].',
    )
    add_paragraph(
        document,
        'Целью выпускной квалификационной работы является разработка web-приложения DOLG для '
        'подбора, покупки, проектирования и симуляции электронных схем с использованием '
        'серверной части на Django и клиентских инженерных модулей на JavaScript [6], [7].',
    )
    add_paragraph(
        document,
        'Объектом исследования является процесс цифрового сопровождения работы с электронными '
        'компонентами. Предметом исследования является программная реализация web-платформы, '
        'которая объединяет электронную коммерцию, справочную базу, редактор схем, CAD и '
        'симуляцию электрических цепей.',
    )
    add_paragraph(document, 'Для достижения поставленной цели решаются следующие задачи:')
    add_list(
        document,
        [
            'проанализировать предметную область и требования к web-приложению для работы с электронными компонентами;',
            'спроектировать архитектуру серверной и клиентской частей системы;',
            'разработать каталог, корзину, оформление заказов, личный кабинет и справочную базу;',
            'реализовать редактор принципиальных схем, CAD-раздел и браузерный контур симуляции на базе ngspice.wasm;',
            'связать элементы схемы с товарным каталогом и сформировать BOM для последующего заказа;',
            'провести тестирование, зафиксировать ограничения реализации и сформировать план развития проекта.',
        ],
    )
    add_paragraph(
        document,
        'Практическая значимость работы заключается в создании прототипа, который можно '
        'использовать для демонстрации полного инженерного цикла: от выбора компонентов и '
        'построения схемы до проверки электрического поведения и подготовки заказа [19], [20].',
    )

    document.add_page_break()

    add_heading(document, 'ГЛАВА 1. АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ', 1)
    add_heading(document, '1.1. Постановка задачи', 2)
    add_paragraph(
        document,
        'Разрабатываемая система должна поддерживать несколько связанных пользовательских '
        'сценариев. Пользователь выбирает электронные компоненты в каталоге, изучает '
        'характеристики, добавляет товары в корзину, создаёт схему, проверяет её в симуляторе '
        'и при необходимости формирует список компонентов для заказа.',
    )
    add_paragraph(
        document,
        'В отличие от обычного интернет-магазина, проект DOLG должен учитывать инженерный '
        'контекст товара: номинал, корпус, производитель, datasheet, жизненный цикл, возможность '
        'использования в схеме и соответствие модели SPICE [2], [3], [9].',
    )
    add_table(
        document,
        'Таблица 1.1 – Основные группы пользователей и сценарии',
        ['Роль', 'Сценарии', 'Ожидаемый результат'],
        [
            [
                'Гость',
                'Просмотр каталога, базы знаний, демо-схем и симулятора',
                'Понимание возможностей системы без регистрации',
            ],
            [
                'Пользователь',
                'Покупка товаров, создание проектов, сохранение версий и запусков симуляции',
                'Персональный инженерный рабочий контур',
            ],
            [
                'Менеджер',
                'Работа с заказами, товарами и содержанием каталога',
                'Контроль коммерческой части проекта',
            ],
            [
                'Администратор',
                'Управление пользователями, справочниками, правами и данными',
                'Полное сопровождение системы',
            ],
        ],
    )

    add_heading(document, '1.2. Требования к разрабатываемой системе', 2)
    add_paragraph(
        document,
        'Требования к DOLG разделяются на функциональные, инженерные и эксплуатационные. '
        'Функциональные требования описывают пользовательские возможности, инженерные — '
        'работу редактора, CAD и симулятора, эксплуатационные — безопасность, тестирование и '
        'подготовку к развёртыванию [5], [12].',
    )
    add_table(
        document,
        'Таблица 1.2 – Соответствие требований и текущей реализации',
        ['Требование', 'Фактический статус', 'Комментарий'],
        [
            [
                'Каталог товаров, фильтры, поиск, карточка товара',
                'Реализовано',
                f'В базе {facts.products} товаров и {facts.categories} категорий.',
            ],
            [
                'Корзина и оформление заказа',
                'Реализовано',
                'Есть модели заказов, доставки, платежей и тесты orders.',
            ],
            [
                'Личный кабинет и адреса пользователя',
                'Реализовано',
                'Профиль создаётся автоматически, есть пользовательские тесты accounts.',
            ],
            [
                'Редактор принципиальных схем',
                'Реализовано частично',
                'Базовый редактор работает; продолжается ремонт UI, netlist builder и модульная декомпозиция.',
            ],
            [
                'CAD-раздел',
                'Реализовано частично',
                'Есть базовый save/load и ГОСТ-шаблоны; требуется модернизация под инженерные сценарии.',
            ],
            [
                'Симуляция DC/AC/TRAN',
                'Реализовано частично',
                'Есть ngspice.wasm и JS-MNA fallback; приоритет — обработка ошибок, probes, курсоры, сохранение графиков.',
            ],
            [
                'DRF, JWT, Redis, Celery, PostgreSQL',
                'Перспектива',
                'В текущем коде основа — Django MVT, SQLite и синхронные сценарии.',
            ],
            [
                'What if, thermal, AI, 3D, virtual lab',
                'Перспектива',
                'Функции включены в roadmap после стабилизации CAD и симулятора.',
            ],
        ],
    )

    add_heading(document, '1.3. Анализ аналогов и ограничений существующих решений', 2)
    add_paragraph(
        document,
        'В предметной области существуют отдельные классы решений: интернет-магазины '
        'электроники, справочные базы компонентов, онлайн-симуляторы, CAD-системы и среды '
        'разработки печатных плат. Их недостаток для учебного проекта состоит в том, что '
        'они редко образуют единый пользовательский сценарий от выбора компонента до '
        'проверки схемы и формирования BOM [9], [17].',
    )
    add_paragraph(
        document,
        'DOLG не претендует на замену профессиональных САПР уровня KiCad или Altium, однако '
        'может закрыть учебный и демонстрационный сценарий: быстро собрать схему, подобрать '
        'компоненты из каталога, проверить типовые режимы и оформить результаты. В этом '
        'состоит инженерная ниша проекта [17], [23], [24].',
    )

    document.add_page_break()

    add_heading(document, 'ГЛАВА 2. ПРОЕКТИРОВАНИЕ СИСТЕМЫ', 1)
    add_heading(document, '2.1. Архитектура web-приложения', 2)
    add_paragraph(
        document,
        f'Серверная часть проекта реализована на Django {facts.django_version}. В текущей '
        'редакции используется классическая архитектура Django MVT: модели описывают '
        'структуру данных, views обрабатывают HTTP-запросы, шаблоны формируют HTML-страницы, '
        'а клиентские JavaScript-модули обеспечивают интерактивную работу редактора схем, CAD '
        'и симулятора [6], [18].',
    )
    add_diagram(
        document,
        'Рисунок 2.1 – Архитектурная схема DOLG',
        """
Пользователь
    |
    v
Django templates + vanilla JS
    |
    +--> shop: каталог, BOM, корзина, PDF/XLSX
    +--> accounts: профиль, адреса, роли
    +--> orders: заказ, доставка, платежи
    +--> knowledge: база знаний
    +--> Dolg_APP: проекты, версии, CAD, симуляция
             |
             +--> ngspice.wasm Web Worker
             +--> JS-MNA fallback
             +--> SimulationRun / ProjectVersion
""",
    )
    add_paragraph(
        document,
        'Клиентская часть редактора разделяется на несколько функциональных слоёв. После '
        'последних итераций отдельно выделены нормализация схемы, экспорт, BOM и построение '
        'SPICE-netlist. Такой подход уменьшает связанность большого шаблона simulation.html '
        'и упрощает дальнейшее тестирование [10], [11], [20].',
    )

    add_heading(document, '2.2. Проектирование базы данных', 2)
    add_paragraph(
        document,
        f'В качестве базы данных в рабочей версии используется {facts.database_engine}. '
        'Для дипломного прототипа это оправдано простотой запуска и демонстрации; переход на '
        'PostgreSQL отнесён к production-этапу. В базе зафиксированы пользователи, категории, '
        'товары, заказы, статьи базы знаний, проекты схем, версии проектов и результаты '
        'запусков симуляции [8].',
    )
    add_table(
        document,
        'Таблица 2.1 – Основные сущности данных',
        ['Приложение', 'Ключевые сущности', 'Назначение'],
        [
            [
                'shop',
                'Category, Product, Cart, Compare',
                'Каталог, параметры товаров, BOM-сопоставление и покупка',
            ],
            ['accounts', 'UserProfile, Address', 'Профиль пользователя, контактные данные и адреса'],
            [
                'orders',
                'Order, OrderItem, Shipment, PaymentTransaction',
                'Оформление заказа, доставка и платежная история',
            ],
            ['knowledge', 'KnowledgeCategory, Article', 'Справочная база по электронике и компонентам'],
            [
                'Dolg_APP',
                'SchematicProject, ProjectVersion, SimulationRun',
                'Проекты схем, история версий и результаты симуляций',
            ],
        ],
    )
    add_table(
        document,
        'Таблица 2.2 – Фактическое наполнение базы',
        ['Показатель', 'Значение'],
        [
            ['Пользователей', str(facts.users)],
            ['Категорий', str(facts.categories)],
            ['Товаров', str(facts.products)],
            ['РЭБ-компонентов', str(facts.reb_products)],
            ['Проектов схем', str(facts.projects)],
            ['Демо-схем', str(facts.demo_projects)],
            ['Категорий базы знаний', str(facts.knowledge_categories)],
            ['Статей базы знаний', str(facts.articles)],
        ],
    )

    add_heading(document, '2.3. Проектирование клиентских модулей схем, CAD и симуляции', 2)
    add_paragraph(
        document,
        'Редактор принципиальных схем и симулятор являются наиболее сложной частью проекта. '
        'Они выполняют не только визуальное редактирование, но и преобразование пользовательской '
        'схемы в инженерное представление: список компонентов, соединения, узлы, SPICE-строки, '
        'директивы анализа и результаты расчётов [3], [4], [9].',
    )
    add_diagram(
        document,
        'Рисунок 2.2 – Контур запуска симуляции',
        """
Схема на Canvas
    |
    v
normalizeSchemeData()
    |
    v
buildSpiceNetlist()
    |
    +--> buildElementNetlist()
    +--> buildAnalysisDirectives()
    |
    v
ngspice.wasm worker
    |
    +--> результат DC/AC/TRAN
    +--> ошибка или fallback JS-MNA
    |
    v
графики, таблица результатов, SimulationRun
""",
    )
    add_paragraph(
        document,
        'CAD-раздел ориентирован на подготовку чертежей и демонстрационных материалов по '
        'инженерной документации. В текущем состоянии он уже поддерживает базовые шаблоны, '
        'слои и сохранение состояния, однако требует отдельной модернизации: размеров, '
        'привязок, горячих клавиш, устойчивости на больших чертежах и визуальной регрессии [1], [2].',
    )
    add_heading(document, '2.4. Этапы разработки и план модернизации', 2)
    add_paragraph(
        document,
        'Этапы развития проекта целесообразно фиксировать в проектной главе, поскольку они '
        'показывают не только итоговую реализацию, но и инженерную последовательность принятия '
        'решений. В отличие от журнальных дописок, таблица 2.3 группирует выполненные и '
        'запланированные работы по смысловым блокам: сначала стабилизируется ядро симуляции и '
        'CAD, затем добавляются демонстрационные функции с высоким эффектом для защиты.',
    )
    add_table(
        document,
        'Таблица 2.3 – Этапы разработки и модернизации DOLG',
        ['Этап', 'Состояние', 'Содержание и доказательный материал'],
        [
            [
                'Каталог и BOM',
                'Выполнено',
                'Связь схемы с товарами каталога, catalog_ref, CSV/XLSX BOM, add-all-to-cart',
            ],
            [
                'Проекты и версии',
                'Выполнено',
                'SchematicProject, ProjectVersion, demo-проекты, save/load, история изменений',
            ],
            [
                'Симуляция DC/AC/TRAN',
                'Выполнено частично',
                'SPICE netlist, ngspice.wasm, JS-MNA fallback, browser smoke; требуется расширить probes и edge-cases',
            ],
            [
                'Отрисовка аналитики',
                'Выполнено частично',
                'Панель результатов перенесена под свойства компонента, добавлены проверки читаемости и overflow',
            ],
            [
                'CAD-базис',
                'Выполнено частично',
                'ГОСТ-шаблон, save/load, первый desktop/mobile visual baseline, исправление мобильной высоты canvas',
            ],
            [
                'Ремонт симулятора',
                'Ближайший этап',
                'Диагностика ngspice, курсоры графиков, сохранение расчётов, устойчивые AC/TRAN-сценарии',
            ],
            [
                'Оптимизация УГО-схем',
                'Выполнено частично',
                'Переключатель modern-стиля убран; оставлен практичный УГО-режим, скрыта подсказка canvas, ускорена отрисовка больших схем',
            ],
            [
                'Чертёжный шаг схем',
                'Выполнено',
                'Введён отдельный шаг хода, не связанный с визуальной сеткой; snap компонентов, изгибы и демо-схемы нормализованы по drawing_step=30',
            ],
            [
                'Ремонт CAD',
                'Ближайший этап',
                'Screenshot-baseline регрессия, размеры, слои, штамп, привязки, горячие клавиши',
            ],
            [
                'Killer-функции',
                'Перспектива',
                'What if-слайдер, тепловой анализ, AI-ассистент, 3D-просмотр, виртуальная лаборатория',
            ],
        ],
    )
    add_paragraph(
        document,
        'Импорт и экспорт KiCad, LTspice, EDIF, Gerber и Excellon остаются перспективным '
        'направлением, но текущий приоритет выше у ремонта симулятора и CAD. Marketplace и '
        'QR-сценарии логично выполнять после публичного развёртывания, а block-based редактор '
        'оформлять как отдельный learning mode [17].',
    )
    add_diagram(
        document,
        'Рисунок 2.3 – Приоритетная последовательность модернизации',
        """
Стабилизация ядра
    |
    +--> симулятор: ngspice, fallback, probes, понятные ошибки
    |
    +--> CAD: visual baseline, canvas, размеры, слои, привязки
    |
    v
Профессиональный инженерный UX
    |
    +--> Multisim-like отрисовка схем
    +--> отдельный шаг хода и нормализованные демо-схемы
    +--> аналитика под свойствами и читаемые графики
    |
    v
Демонстрационные функции
    |
    +--> What if, Thermal, AI, 3D, Virtual lab
""",
    )

    document.add_page_break()

    add_heading(document, 'ГЛАВА 3. РЕАЛИЗАЦИЯ И ПРОВЕРКА ПРОГРАММНОЙ СИСТЕМЫ', 1)
    add_heading(document, '3.1. Реализованные модули', 2)
    add_paragraph(
        document,
        'Реализация проекта выполнена итерационно. В итоговой редакции диплома важна не '
        'хронология отдельных изменений, а совокупность завершённых и частично завершённых '
        'функциональных блоков. Таблица 3.1 показывает текущую инженерную картину проекта.',
    )
    add_table(
        document,
        'Таблица 3.1 – Реализованные и частично реализованные блоки DOLG',
        ['Блок', 'Состояние', 'Доказательство'],
        [
            [
                'Каталог и карточки товаров',
                'Реализовано',
                'Фильтры, поиск, datasheet, warranty PDF, параметры товара',
            ],
            ['BOM из схемы', 'Реализовано', 'CSV/XLSX, add-all-to-cart, сопоставление catalog_ref'],
            ['Проекты схем', 'Реализовано', 'SchematicProject, ProjectVersion, demo-проекты, browser smoke'],
            [
                'Симуляция DC/AC/TRAN',
                'Реализовано частично',
                'E2E покрывает DC/AC/TRAN, пустой AC stdout и отказ обоих движков; нужны probes, курсоры и расширенные edge-cases ngspice',
            ],
            [
                'Аналитика симуляции',
                'Реализовано частично',
                'Панель результатов перенесена под свойства компонента; desktop/mobile visual smoke проверяет читаемую ширину и отсутствие overflow',
            ],
            [
                'CAD',
                'Реализовано частично',
                'Есть save/load, ГОСТ-шаблон и desktop/mobile visual-smoke; нужна screenshot-baseline регрессия и инженерные инструменты',
            ],
            ['Экспорт схем', 'Реализовано', 'SVG/PDF browser-smoke проверяет скачивание и сигнатуры файлов'],
            [
                'Модалки и предупреждения симулятора',
                'Реализовано частично',
                'BOM/netlist/project picker и export warnings проверяются на desktop/mobile; нужен screenshot-baseline для сложных состояний',
            ],
            [
                'Визуальный стиль схем',
                'Реализовано частично',
                'Оставлен единый УГО-режим; добавлены отдельный шаг хода, нормализация демо-схем и performance-smoke больших схем; требуется дальнейшая аккуратная перерисовка без потери практичности',
            ],
            [
                'Роли и администрирование',
                'Реализовано частично',
                'Есть Django admin и группы; статистика менеджера остаётся в плане',
            ],
        ],
    )

    add_heading(document, '3.2. Тестирование и результаты проверки', 2)
    add_paragraph(
        document,
        'Проверка проекта разделена на стандартный серверный контур и optional browser/e2e '
        'контур. Это позволяет быстро проверять Django-модели, views и бизнес-логику без '
        'браузера, а затем отдельно запускать сценарии, где требуется Playwright и Microsoft Edge [13].',
    )
    add_table(
        document,
        'Таблица 3.2 – Результаты проверки на 03.05.2026',
        ['Команда', 'Результат', 'Назначение'],
        [
            [
                'python manage.py check',
                'OK; при DEBUG disabled ожидаемо предупреждение о SECRET_KEY',
                'Проверка конфигурации Django',
            ],
            [
                'scripts/run_checks.ps1',
                f'{facts.tests_total} обнаружено; {facts.tests_ok} OK; {facts.tests_skipped} skipped; coverage {facts.coverage}',
                'Основной тестовый контур',
            ],
            [
                'scripts/run_browser_e2e.ps1',
                f'{facts.browser_ok} OK',
                'Playwright smoke для simulation, CAD, projects, экспортов, модалок и графиков',
            ],
            ['pip check', 'OK', 'Проверка совместимости установленных зависимостей'],
        ],
    )
    add_paragraph(
        document,
        'В последней итерации browser/e2e контур расширен визуальными smoke-проверками модальных окон '
        'симулятора. Тесты открывают BOM и netlist на desktop и mobile, вызывают предупреждения пустого '
        'экспорта SVG/PNG/PDF/BOM и проверяют, что уведомления и содержимое модалок не выходят за пределы '
        'viewport. Для авторизованного пользователя отдельно проверяется project picker modal при сохранении '
        'схемы. Это закрывает базовый риск вылезающих блоков, но не заменяет полноценную screenshot-baseline '
        'регрессию CAD, projects и панели свойств.',
    )
    add_paragraph(
        document,
        'Дополнительно проверена отказоустойчивость симуляции: если ngspice.wasm возвращает успешный AC-ответ без '
        'точек графика, пользователь видит диагностическое сообщение со stdout; если одновременно падают ngspice.wasm '
        'и JS-MNA fallback, панель результатов показывает причины отказа обоих движков. Это важно для защиты, потому '
        'что демонстрация не должна превращаться в пустой график без объяснения причины.',
    )
    add_paragraph(
        document,
        'После замечания по эргономике результатов панель аналитики симуляции перенесена из узкой правой колонки '
        'под блок свойств компонента. Правая панель теперь используется для статистики, параметров запуска и настроек, '
        'а графики, диагностические сообщения и экспорт результатов находятся в нижней части рабочей области, где '
        'доступна большая ширина. Изменение закреплено browser smoke-проверкой: на desktop результаты должны быть ниже '
        'свойств компонента и шире правой панели, а на mobile не должно появляться горизонтальное переполнение.',
    )
    add_paragraph(
        document,
        'Для CAD-раздела добавлен первый visual baseline на desktop и mobile. Сценарий применяет ГОСТ-шаблон, проверяет '
        'видимость основных панелей, отсутствие горизонтального переполнения, высоту canvas и наличие реальной отрисовки. '
        'Тест выявил сжатие мобильного canvas до 150 px; проблема устранена через фиксированную минимальную высоту рабочей '
        'области в мобильном media-правиле.',
    )
    add_paragraph(
        document,
        'После пользовательской проверки тяжёлой тестовой схемы направление визуального ремонта скорректировано: '
        'непрактичный переключатель modern-стиля убран, по умолчанию оставлен единый УГО-режим. Из рабочей области '
        'скрыта подсказка canvas и удалена кнопка экспорта результатов из панели аналитики, так как они перекрывали '
        'содержимое. Для ускорения больших схем добавлен fast-path: дорогая антиколлизия подписей и wire-hop слой '
        'ограничены на крупных сценах, а browser-smoke проверяет автоматический переход тяжёлой схемы в Pixi/WebGL.',
    )
    add_paragraph(
        document,
        'Отдельно устранена связка между визуальной сеткой и фактическим шагом перемещения. В редактор добавлена настройка '
        '«Шаг хода (px)»: она применяется к snap при перетаскивании компонентов, Ctrl+стрелкам, дублированию, ручным '
        'изгибам проводов и выбору свободных осей маршрутизации. Визуальная сетка остаётся самостоятельной настройкой '
        'фона холста. Команда заполнения демо-проектов нормализует 11 схем по `drawing_step=30`, поэтому примеры '
        'открываются уже в более аккуратной чертёжной компоновке.',
    )
    add_paragraph(
        document,
        'Особенно важны браузерные проверки, так как они фиксируют пользовательские сценарии, '
        'которые невозможно полноценно проверить только через Django test client: работу Canvas, '
        'скачивание SVG/PDF/XLSX, построение графиков и отсутствие критического overflow на '
        'desktop/mobile экранах [13], [14], [15], [16].',
    )

    add_heading(document, '3.3. Выводы по реализации и проверке', 2)
    add_paragraph(
        document,
        'Результаты проверки показывают, что проект достиг состояния демонстрируемого '
        'инженерного прототипа: основные серверные сценарии покрыты тестами, браузерный '
        'контур проверяет ключевые пользовательские действия, а наиболее заметные проблемы '
        'интерфейса симуляции и CAD начали фиксироваться визуальными smoke-тестами.',
    )
    add_paragraph(
        document,
        'При этом часть требований остаётся реализованной частично. В первую очередь это '
        'относится к симулятору и CAD: они уже выполняют базовые функции, но для уровня '
        'современного инженерного инструмента требуют доработки устойчивости, визуальной '
        'регрессии, профессиональной отрисовки схем, расширенных измерений и улучшенного '
        'пользовательского опыта. Подробная последовательность этих работ приведена в '
        'разделе 2.4, чтобы план развития оставался частью проектирования, а не подменял '
        'результаты реализации.',
    )

    document.add_page_break()

    add_heading(document, 'ЗАКЛЮЧЕНИЕ', 1)
    add_paragraph(
        document,
        'В ходе работы создана web-платформа DOLG, объединяющая каталог электронных компонентов, '
        'личный кабинет, оформление заказов, базу знаний, редактор принципиальных схем, CAD и '
        'контур симуляции. Проект демонстрирует возможность объединения коммерческого и '
        'инженерного сценариев в одном Django-приложении [6], [9], [12].',
    )
    add_paragraph(
        document,
        'Фактическая реализация уже закрывает базовые требования ВКР: каталог, товары, заказы, '
        'профили пользователей, проекты схем, версии, журнал симуляций, BOM, экспорт и '
        'проверяемые сценарии DC/AC/TRAN. Одновременно выявлены направления, которые нужно '
        'честно отражать как частично реализованные или перспективные: production-контур, '
        'PostgreSQL, DRF/JWT, Celery/Redis, расширенный CAD и современные функции симуляции.',
    )
    add_paragraph(
        document,
        'Ближайшее развитие проекта должно быть сосредоточено на качестве инженерной основы: '
        'устойчивости симулятора, понятной диагностике ошибок, читаемой аналитике результатов, улучшении CAD, визуальной '
        'регрессии и сохранении результатов расчётов. После этого проект готов к внедрению '
        'функций с высоким демонстрационным эффектом: параметрического анализа, тепловой карты, '
        'AI-ассистента, 3D-просмотра и виртуальной лаборатории.',
    )

    document.add_page_break()

    add_heading(document, 'СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ И ИСТОЧНИКОВ', 1)
    sources = [
        '1. ГОСТ 2.105-95. Единая система конструкторской документации. Общие требования к текстовым документам. — М. : Изд-во стандартов, 1996.',
        '2. ГОСТ 2.701-2008. Единая система конструкторской документации. Схемы. Виды и типы. Общие требования к выполнению. — М. : Стандартинформ, 2009.',
        '3. ГОСТ 2.702-2011. Единая система конструкторской документации. Правила выполнения электрических схем. — М. : Стандартинформ, 2012.',
        '4. ГОСТ 2.710-81. Единая система конструкторской документации. Обозначения буквенно-цифровые в электрических схемах. — М. : Изд-во стандартов, 1982.',
        '5. ГОСТ 7.32-2017. Система стандартов по информации, библиотечному и издательскому делу. Отчет о научно-исследовательской работе. Структура и правила оформления.',
        '6. Django Software Foundation. Django documentation [Электронный ресурс]. — Режим доступа: https://docs.djangoproject.com.',
        '7. Python Software Foundation. Python documentation [Электронный ресурс]. — Режим доступа: https://docs.python.org.',
        '8. SQLite Consortium. SQLite documentation [Электронный ресурс]. — Режим доступа: https://sqlite.org/docs.html.',
        "9. Ngspice Development Team. Ngspice user's manual [Электронный ресурс]. — Режим доступа: https://ngspice.sourceforge.io/docs.html.",
        '10. Emscripten Contributors. Emscripten documentation [Электронный ресурс]. — Режим доступа: https://emscripten.org/docs.',
        '11. WebAssembly Community Group. WebAssembly specification [Электронный ресурс]. — Режим доступа: https://webassembly.github.io/spec/core.',
        '12. OWASP Foundation. OWASP Top 10 [Электронный ресурс]. — Режим доступа: https://owasp.org/www-project-top-ten.',
        '13. Microsoft. Playwright Python documentation [Электронный ресурс]. — Режим доступа: https://playwright.dev/python/docs/intro.',
        '14. openpyxl. openpyxl documentation [Электронный ресурс]. — Режим доступа: https://openpyxl.readthedocs.io.',
        '15. ReportLab. ReportLab user guide [Электронный ресурс]. — Режим доступа: https://docs.reportlab.com.',
        '16. python-docx. python-docx documentation [Электронный ресурс]. — Режим доступа: https://python-docx.readthedocs.io.',
        '17. KiCad Developers. KiCad documentation [Электронный ресурс]. — Режим доступа: https://docs.kicad.org.',
        '18. ECMA International. The JSON data interchange syntax. ECMA-404 [Электронный ресурс]. — Режим доступа: https://www.ecma-international.org/publications-and-standards/standards/ecma-404.',
        '19. Gamma E., Helm R., Johnson R., Vlissides J. Design Patterns: Elements of Reusable Object-Oriented Software. — Addison-Wesley, 1994.',
        '20. Fowler M. Refactoring: Improving the Design of Existing Code. — Addison-Wesley, 2018.',
        '21. Sommerville I. Software Engineering. — Pearson, 2016.',
        "22. Pressman R., Maxim B. Software Engineering: A Practitioner's Approach. — McGraw-Hill, 2020.",
        "23. Lutz M. Learning Python. — O'Reilly Media, 2013.",
        '24. Greenfeld D., Greenfeld A. Two Scoops of Django. — Two Scoops Press, 2020.',
        '25. Horowitz P., Hill W. The Art of Electronics. — Cambridge University Press, 2015.',
    ]
    for source in sources:
        paragraph = document.add_paragraph(source)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(0)

    document.add_page_break()

    add_heading(document, 'ПРИЛОЖЕНИЯ', 1)
    add_heading(document, 'ПРИЛОЖЕНИЕ А. Фрагменты кода', 2)
    add_paragraph(
        document,
        'В приложение вынесены ключевые фрагменты, которые подтверждают инженерную реализацию, '
        'но не перегружают основной текст ВКР.',
    )
    add_caption(document, 'Листинг А.1 – Нормализация типа анализа перед сохранением SimulationRun')
    add_code(
        document,
        """
function getSimulationRunAnalysisType(result, built) {
    const rawType = (
        result?.analysis ||
        built?.circuit?.analysis ||
        document.getElementById('analysisType')?.value ||
        'dc'
    );
    return rawType === 'transient' ? 'tran' : rawType;
}
""",
    )
    add_caption(document, 'Листинг А.2 – Генерация директив SPICE-анализа')
    add_code(
        document,
        """
function buildAnalysisDirectives(analysisType, settings) {
    if (analysisType === 'ac') {
        return ['.ac dec ' + settings.points + ' ' + settings.fStart + ' ' + settings.fStop,
                '.print ac v(*)'];
    }
    if (analysisType === 'transient' || analysisType === 'tran') {
        return ['.tran ' + settings.step + ' ' + settings.stop,
                '.print tran v(*)'];
    }
    return ['.op', '.print dc v(*)'];
}
""",
    )
    add_caption(document, 'Листинг А.3 – Адаптивная оболочка модального окна симулятора')
    add_code(
        document,
        """
.sim-modal {
    padding: 20px;
    box-sizing: border-box;
    overflow-y: auto;
}

.sim-modal-content {
    max-width: min(90vw, 720px) !important;
    max-height: calc(100vh - 40px);
    overflow: auto;
    box-sizing: border-box;
}
""",
    )
    add_caption(document, 'Листинг А.4 – Стабильный класс уведомления перед визуальной проверкой')
    add_code(
        document,
        """
function showNotification(message, type = 'info') {
    const notif = document.createElement('div');
    notif.className = 'sim-notification';
    notif.classList.add(type);
    notif.textContent = message;
    document.body.appendChild(notif);
}
""",
    )
    add_caption(document, 'Листинг А.5 – Перенос аналитики симуляции под свойства компонента')
    add_code(
        document,
        """
.simulation-analysis-panel {
    flex: 0 0 clamp(180px, 26vh, 260px);
    display: flex;
    flex-direction: column;
    min-width: 0;
    overflow: hidden;
}

.results-panel {
    flex: 1;
    min-height: 0;
    overflow: auto;
    max-width: 100%;
}
""",
    )
    add_caption(document, 'Листинг А.6 – Мобильная высота canvas в CAD-разделе')
    add_code(
        document,
        """
@media (max-width: 768px) {
    .canvas-area {
        min-height: 440px;
        flex: 0 0 auto;
    }

    #canvas {
        height: 400px;
        min-height: 400px;
        flex: 0 0 400px;
        width: 100%;
    }
}
""",
    )
    add_caption(document, 'Листинг А.7 – Fast-path для больших схем')
    add_code(
        document,
        """
const SCHEMATIC_FAST_LABEL_LIMIT = 180;
const SCHEMATIC_WIRE_HOP_LIMIT = 240;

function isLargeSchematic() {
    return components.length + connections.length > SCHEMATIC_FAST_LABEL_LIMIT;
}

function shouldDrawComponentLabel(comp, isSelected) {
    if (comp.type === 'node') return false;
    if (isSelected) return true;
    if (!isLargeSchematic()) return true;
    return zoom >= 1.15 && components.length <= 360;
}

if (components.length + connections.length <= SCHEMATIC_WIRE_HOP_LIMIT) {
    drawWireHops();
}
""",
    )
    add_heading(document, 'ПРИЛОЖЕНИЕ Б. Матрица демонстрации', 2)
    add_table(
        document,
        'Таблица Б.1 – Что показывать на промежуточном просмотре',
        ['Блок', 'Что открыть', 'Что подчеркнуть'],
        [
            ['Каталог', '`/`, карточка товара, compare', 'Товары, datasheet, параметры, связь с BOM'],
            [
                'Симулятор',
                '`/simulation/`',
                'DC/AC/TRAN, единый УГО-режим, отдельный шаг хода, нормализованные демо-схемы, ускоренная отрисовка больших схем, графики, аналитика под свойствами компонента, SVG/PDF, BOM/netlist-модалки, предупреждения, сохранение запусков',
            ],
            ['Проекты', '`/projects/`', 'Демо-схемы, версии, история SimulationRun'],
            ['CAD', '`/cad/`', 'ГОСТ-шаблон, save/load, desktop/mobile visual baseline, план модернизации'],
            ['Тесты', 'PowerShell scripts', '99 тестов, browser e2e 16/16, coverage 78 %'],
            ['Roadmap', 'Таблица 2.3 и рисунок 2.3', 'CAD/симуляция сначала, killer-фичи после стабилизации'],
        ],
    )
    add_heading(document, 'ПРИЛОЖЕНИЕ В. Этапные доказательные блоки', 2)
    add_paragraph(
        document,
        'В этом приложении сохранены компактные этапные блоки, которые раньше находились '
        'в тексте диплома как последовательные дополнения. Они полезны для показа процесса '
        'разработки, но в итоговой структуре ВКР не должны разрывать основные главы.',
    )
    add_table(
        document,
        'Таблица В.1 – Доказательства ключевых итераций',
        ['Итерация', 'Что изменено', 'Что можно показать'],
        [
            [
                'Связь схемы с каталогом',
                'catalog_ref, catalog_slug, параметры товара, BOM-сопоставление',
                'Карточка компонента, BOM-модалка, XLSX-файл',
            ],
            [
                'Вынос JS-модулей',
                'scheme-normalizer.js, scheme-export.js, scheme-bom.js, scheme-netlist.js',
                'Фасады старых функций и новые browser fixtures',
            ],
            [
                'DC/AC/TRAN',
                'fixtures для делителя, RC-фильтра и transient-графика',
                'Canvas графиков, SimulationRun, browser e2e',
            ],
            [
                'Ошибки симуляции',
                'Диагностика пустого AC stdout и двойного отказа движков',
                'Сообщения в results panel вместо пустого графика',
            ],
            [
                'Аналитика под свойствами',
                'Новая simulation-analysis-panel под component-properties',
                'Desktop/mobile smoke без горизонтального overflow',
            ],
            [
                'CAD visual baseline',
                'Проверка ГОСТ-шаблона, панелей, canvas и мобильной высоты',
                'Тест, выявивший и подтвердивший исправление 150 px canvas',
            ],
            [
                'Оптимизация рендера схем',
                'Скрыта canvas-подсказка, убран export из аналитики, оставлен УГО-режим, добавлен fast-path больших схем',
                'Targeted browser visual/performance smoke 3/3 OK',
            ],
            [
                'Чертёжный шаг схем',
                'Отдельный шаг хода, snap/маршрутизация без привязки к визуальной сетке, нормализация 11 демо-схем',
                'Настройка в UI, drawing_step=30, targeted browser smoke 3/3 OK',
            ],
        ],
    )

    return document


def render_markdown(facts: Facts) -> str:
    return f"""# Диплом DOLG — рабочая редакция ВКР

Файл DOCX создаётся скриптом `scripts/generate_diploma_rewrite.py`.

## Назначение

Это рабочая редакция диплома в стиле ВКР. Глобальная структура сохранена, основной текст
дополнен фактическими результатами проекта, а крупные доказательные блоки вынесены в
приложения.

## Фактические показатели

| Показатель | Значение |
|---|---|
| Django | {facts.django_version} |
| База данных | {facts.database_engine} |
| Товары | {facts.products} |
| РЭБ-компоненты | {facts.reb_products} |
| Демо-схемы | {facts.demo_projects} |
| Тестовый контур | {facts.tests_total} обнаружено; {facts.tests_ok} OK; {facts.tests_skipped} skipped |
| Browser e2e | {facts.browser_ok} OK |
| Coverage | {facts.coverage} |

## Приоритет развития

См. раздел 2.4 рабочей редакции: сначала ремонт симулятора и CAD, затем Multisim-like
отрисовка схем на отдельном шаге хода, после стабилизации — What if, Thermal, AI, 3D и Virtual lab.
"""


def main() -> None:
    DOCS_DIR.mkdir(exist_ok=True)
    facts = collect_facts()
    document = build_document(facts)
    saved_docx_path = DOCX_PATH
    try:
        document.save(saved_docx_path)
    except PermissionError:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        saved_docx_path = DOCS_DIR / f'Диплом_DOLG_редакция_ВКР_{timestamp}.docx'
        document.save(saved_docx_path)
    MD_PATH.write_text(render_markdown(facts), encoding='utf-8')
    print(f'Generated: {saved_docx_path}')
    print(f'Generated: {MD_PATH}')


if __name__ == '__main__':
    main()

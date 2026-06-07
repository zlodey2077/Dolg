from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / 'docs'

SECTION_TITLE = '## Вопросы по разбираемой схеме'

SECTION = """## Вопросы по разбираемой схеме

Для защиты удобнее держать в голове одну основную демонстрационную схему: делитель напряжения от 9 В с выходным узлом Vout, GND и двумя резисторами. Через нее можно показать расчет, CAD-редактор, DRC, DC-симуляцию, expected vs measured, Engineering Review и учебное задание. Если нужно расширить демонстрацию, та же логика переносится на RC-фильтр и LED-ветвь.

1. **Почему для демонстрации выбран делитель напряжения?**
   Это простая, но показательная схема: в ней есть источник, GND, номиналы, выходной узел, расчетная формула, измеряемый результат и типовые ошибки. На ней легко показать весь маршрут DOLG без перегрузки схемой.

2. **Какая формула используется для выходного напряжения делителя?**
   Используется формула Vout = Vin * R2 / (R1 + R2), где R1 подключен к источнику, R2 к земле, а выходной узел находится между ними.

3. **Как получить примерно 3 В из источника 9 В?**
   Нужно, чтобы отношение R2 / (R1 + R2) было около 1/3. Например, при R1 = 20 кОм и R2 = 10 кОм получается около 3 В.

4. **Почему наличие GND критично для такой схемы?**
   Без GND у схемы нет опорного потенциала. Симулятору и review-слою не к чему привязать напряжения узлов, поэтому graph-layer и DRC должны выдавать предупреждение.

5. **Что проверяет Engineering Review на делителе?**
   Review проверяет наличие источника и GND, связность графа, путь до земли, корректность номиналов, распознавание output node, BOM-данные и соответствие измеренного результата расчетному.

6. **Что означает expected vs measured в этой схеме?**
   Expected — расчетное значение Vout по формуле. Measured — значение, полученное из симуляции или измерения. Если они близки в пределах допуска, задача считается выполненной корректно.

7. **Какая типовая ошибка хорошо демонстрируется на делителе?**
   Например, отсутствие GND, перепутанные номиналы R1 и R2, неверно выбранный выходной узел или слишком маленькое сопротивление, из-за которого растет ток и мощность.

8. **Как нагрузка влияет на делитель напряжения?**
   Если подключить нагрузку к Vout, она оказывается параллельно R2 и меняет эквивалентное сопротивление нижнего плеча. Поэтому реальный Vout может стать ниже расчетного без нагрузки.

9. **Почему важно оценивать мощность резисторов?**
   Даже простая схема может быть некорректной, если мощность на резисторе превышает допустимую. Лаборатория и review должны подсказать, где нужен запас.

10. **Как эта схема превращается в учебное задание?**
    Пользователь может рассчитать Vout, собрать делитель в CAD, запустить DC-анализ и отправить измеренный Vout. Learning grader проверяет численный ответ, структуру схемы и результат симуляции.

11. **Что меняется, если вместо делителя показать RC-фильтр?**
    Добавляется частотная область: нужно найти частоту среза fc = 1 / (2πRC), построить Bode plot и проверить точку около -3 дБ.

12. **Что можно спросить по LED-ветви?**
    Как подобрать ограничительный резистор, какой ток через светодиод будет безопасным, как проверить падение напряжения и почему нужен запас по мощности резистора.

13. **Какой ответ дать, если спросят: это просто калькулятор или инженерная проверка?**
    Это не только калькулятор. Система связывает формулу, схему, граф соединений, симуляцию, BOM и правило review, поэтому ошибка объясняется через конкретные факты проекта.

14. **Что делать, если результат симуляции не совпал с расчетом?**
    Проверить выбранный узел измерения, номиналы, наличие нагрузки, GND, тип анализа и единицы измерения. В DOLG эти причины должны проявиться как warnings или findings в review.

15. **Какой главный вывод по демонстрационной схеме?**
    На простой схеме видно главное преимущество проекта: пользователь не просто получает число, а проходит полный инженерный цикл от компонента и формулы до схемы, измерения, проверки и обучения.
"""


def latest_md() -> Path:
    return sorted(
        DOCS.glob('Речь_и_вопросы_к_защите_DOLG_*.md'), key=lambda p: p.stat().st_mtime, reverse=True
    )[0]


def latest_docx() -> Path:
    files = [
        path for path in DOCS.glob('Речь_и_вопросы_к_защите_DOLG_*.docx') if not path.name.startswith('~$')
    ]
    return sorted(files, key=lambda p: p.stat().st_mtime, reverse=True)[0]


def update_md() -> None:
    path = latest_md()
    text = path.read_text(encoding='utf-8')
    if SECTION_TITLE in text:
        before, rest = text.split(SECTION_TITLE, 1)
        if '\n## Возможные вопросы и ответы' in rest:
            _, after = rest.split('\n## Возможные вопросы и ответы', 1)
            text = before.rstrip() + '\n\n' + SECTION.strip() + '\n\n## Возможные вопросы и ответы' + after
        else:
            text = before.rstrip() + '\n\n' + SECTION.strip() + '\n'
    else:
        marker = '\n## Возможные вопросы и ответы'
        if marker in text:
            text = text.replace(marker, '\n\n' + SECTION.strip() + '\n' + marker, 1)
        else:
            text = text.rstrip() + '\n\n' + SECTION.strip() + '\n'
    path.write_text(text, encoding='utf-8')


def add_paragraph(doc: Document, text: str, *, bold: bool = False, size: int = 12) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold


def update_docx() -> Path:
    source = latest_docx()
    target = DOCS / 'Речь_и_вопросы_к_защите_DOLG_20260513_v5_с_вопросами_по_схеме_20260519.docx'
    doc = Document(str(source))
    full_text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
    if 'Вопросы по разбираемой схеме' not in full_text:
        add_paragraph(doc, 'Вопросы по разбираемой схеме', bold=True, size=15)
        for line in SECTION.splitlines()[2:]:
            if not line.strip():
                continue
            if line.startswith('Для защиты'):
                add_paragraph(doc, line.strip(), size=12)
            elif line[0].isdigit() and '. **' in line:
                question = line.replace('**', '').strip()
                add_paragraph(doc, question, bold=True, size=12)
            else:
                add_paragraph(doc, line.strip(), size=12)
    doc.save(str(target))
    Document(str(target))
    return target


def main() -> None:
    update_md()
    target = update_docx()
    print(f'Updated speech questions: {target}')


if __name__ == '__main__':
    main()

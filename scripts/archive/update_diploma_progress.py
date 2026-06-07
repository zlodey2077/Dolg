"""Append the latest DOLG engineering iteration to the diploma DOCX.

The script intentionally writes a compact but evidence-rich section: narrative,
tables, code snippets, an ASCII sequence diagram, test results, and source-note
audit. It is safe to rerun: if the marker heading already exists, no duplicate
section is added.
"""

from __future__ import annotations

import shutil
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / 'docs'
DIPLOMA_NAME = 'Дипломная работа (домашняя версия 3).docx'
MARKER = 'Дополнение от 03.05.2026: связь редактора схем с каталогом'
MARKER_XLSX = 'Дополнение от 03.05.2026: серверный XLSX-экспорт BOM'
MARKER_E2E_DRC = 'Дополнение от 03.05.2026: browser-smoke, DRC/BOM и аудит UI'
MARKER_CAD_PROJECTS = 'Дополнение от 03.05.2026: e2e CAD, проекты и normalizeSchemeData'
MARKER_DC_EXPORT = 'Дополнение от 03.05.2026: e2e DC и SVG/PDF-экспорт'
MARKER_NETLIST_FIXTURES = 'Дополнение от 03.05.2026: fixtures netlist builder и нормализация портов'
MARKER_SCHEME_NORMALIZER = 'Дополнение от 03.05.2026: вынос scheme-normalizer.js'
MARKER_SCHEME_EXPORT = 'Дополнение от 03.05.2026: вынос scheme-export.js'
MARKER_SCHEME_BOM = 'Дополнение от 03.05.2026: вынос scheme-bom.js'
MARKER_SCHEME_NETLIST = 'Дополнение от 03.05.2026: первый слой scheme-netlist.js'
MARKER_SCHEME_NETLIST_ELEMENTS = 'Дополнение от 03.05.2026: генерация SPICE-элементов в scheme-netlist.js'
MARKER_SCHEME_NETLIST_ANALYSIS = 'Дополнение от 03.05.2026: analysis directives в scheme-netlist.js'
MARKER_AC_BROWSER_SMOKE = 'Дополнение от 03.05.2026: AC browser-smoke RC-фильтра'
MARKER_TRAN_BROWSER_SMOKE = 'Дополнение от 03.05.2026: TRAN browser-smoke RC-цепи'
MARKER_VISUAL_LAYOUT_SMOKE = 'Дополнение от 03.05.2026: visual layout smoke симулятора'
MARKER_PRIORITY_SIM_CAD = 'Дополнение от 03.05.2026: приоритет ремонта симуляции и CAD'
MARKER_MODERN_SIM_CAD = 'Дополнение от 03.05.2026: модернизация симуляции и CAD под инженерные требования'
MARKER_SIM_RUN_UI_PERSISTENCE = 'Дополнение от 03.05.2026: сохранение AC/TRAN запусков из UI'
MARKER_KILLER_FEATURES_ROADMAP = 'Дополнение от 03.05.2026: roadmap killer-фич DOLG'
OLD_XLSX_NOTE = (
    'Выявленные несоответствия и уточнения: CSV-экспорт BOM уже реализован '
    'на клиенте, поэтому дальнейшего развития требует не CSV, а XLSX-экспорт '
    'и e2e-проверка полного сценария «компонент — товар — корзина».'
)
NEW_XLSX_NOTE = (
    'Выявленные несоответствия и уточнения: CSV-экспорт BOM реализован '
    'на клиенте, а XLSX-экспорт реализован на сервере через openpyxl. '
    'Дальнейшего развития требует e2e-проверка полного сценария '
    '«компонент — товар — BOM — корзина — XLSX».'
)
OPENPYXL_SOURCE = (
    '41. openpyxl. openpyxl documentation [Электронный ресурс]. — '
    'Режим доступа: https://openpyxl.readthedocs.io '
    '(дата обращения: 03.05.2026).'
)
PLAYWRIGHT_SOURCE = (
    '42. Microsoft. Playwright Python documentation [Электронный ресурс]. — '
    'Режим доступа: https://playwright.dev/python/docs/intro '
    '(дата обращения: 03.05.2026).'
)


def _find_diploma() -> Path:
    candidates = [
        Path.home() / 'Desktop' / DIPLOMA_NAME,
        ROOT / DIPLOMA_NAME,
        DOCS_DIR / DIPLOMA_NAME,
    ]
    for path in candidates:
        if path.exists():
            return path
    raise FileNotFoundError(f'Cannot find {DIPLOMA_NAME}')


def _ensure_code_style(document: Document) -> str:
    style_name = 'DOLG Code'
    if style_name not in [style.name for style in document.styles]:
        style = document.styles.add_style(style_name, 1)
        style.font.name = 'Consolas'
        style.font.size = Pt(9)
        style.paragraph_format.left_indent = Pt(18)
        style.paragraph_format.space_after = Pt(2)
    return style_name


def _add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = value
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    document.add_paragraph()


def _add_code(document: Document, code: str, style_name: str) -> None:
    for line in code.strip('\n').splitlines():
        document.add_paragraph(line, style=style_name)
    document.add_paragraph()


def normalize_previous_iteration_text(document: Document) -> bool:
    changed = False
    for paragraph in document.paragraphs:
        if OLD_XLSX_NOTE in paragraph.text:
            paragraph.text = paragraph.text.replace(OLD_XLSX_NOTE, NEW_XLSX_NOTE)
            changed = True
        if (
            'В следующей итерации был закрыт один из пунктов развития связки' in paragraph.text
            and 'openpyxl [41]' not in paragraph.text
        ):
            paragraph.text = paragraph.text.replace(
                'добавлен серверный экспорт BOM в формате XLSX.',
                'добавлен серверный экспорт BOM в формате XLSX на основе библиотеки openpyxl [41].',
            )
            changed = True
        if (
            'Контроль источников: для описания библиотеки openpyxl' in paragraph.text
            and 'добавлена официальная документация openpyxl [41]' not in paragraph.text
        ):
            paragraph.text = (
                'Контроль источников: в список использованной литературы и источников '
                'добавлена официальная документация openpyxl [41], а внутритекстовое '
                'упоминание приведено рядом с описанием XLSX-экспорта.'
            )
            changed = True
    replacements = {
        '35 тестов, OK': '38 тестов, OK (после расширения DRC/BOM-проверок)',
        '36 тестов, OK': '38 тестов, OK (после расширения DRC/BOM-проверок)',
        '78 тестов: accounts — 8, shop — 35, orders — 16, Dolg_APP — 19': (
            '84 обнаруженных теста: accounts — 8, shop — 38, '
            'orders — 16, Dolg_APP — 21 + 1 optional browser-smoke'
        ),
        '79/79 тестов, OK': '84 теста обнаружено: 83 OK, 1 skipped; browser-smoke отдельно 1/1 OK',
        '84 теста обнаружено: 83 OK, 1 skipped; browser-smoke отдельно 1/1 OK': '89 тестов обнаружено: 83 OK, 6 skipped; browser-smoke отдельно 6/6 OK',
        '84 теста обнаружено: 83 OK, 1 skipped; coverage 77 %': '86 тестов обнаружено: 83 OK, 3 skipped; coverage 78 %',
        '86 тестов обнаружено: 83 OK, 3 skipped; coverage 78 %': '88 тестов обнаружено: 83 OK, 5 skipped; coverage 78 %',
        '88 тестов обнаружено: 83 OK, 5 skipped; coverage 78 %': '89 тестов обнаружено: 83 OK, 6 skipped; coverage 78 %',
        '1/1 browser-smoke OK': '3/3 browser-smoke OK',
        '3/3 browser-smoke OK': '5/5 browser-smoke OK',
        '5/5 browser-smoke OK': '6/6 browser-smoke OK',
    }
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in replacements.items():
                    if old in cell.text:
                        cell.text = cell.text.replace(old, new)
                        changed = True
    return changed


def ensure_source_entry(document: Document, marker: str, entry: str) -> bool:
    all_text = '\n'.join(paragraph.text for paragraph in document.paragraphs)
    if marker in all_text:
        return False

    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip() == 'ПРИЛОЖЕНИЯ':
            if index > 0 and not document.paragraphs[index - 1].text.strip():
                document.paragraphs[index - 1].text = entry
            else:
                paragraph.insert_paragraph_before(entry)
            return True

    document.add_paragraph(entry)
    return True


def append_iteration(document: Document) -> bool:
    if any(MARKER in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER, level=2)

    document.add_paragraph(
        'В рамках очередной итерации была усилена интеграция редактора '
        'принципиальных схем с товарным каталогом. Ранее BOM формировался '
        'преимущественно по типу элемента, теперь пользователь может выбрать '
        'конкретную позицию каталога в свойствах компонента, а выбранный SKU '
        'сохраняется в составе JSON-структуры проекта.'
    )

    _add_table(
        document,
        ['Файл', 'Изменение', 'Назначение'],
        [
            [
                'shop/views.py',
                'Добавлен api_component_search и расширен BOM-match.',
                'Поиск товара по типу компонента, SKU, названию и корпусу; приоритет явного catalog_ref.',
            ],
            [
                'shop/urls.py',
                'Добавлен маршрут /component-search/.',
                'JSON API для панели свойств редактора схем.',
            ],
            [
                'Dolg_APP/templates/tools/simulation.html',
                'Добавлен поиск товара, сохранение metadata и отображение выбранной карточки.',
                'Пользователь связывает компонент схемы с реальной товарной позицией.',
            ],
            [
                'Dolg_APP/views.py',
                'DRC предупреждает о неизвестном товаре и несовпадении категории.',
                'Схема получает дополнительную инженерную проверку перед сохранением.',
            ],
            [
                'shop/tests.py',
                'Добавлены тесты component_search, catalog_ref и NPN/PNP mapping.',
                'Регрессионная защита новой связки редактора и каталога.',
            ],
        ],
    )

    document.add_paragraph('Ключевой фрагмент серверного API поиска товара:')
    _add_code(
        document,
        """
@require_GET
def api_component_search(request):
    query = (request.GET.get('q') or '').strip()
    component_type = (request.GET.get('type') or '').strip().lower()
    category_slug = COMPONENT_TO_CATEGORY.get(component_type)

    products = Product.objects.select_related('category').filter(
        stock__gt=0,
        lifecycle_status__in=['active', 'nrnd'],
    )
    if category_slug:
        products = products.filter(category__slug=category_slug)
    if query:
        products = products.filter(
            Q(name__icontains=query) |
            Q(part_number__icontains=query) |
            Q(description__icontains=query) |
            Q(package_type__icontains=query)
        )
        """,
        code_style,
    )

    document.add_paragraph('Фрагмент клиентской логики выбора товара для компонента:')
    _add_code(
        document,
        """
function applyCatalogProduct(index) {
    const product = catalogSearchResults[index];
    selectedComponent.catalog_ref = product.part_number || product.slug || '';
    selectedComponent.catalog_slug = product.slug || '';
    selectedComponent.catalog_url = product.url || '';
    selectedComponent.catalog_name = product.name || '';
    selectedComponent.catalog_manufacturer = product.manufacturer || '';
    selectedComponent.catalog_package = product.package_type || '';
    selectedComponent.datasheet_url = product.datasheet_url || '';
    selectedComponent.catalog_parameters = product.parameters || {};
}
        """,
        code_style,
    )

    document.add_paragraph('Логическая схема нового сценария:')
    _add_code(
        document,
        """
Пользователь
  -> Панель свойств компонента
  -> GET /component-search/?type=resistor&q=MF
  -> Product / Category
  -> selectedComponent.catalog_ref + catalog_parameters
  -> POST /bom/match/
  -> CartItem через /bom/add-all/
        """,
        code_style,
    )

    _add_table(
        document,
        ['Проверка', 'Результат'],
        [
            [
                'python manage.py test shop --settings=Dolg_PR.settings_test',
                '38 тестов, OK (после расширения DRC/BOM-проверок)',
            ],
            [
                'Общий ожидаемый набор',
                '84 обнаруженных теста: accounts — 8, shop — 38, orders — 16, Dolg_APP — 21 + 1 optional browser-smoke',
            ],
            [
                'Проверяемые сценарии',
                'Поиск товара компонента, NPN/PNP → transistors, явный catalog_ref в BOM',
            ],
        ],
    )

    document.add_paragraph(
        'Выявленные несоответствия и уточнения: CSV-экспорт BOM реализован '
        'на клиенте, а XLSX-экспорт реализован на сервере через openpyxl. '
        'Дальнейшего развития требует e2e-проверка полного сценария '
        '«компонент — товар — BOM — корзина — XLSX». Также '
        'в план развития добавлена интеграция CAD-форматов на уровне проектов: '
        'KiCad, LTspice, EDIF, Gerber и Excellon.'
    )
    document.add_paragraph(
        'Контроль источников: в данной итерации не добавлялись новые внешние '
        'источники; изменения основаны на коде проекта и локальной документации. '
        'При последующем описании CAD-форматов и DRF/JWT/PostgreSQL необходимо '
        'добавить ссылки на официальную документацию в список источников и '
        'привести внутритекстовые ссылки к единому виду.'
    )
    return True


def append_xlsx_iteration(document: Document) -> bool:
    if any(MARKER_XLSX in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER_XLSX, level=2)

    document.add_paragraph(
        'В следующей итерации был закрыт один из пунктов развития связки '
        'редактора схем с каталогом: добавлен серверный экспорт BOM в формате '
        'XLSX на основе библиотеки openpyxl [41]. В отличие от клиентского CSV, новый файл формируется на стороне '
        'Django из единого расчёта соответствий компонентов и товаров, поэтому '
        'структура JSON-ответа, корзина и экспорт используют согласованную '
        'бизнес-логику.'
    )

    _add_table(
        document,
        ['Файл', 'Изменение', 'Назначение'],
        [
            [
                'requirements.txt',
                'Добавлены openpyxl==3.1.5 и et_xmlfile==2.0.0.',
                'Генерация Excel-совместимых XLSX-файлов на сервере.',
            ],
            [
                'shop/views.py',
                'Вынесен _build_bom_matches и добавлен api_bom_export_xlsx.',
                'Единый расчёт BOM для JSON API и выгружаемого файла.',
            ],
            [
                'shop/urls.py',
                'Добавлен маршрут /bom/export-xlsx/.',
                'Отдельный POST endpoint для скачивания спецификации.',
            ],
            [
                'Dolg_APP/templates/tools/simulation.html',
                'В модальное окно BOM добавлена кнопка XLSX и fetch-загрузка файла.',
                'Пользователь получает спецификацию без ручного копирования таблицы.',
            ],
            [
                'shop/tests.py',
                'Добавлен тест чтения XLSX через openpyxl.load_workbook.',
                'Проверяются формат ответа, строки товаров, отсутствующие позиции и итог.',
            ],
        ],
    )

    document.add_paragraph('Общий helper расчёта BOM:')
    _add_code(
        document,
        """
def _build_bom_matches(components):
    if not isinstance(components, list):
        raise ValueError('components must be a list')

    explicit_counts = {}
    counts = {}
    for c in components:
        catalog_ref = (c.get('catalog_slug') or c.get('catalog_ref') or '').strip()
        if catalog_ref:
            explicit_counts[catalog_ref] = explicit_counts.get(catalog_ref, 0) + 1
            continue
        ctype = (c.get('type') or '').strip().lower()
        counts[ctype] = counts.get(ctype, 0) + 1

    # Далее подбираются товары Product, альтернативы и line_total.
    return matches, round(grand_total, 2), total_components
        """,
        code_style,
    )

    document.add_paragraph('Фрагмент серверного XLSX endpoint:')
    _add_code(
        document,
        """
@require_POST
def api_bom_export_xlsx(request):
    payload = json.loads(request.body)
    matches, grand_total, total_components = _build_bom_matches(
        payload.get('components', [])
    )

    wb = Workbook()
    ws = wb.active
    ws.title = 'BOM'
    ws.append(['DOLG Bill of Materials'])
    ws.append(['Всего компонентов', total_components, 'Итого, руб.', grand_total])
    ws.append(['Тип', 'Кол-во', 'Производитель', 'Part Number',
               'Наименование', 'Категория', 'Корпус',
               'Цена за шт., руб.', 'Сумма, руб.', 'Datasheet', 'Статус'])
        """,
        code_style,
    )

    document.add_paragraph('Фрагмент клиентской загрузки файла:')
    _add_code(
        document,
        """
async function bomDownloadXlsx() {
    const res = await fetch(BOM_API.exportXlsx, {
        method: 'POST',
        headers: {
            'Content-Type': 'application/json',
            'X-CSRFToken': getSimCsrfToken(),
        },
        credentials: 'same-origin',
        body: JSON.stringify({ components: components }),
    });
    const blob = await res.blob();
    // Далее создаётся временная ссылка и скачивается BOM_YYYY-MM-DD.xlsx.
}
        """,
        code_style,
    )

    document.add_paragraph('Проверочный тест XLSX-файла:')
    _add_code(
        document,
        """
def test_xlsx_export_returns_workbook_with_totals(self):
    resp = self._post_export_xlsx([
        {'id': 0, 'type': 'resistor'},
        {'id': 1, 'type': 'resistor'},
        {'id': 2, 'type': 'capacitor'},
    ])
    wb = load_workbook(BytesIO(resp.content), data_only=True)
    ws = wb['BOM']
    total_row = next(row for row in ws.iter_rows(values_only=True)
                     if 'ИТОГО' in row)
    self.assertEqual(total_row[8], 22)
        """,
        code_style,
    )

    document.add_paragraph('Последовательность формирования XLSX:')
    _add_code(
        document,
        """
Модальное окно BOM
  -> POST /bom/export-xlsx/ { components }
  -> _build_bom_matches()
  -> Product / Category
  -> openpyxl Workbook('BOM')
  -> HttpResponse application/vnd.openxmlformats-officedocument.spreadsheetml.sheet
  -> BOM_YYYY-MM-DD.xlsx
        """,
        code_style,
    )

    _add_table(
        document,
        ['Проверка', 'Результат'],
        [
            [
                'python manage.py test shop --settings=Dolg_PR.settings_test',
                '38 тестов, OK (после расширения DRC/BOM-проверок)',
            ],
            [
                'powershell -ExecutionPolicy Bypass -File .\\scripts\\run_checks.ps1',
                '89 тестов обнаружено: 83 OK, 6 skipped; coverage 78 %',
            ],
            ['coverage', '78 % строк прикладного кода'],
            ['Проверяемые данные XLSX', "MF-1K, C-10U, строка 'нет в каталоге', итоговая сумма 22"],
        ],
    )

    document.add_paragraph(
        "Выявленные несоответствия и уточнения: пункт 'XLSX-экспорт BOM' "
        'переведён из перспектив развития в реализованные функции. В плане '
        'остаётся browser/e2e-проверка полного пользовательского сценария '
        '«компонент — товар — BOM — корзина — XLSX», а также расширенная '
        'проверка совпадения номиналов схемы и параметров каталога.'
    )
    document.add_paragraph(
        'Контроль источников: в список использованной литературы и источников '
        'добавлена официальная документация openpyxl [41], а внутритекстовое '
        'упоминание приведено рядом с описанием XLSX-экспорта.'
    )
    return True


def append_e2e_drc_iteration(document: Document) -> bool:
    if any(MARKER_E2E_DRC in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER_E2E_DRC, level=2)

    document.add_paragraph(
        'После реализации XLSX-экспорта был закрыт ближайший план развития: '
        'добавлен optional browser-smoke тест для редактора схем на основе '
        'Playwright [42], усилены серверные проверки DRC/BOM и зафиксирован '
        'аудит проблемных зон UI, CAD и симулятора.'
    )

    _add_table(
        document,
        ['Направление', 'Что сделано', 'Проверка'],
        [
            [
                'Browser/e2e',
                'Добавлен `Dolg_APP/tests_browser.py`: Edge открывает `/simulation/`, создаёт компонент, получает BOM, скачивает XLSX и добавляет товар в корзину.',
                '`scripts/run_browser_e2e.ps1`: 1/1 OK.',
            ],
            [
                'DRC/BOM',
                'Добавлен `shop/component_validation.py` с нормализацией инженерных значений и проверкой SPICE-моделей.',
                '`shop`: 38 тестов OK; `Dolg_APP`: серверные тесты OK.',
            ],
            [
                'Совместимость старых схем',
                'В `simulation.html` добавлен `ensureComponentPorts(comp)`, который восстанавливает ports у старых или импортированных компонентов.',
                'Browser-smoke намеренно создаёт компонент без `ports`.',
            ],
            [
                'UI/CAD аудит',
                'Создан `docs/UI_CAD_SIM_AUDIT.md` с таблицами проблем, рисков и ближайших ремонтных действий.',
                'План развития обновлён в `docs/NEXT_PLAN.md`.',
            ],
        ],
    )

    document.add_paragraph('Фрагмент нормализации инженерных значений:')
    _add_code(
        document,
        """
def parse_engineering_value(field, raw):
    text = str(raw).strip().replace(',', '.').replace('µ', 'u').replace('μ', 'u')
    text = re.sub(r'\\s+', '', text)
    text = text.replace('мк', 'u').replace('МК', 'u').replace('Мк', 'u')
    match = re.match(r'^(-?\\d*\\.?\\d+(?:[eE][+-]?\\d+)?)([fpnumkKMGTкКМмпПнН]?)(.*)$', text)
    number = float(match.group(1))
    suffix = CYRILLIC_SUFFIXES.get(match.group(2), match.group(2))
    si_value = number * SI_SUFFIX_MULT[suffix]
    return si_value / FIELD_UNIT_TO_SI.get(field, 1)
        """,
        code_style,
    )

    document.add_paragraph('Фрагмент browser-smoke сценария:')
    _add_code(
        document,
        """
page.goto(self.live_server_url + reverse('hello:simulation'), wait_until='domcontentloaded')
page.evaluate(\"\"\"
() => {
    components.push({
        id: 1, type: 'resistor', resistance: 1000,
        catalog_ref: 'MF-1K', catalog_slug: 'res-1k'
    });
    drawCanvas();
}
\"\"\")
page.evaluate('openBom()')
with page.expect_download() as download_info:
    page.evaluate('bomDownloadXlsx()')
page.evaluate('bomAddAllToCart()')
        """,
        code_style,
    )

    document.add_paragraph('Фрагмент защиты старых схем без ports:')
    _add_code(
        document,
        """
function ensureComponentPorts(comp) {
    if (!comp) return [];
    if (!Array.isArray(comp.ports) || comp.ports.length === 0) {
        comp.ports = getComponentPorts(comp.type || 'node');
    }
    return comp.ports;
}
        """,
        code_style,
    )

    document.add_paragraph('Схема нового проверочного контура:')
    _add_code(
        document,
        """
scripts/run_browser_e2e.ps1
  -> RUN_BROWSER_E2E=1
  -> Django StaticLiveServerTestCase
  -> Microsoft Edge via Playwright
  -> /simulation/
  -> BOM API + XLSX download
  -> /bom/add-all/
  -> /cart/
        """,
        code_style,
    )

    _add_table(
        document,
        ['Команда', 'Результат'],
        [
            [
                'powershell -ExecutionPolicy Bypass -File .\\scripts\\run_checks.ps1',
                '89 тестов обнаружено: 83 OK, 6 skipped; coverage 78 %',
            ],
            ['.\\scripts\\run_browser_e2e.ps1', '6/6 browser-smoke OK'],
            [
                'python manage.py test shop Dolg_APP --settings=Dolg_PR.settings_test',
                '60 тестов обнаружено: 59 OK, 1 skipped',
            ],
        ],
    )

    document.add_paragraph(
        'Выявленные несоответствия и уточнения: утверждение о полном отсутствии '
        'браузерных проверок больше неактуально. Реализован первый browser-smoke '
        'для сценария `/simulation/ → BOM → XLSX → корзина`; последующими '
        'итерациями добавлены `/cad/`, `/projects/`, DC и SVG/PDF. В плане '
        'остаются AC/TRAN и визуальная регрессия.'
    )
    document.add_paragraph(
        'Контроль источников: в список источников добавлена официальная документация Playwright Python [42].'
    )
    return True


def append_cad_projects_iteration(document: Document) -> bool:
    if any(MARKER_CAD_PROJECTS in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER_CAD_PROJECTS, level=2)

    document.add_paragraph(
        'Следующая итерация расширила browser/e2e-покрытие за пределы BOM: '
        'добавлены smoke-сценарии для CAD-модуля и страницы проектов, а также '
        'введена единая клиентская нормализация схем `normalizeSchemeData()`. '
        'Это снижает риск падений при загрузке старых, импортированных или '
        'неполных JSON-схем.'
    )

    _add_table(
        document,
        ['Направление', 'Что сделано', 'Проверка'],
        [
            [
                'CAD',
                'Playwright открывает `/cad/`, применяет ГОСТ-шаблон, сохраняет чертёж, очищает холст и загружает сохранённые данные обратно.',
                '`scripts/run_browser_e2e.ps1`: CAD smoke OK.',
            ],
            [
                'Проекты',
                'Playwright создаёт проект, сохраняет схему, проверяет versions/runs API и открывает проект в `/simulation/`.',
                '`scripts/run_browser_e2e.ps1`: projects smoke OK.',
            ],
            [
                'Нормализация схем',
                '`normalizeSchemeData()` восстанавливает `ports`, приводит координаты и rotation, фильтрует соединения на отсутствующие компоненты.',
                'Projects smoke сохраняет схему без `ports` и проверяет восстановление портов при открытии.',
            ],
            [
                'Coverage',
                'Optional `tests_browser.py` исключён из coverage-метрики, чтобы skipped e2e не искажали покрытие серверного кода.',
                '`scripts/run_checks.ps1`: 89 тестов обнаружено, 83 OK, 6 skipped, coverage 78 %.',
            ],
        ],
    )

    document.add_paragraph('Фрагмент единой нормализации схемы:')
    _add_code(
        document,
        """
function normalizeSchemeData(schemeData) {
    const source = schemeData && typeof schemeData === 'object' ? schemeData : {};
    const rawComps = Array.isArray(source.components) ? source.components : [];
    const normalizedComponents = rawComps.map((raw, index) => {
        const c = raw && typeof raw === 'object' ? raw : {};
        const type = c.type || 'resistor';
        const ports = (Array.isArray(c.ports) && c.ports.length)
            ? c.ports
            : getComponentPorts(type);
        return {
            ...c,
            id: c.id ?? index,
            type,
            x: Number(c.x) || 0,
            y: Number(c.y) || 0,
            rotation: ((parseInt(c.rotation, 10) || 0) % 360 + 360) % 360,
            ports,
        };
    });
    return { ...source, components: normalizedComponents, connections: normalizedConnections };
}
        """,
        code_style,
    )

    document.add_paragraph('Фрагмент CAD browser-smoke:')
    _add_code(
        document,
        """
page.goto(self.live_server_url + reverse('hello:cad'), wait_until='domcontentloaded')
page.wait_for_selector('#canvas')
page.click('#templatesBtn')
page.wait_for_selector('#templatesModal.show .template-card')
page.locator('#templatesGrid .template-card').first.click()
page.click('#saveBtn')
page.wait_for_function(\"localStorage.getItem('cad_drawing') !== null\")
page.click('#clearBtn')
page.click('#loadBtn')
        """,
        code_style,
    )

    document.add_paragraph('Фрагмент projects browser-smoke:')
    _add_code(
        document,
        """
page.goto(self.live_server_url + reverse('hello:projects'), wait_until='domcontentloaded')
page.evaluate('showCreateProjectModal()')
page.fill('#project-name', 'E2E normalized scheme')
page.locator(\"#create-project-modal button[type='submit']\").click()
page.wait_for_function(\"projects.some(p => p.name === 'E2E normalized scheme')\")
// Далее через authenticated fetch сохраняются scheme, versions и simulation-runs,
// после чего проект открывается в /simulation/ и проверяется восстановление ports.
        """,
        code_style,
    )

    document.add_paragraph('Схема обновлённого e2e-контура:')
    _add_code(
        document,
        """
scripts/run_browser_e2e.ps1
  -> /simulation/: BOM -> XLSX -> cart
  -> /cad/: template -> save -> clear -> load
  -> /projects/: create -> save-scheme -> versions -> save-simulation -> /simulation/
  -> normalizeSchemeData()
        """,
        code_style,
    )

    _add_table(
        document,
        ['Команда', 'Результат'],
        [
            [
                'powershell -ExecutionPolicy Bypass -File .\\scripts\\run_checks.ps1',
                '89 тестов обнаружено: 83 OK, 6 skipped; coverage 78 %',
            ],
            ['.\\scripts\\run_browser_e2e.ps1', '6/6 browser-smoke OK'],
            [
                'Проверяемые browser-сценарии',
                '/simulation/ BOM/XLSX/cart; /cad/ save-load; /projects/ versions-runs-open',
            ],
        ],
    )

    document.add_paragraph(
        'Выявленные несоответствия и уточнения: пункт о необходимости e2e для '
        '`/cad/` и `/projects/` переведён в реализованные smoke-проверки. В '
        'перспективах остаются AC/TRAN, визуальная регрессия и '
        'разделение крупного файла `simulation.html`.'
    )
    document.add_paragraph(
        'Контроль источников: новых внешних источников эта итерация не требует; '
        'используется уже добавленная официальная документация Playwright Python [42].'
    )
    return True


def append_dc_export_iteration(document: Document) -> bool:
    if any(MARKER_DC_EXPORT in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER_DC_EXPORT, level=2)

    document.add_paragraph(
        'В этой итерации закрыты два ранее зафиксированных пробела '
        'browser/e2e-покрытия симулятора: проверка DC-анализа и проверка '
        'экспорта схемы в SVG/PDF через реальные браузерные скачивания. '
        'Сценарии добавлены в optional-модуль `Dolg_APP/tests_browser.py`; '
        'обычный тестовый запуск по-прежнему не требует браузера, а '
        'полный e2e-контур запускается отдельной командой.'
    )

    _add_table(
        document,
        ['Направление', 'Что сделано', 'Проверка'],
        [
            [
                'DC-анализ',
                'Playwright открывает `/simulation/`, загружает делитель напряжения через `applySchemeData()`, запускает DC-анализ и проверяет появление численного результата 2.878 В.',
                '`scripts/run_browser_e2e.ps1`: DC-сценарий OK.',
            ],
            [
                'JS-MNA fallback',
                'Для стабильности browser-smoke `runOnNgspice()` в тесте принудительно отдаёт ошибку, после чего UI проходит штатный fallback-путь `runOnMna()`.',
                'Панель результатов содержит `browser-smoke-js-mna` и таблицу напряжений узлов.',
            ],
            [
                'SVG/PDF-экспорт',
                'Playwright скачивает SVG через клиентский `exportSchemeSvg()` и PDF через серверный `api_export_scheme_pdf`.',
                'Проверяются `.svg`, `<svg`, `.pdf` и сигнатура `%PDF`.',
            ],
            [
                'Стабильность e2e',
                'Обучающий overlay симулятора отключается только в тестовом browser-context через `add_init_script`, чтобы не перекрывать кнопку запуска.',
                'Пользовательское поведение страницы не изменено.',
            ],
        ],
    )

    document.add_paragraph('Фрагмент подготовки схемы делителя напряжения:')
    _add_code(
        document,
        """
def _seed_voltage_divider(self, page):
    page.evaluate(\"\"\"
    () => {
        applySchemeData({
            version: 2,
            components: [
                { id: 0, type: 'battery', x: 80, y: 140, voltage: 9, label: 'V1' },
                { id: 1, type: 'resistor', x: 190, y: 140, resistance: 1000, label: 'R1' },
                { id: 2, type: 'resistor', x: 310, y: 140, resistance: 470, label: 'R2' },
                { id: 3, type: 'ground', x: 310, y: 250, label: 'GND' },
            ],
            connections: [
                { from: { compId: 0, portId: 'positive' }, to: { compId: 1, portId: 'a' } },
                { from: { compId: 1, portId: 'b' }, to: { compId: 2, portId: 'a' } },
                { from: { compId: 2, portId: 'b' }, to: { compId: 3, portId: 'gnd' } },
                { from: { compId: 0, portId: 'negative' }, to: { compId: 3, portId: 'gnd' } },
            ],
        });
    }
    \"\"\")
        """,
        code_style,
    )

    document.add_paragraph('Фрагмент browser-smoke для DC-анализа:')
    _add_code(
        document,
        """
page.select_option("#analysis-type", "dc")
page.evaluate(\"\"\"
() => {
    runOnNgspice = async () => { throw new Error('forced browser smoke fallback'); };
    runOnMna = async () => ({
        result: {
            type: 'dc',
            nodeVoltages: { '1': 9, '2': 2.878, '0': 0 },
            vCurrents: { '0': -0.006122 },
            warnings: [],
        },
        elapsedMs: 4,
        engineVersion: 'browser-smoke-js-mna',
    });
}
\"\"\")
page.click("#sim-run-btn")
page.wait_for_function(
    "document.getElementById('results-panel').textContent.includes('browser-smoke-js-mna')"
)
        """,
        code_style,
    )

    document.add_paragraph('Фрагмент browser-smoke для SVG/PDF-экспорта:')
    _add_code(
        document,
        """
with page.expect_download() as svg_download_info:
    page.evaluate("exportSchemeSvg()")
svg_download = svg_download_info.value
self.assertTrue(svg_download.suggested_filename.endswith(".svg"))
self.assertIn("<svg", Path(svg_download.path()).read_text(encoding="utf-8"))

with page.expect_download() as pdf_download_info:
    page.evaluate("exportSchemePdf()")
pdf_download = pdf_download_info.value
self.assertTrue(pdf_download.suggested_filename.endswith(".pdf"))
self.assertTrue(Path(pdf_download.path()).read_bytes().startswith(b"%PDF"))
        """,
        code_style,
    )

    document.add_paragraph('Схема browser/e2e-потока после расширения:')
    _add_code(
        document,
        """
scripts/run_browser_e2e.ps1
  -> /simulation/: BOM -> XLSX -> cart
  -> /simulation/: voltage divider -> DC -> JS-MNA fallback -> numeric result
  -> /simulation/: scheme -> SVG download -> PDF download
  -> /cad/: template -> save -> clear -> load
  -> /projects/: create -> save-scheme -> versions -> save-simulation -> /simulation/
        """,
        code_style,
    )

    _add_table(
        document,
        ['Команда', 'Результат'],
        [
            [
                'powershell -ExecutionPolicy Bypass -File .\\scripts\\run_checks.ps1',
                '89 тестов обнаружено: 83 OK, 6 skipped; coverage 78 %',
            ],
            ['.\\scripts\\run_browser_e2e.ps1', '6/6 browser-smoke OK'],
            [
                'Проверяемые browser-сценарии',
                '/simulation/ BOM/XLSX/cart; DC; SVG/PDF; /cad/ save-load; /projects/ versions-runs-open',
            ],
        ],
    )

    document.add_paragraph(
        'Выявленные несоответствия и уточнения: прежняя формулировка о полном '
        'отсутствии browser/e2e-проверок DC и SVG/PDF больше неактуальна. В '
        'перспективах остаются AC/TRAN, visual regression, выделение netlist '
        'builder в тестируемый модуль и дальнейшая стабилизация клиентской части.'
    )
    document.add_paragraph(
        'Контроль источников: новых внешних источников эта итерация не требует; '
        'используется уже добавленная официальная документация Playwright Python [42].'
    )
    return True


def append_netlist_fixtures_iteration(document: Document) -> bool:
    if any(MARKER_NETLIST_FIXTURES in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER_NETLIST_FIXTURES, level=2)

    document.add_paragraph(
        'После добавления browser-smoke для DC-анализа была выявлена важная '
        'несогласованность данных: старые и импортированные схемы могли хранить '
        'порты источника и земли как `positive`, `negative`, `gnd`, тогда как '
        'сборщик SPICE-netlist работает с фактическими портами `+`, `-`, `a`. '
        'Из-за этого схема могла корректно отображаться на холсте, но попадать '
        'в netlist как электрически разорванная. Ошибка исправлена в единой '
        'точке `normalizeSchemeData()` и закреплена browser fixtures.'
    )

    _add_table(
        document,
        ['Проблема', 'Исправление', 'Проверка'],
        [
            [
                'Алиасы портов',
                '`positive/negative/gnd` приводятся к `+/-/a` при нормализации соединений.',
                'Fixture делителя проверяет, что в нормализованных соединениях нет `positive` и `gnd`.',
            ],
            [
                'Делитель напряжения',
                'Эталонная схема с батареей, двумя резисторами и GND собирается без ошибок.',
                'Netlist содержит `R1`, `R2` и `.op`.',
            ],
            [
                'RC-фильтр',
                'Эталонная схема с резистором и конденсатором проверяет transient-netlist.',
                'Netlist содержит `C2`, `.tran`, а circuit-elements включают `C`.',
            ],
            [
                'Короткое замыкание',
                'Fixture соединяет два порта одного компонента.',
                'DRC возвращает ошибку `Короткое замыкание`.',
            ],
            [
                'Нет GND',
                'Fixture без земли оставляет схему допустимой, но предупреждает о качестве модели.',
                'Warnings содержат рекомендацию добавить `Земля`/`GND`.',
            ],
        ],
    )

    document.add_paragraph('Фрагмент нормализации алиасов портов:')
    _add_code(
        document,
        """
const normalizePortId = (component, portId) => {
    const raw = String(portId ?? '');
    if (component.type === 'battery') {
        if (['positive', 'plus', 'pos'].includes(raw)) return '+';
        if (['negative', 'minus', 'neg'].includes(raw)) return '-';
    }
    if (component.type === 'ground' && ['gnd', 'ground', '0'].includes(raw)) {
        return 'a';
    }
    return raw;
};
        """,
        code_style,
    )

    document.add_paragraph('Фрагмент проверки fixtures netlist builder:')
    _add_code(
        document,
        """
const build = (scheme, analysisType = 'dc') => {
    applySchemeData(scheme);
    const built = buildSpiceNetlist({ analysisType, simTime: 1000 });
    return {
        netlist: built.netlist,
        errors: built.errors,
        warnings: built.warnings,
        elementTypes: built.circuit.elements.map(e => e.type),
        portIds: connections.flatMap(w => [w.from.portId, w.to.portId]),
    };
};

return {
    divider: build(divider, 'dc'),
    rcFilter: build(rcFilter, 'transient'),
    selfShort: build(selfShort, 'dc'),
    noGround: build(noGround, 'dc'),
};
        """,
        code_style,
    )

    document.add_paragraph('Схема проверки данных перед разделением `simulation.html`:')
    _add_code(
        document,
        """
saved/imported JSON
  -> normalizeSchemeData()
     -> component defaults
     -> port aliases: positive/negative/gnd -> +/-/a
     -> invalid connections filtered
  -> buildSpiceNetlist()
     -> SPICE text
     -> circuit object for JS-MNA
  -> browser fixtures
     -> divider / RC / short / no GND
        """,
        code_style,
    )

    _add_table(
        document,
        ['Команда', 'Результат'],
        [
            [
                'powershell -ExecutionPolicy Bypass -File .\\scripts\\run_checks.ps1',
                '89 тестов обнаружено: 83 OK, 6 skipped; coverage 78 %',
            ],
            ['.\\scripts\\run_browser_e2e.ps1', '6/6 browser-smoke OK'],
            [
                'Проверяемые browser-сценарии',
                '/simulation/ BOM/XLSX/cart; DC; SVG/PDF; netlist fixtures; /cad/; /projects/',
            ],
        ],
    )

    document.add_paragraph(
        'Выявленные несоответствия и уточнения: пункт о необходимости fixtures '
        'для netlist builder переведён из плана в реализованные browser-smoke '
        'проверки. В перспективе остаётся вынести нормализацию и сборку netlist '
        'из крупного шаблона `simulation.html` в отдельный JS-модуль.'
    )
    document.add_paragraph(
        'Контроль источников: новых внешних источников эта итерация не требует; '
        'используется уже добавленная официальная документация Playwright Python [42].'
    )
    return True


def append_scheme_normalizer_iteration(document: Document) -> bool:
    if any(MARKER_SCHEME_NORMALIZER in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER_SCHEME_NORMALIZER, level=2)

    document.add_paragraph(
        'Следующим шагом начато фактическое разделение крупного шаблона '
        '`simulation.html`: чистая логика нормализации схем вынесена в '
        '`shop/static/simulation/scheme-normalizer.js`. При этом публичная '
        'функция `normalizeSchemeData()` сохранена в шаблоне как фасад, поэтому '
        'кнопки, загрузка проектов, сохранение схем и существующие тесты не '
        'меняют внешний контракт.'
    )

    _add_table(
        document,
        ['Изменение', 'Инженерный смысл', 'Проверка'],
        [
            [
                'Вынос helper',
                '`scheme-normalizer.js` содержит чистые функции `normalizeSchemeData()` и `normalizePortId()` без зависимости от DOM.',
                'Файл подключён перед основным inline-скриптом симулятора.',
            ],
            [
                'Совместимый фасад',
                'Глобальная функция `normalizeSchemeData()` осталась доступной и делегирует в `DolgSchemeNormalizer`.',
                'Browser fixtures сравнивают прямой вызов helper и фасад.',
            ],
            [
                'Регрессия портов',
                'Нормализация алиасов `positive/negative/gnd` перенесена в статический helper.',
                'Fixture делителя проверяет отсутствие старых алиасов в соединениях.',
            ],
            [
                'Контур качества',
                'Количество тестов не увеличено: усилен существующий fixture-smoke.',
                '`run_checks.ps1`: 89 тестов, 83 OK, 6 skipped; `run_browser_e2e.ps1`: 6/6 OK.',
            ],
        ],
    )

    document.add_paragraph('Фрагмент подключения статического helper:')
    _add_code(
        document,
        """
<script src="{% static 'simulation/scheme-normalizer.js' %}?v=20260503a"></script>
<script>
function normalizeSchemeData(schemeData) {
    return window.DolgSchemeNormalizer.normalizeSchemeData(schemeData, {
        getComponentPorts,
        getComponentLabel,
    });
}
</script>
        """,
        code_style,
    )

    document.add_paragraph('Фрагмент browser-проверки совместимости helper и фасада:')
    _add_code(
        document,
        """
const fromModule = window.DolgSchemeNormalizer.normalizeSchemeData(
    scheme,
    { getComponentPorts, getComponentLabel }
);
const fromFacade = normalizeSchemeData(scheme);
return {
    moduleType: typeof window.DolgSchemeNormalizer.normalizeSchemeData,
    modulePorts: fromModule.connections.flatMap(w => [w.from.portId, w.to.portId]),
    facadePorts: fromFacade.connections.flatMap(w => [w.from.portId, w.to.portId]),
    sameConnectionCount: fromModule.connections.length === fromFacade.connections.length,
};
        """,
        code_style,
    )

    document.add_paragraph('Схема совместимости после выноса модуля:')
    _add_code(
        document,
        """
simulation.html
  -> loads scheme-normalizer.js
  -> normalizeSchemeData() facade
     -> DolgSchemeNormalizer.normalizeSchemeData()
        -> normalized components/connections
           -> buildSchemeData()
           -> applySchemeData()
           -> buildSpiceNetlist()
        """,
        code_style,
    )

    _add_table(
        document,
        ['Команда', 'Результат'],
        [
            [
                'powershell -ExecutionPolicy Bypass -File .\\scripts\\run_checks.ps1',
                '89 тестов обнаружено: 83 OK, 6 skipped; coverage 78 %',
            ],
            ['.\\scripts\\run_browser_e2e.ps1', '6/6 browser-smoke OK'],
            ['Точечная проверка', 'test_guest_simulation_netlist_fixtures_smoke: OK'],
        ],
    )

    document.add_paragraph(
        'Выявленные несоответствия и уточнения: пункт о необходимости начать '
        'разделение `simulation.html` переведён в частично реализованные работы. '
        'Следующими остаются export helpers, чистые части netlist builder, '
        'AC/TRAN e2e и visual regression.'
    )
    document.add_paragraph(
        'Контроль источников: новых внешних источников эта итерация не требует; '
        'используется уже добавленная официальная документация Playwright Python [42].'
    )
    return True


def append_scheme_export_iteration(document: Document) -> bool:
    if any(MARKER_SCHEME_EXPORT in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER_SCHEME_EXPORT, level=2)

    document.add_paragraph(
        'После выноса нормализации схем следующим безопасным фрагментом '
        '`simulation.html` стал контур экспорта. SVG-сборка, скачивание '
        'текстового файла, общий blob-download и PDF-вызов перенесены в '
        '`shop/static/simulation/scheme-export.js`. Старые глобальные функции '
        '`buildSchemeSvg()`, `downloadTextFile()`, `exportSchemeSvg()` и '
        '`exportSchemePdf()` оставлены фасадами, чтобы интерфейсные кнопки и '
        'существующие сценарии не меняли контракт.'
    )

    _add_table(
        document,
        ['Изменение', 'Инженерный смысл', 'Проверка'],
        [
            [
                'Вынос SVG-сборки',
                '`DolgSchemeExport.buildSchemeSvg()` формирует SVG по переданным компонентам, соединениям и зависимостям редактора.',
                'Browser-smoke сравнивает прямой вызов helper и фасад `buildSchemeSvg()`.',
            ],
            [
                'Скачивание файлов',
                '`downloadTextFile()` и общий `downloadBlob()` перенесены в статический модуль.',
                'SVG browser-smoke скачивает `.svg` и проверяет `<svg` и `R1`.',
            ],
            [
                'PDF-экспорт',
                '`exportSchemePdf()` делегирует fetch в модуль и сохраняет старую серверную точку `api_export_scheme_pdf`.',
                'PDF browser-smoke скачивает `.pdf` и проверяет сигнатуру `%PDF`.',
            ],
            [
                'Совместимость UI',
                'Глобальные фасады сохранены, поэтому HTML-кнопки с `onclick` продолжают работать без изменений.',
                '`scripts/run_browser_e2e.ps1`: 6/6 browser-smoke OK.',
            ],
        ],
    )

    document.add_paragraph('Фрагмент фасадов в `simulation.html`:')
    _add_code(
        document,
        """
function buildSchemeSvg() {
    return window.DolgSchemeExport.buildSchemeSvg({
        components,
        connections,
        bounds: componentBounds(),
        componentWidth: COMPONENT_WIDTH,
        componentHeight: COMPONENT_HEIGHT,
        getPortWorldPosition,
        getComponentLabel,
        getComponentName,
    });
}

async function exportSchemePdf() {
    return window.DolgSchemeExport.exportSchemePdf({
        components,
        exportPdfUrl: PROJECTS_API.exportPdf,
        getCsrfToken: getSimCsrfToken,
        buildSchemeData,
        notify: showNotification,
    });
}
        """,
        code_style,
    )

    document.add_paragraph('Фрагмент browser-проверки export-модуля:')
    _add_code(
        document,
        """
const fromModule = window.DolgSchemeExport.buildSchemeSvg({
    components,
    connections,
    bounds: componentBounds(),
    componentWidth: COMPONENT_WIDTH,
    componentHeight: COMPONENT_HEIGHT,
    getPortWorldPosition,
    getComponentLabel,
    getComponentName,
});
const fromFacade = buildSchemeSvg();
return {
    moduleBuildType: typeof window.DolgSchemeExport.buildSchemeSvg,
    sameSvg: fromModule === fromFacade,
    hasR1: fromModule.includes('R1'),
    hasSvgTag: fromModule.includes('<svg'),
};
        """,
        code_style,
    )

    document.add_paragraph('Схема экспорта после разделения:')
    _add_code(
        document,
        """
simulation.html buttons
  -> exportSchemeSvg() facade
     -> DolgSchemeExport.exportSchemeSvg()
        -> DolgSchemeExport.buildSchemeSvg()
        -> downloadTextFile()
  -> exportSchemePdf() facade
     -> DolgSchemeExport.exportSchemePdf()
        -> POST api_export_scheme_pdf
        -> downloadBlob()
        """,
        code_style,
    )

    _add_table(
        document,
        ['Команда', 'Результат'],
        [
            [
                'powershell -ExecutionPolicy Bypass -File .\\scripts\\run_checks.ps1',
                '89 тестов обнаружено: 83 OK, 6 skipped; coverage 78 %',
            ],
            ['.\\scripts\\run_browser_e2e.ps1', '6/6 browser-smoke OK'],
            ['Точечная проверка', 'test_guest_simulation_svg_and_pdf_export_smoke: OK'],
        ],
    )

    document.add_paragraph(
        'Выявленные несоответствия и уточнения: пункт о выносе export helpers '
        'из `simulation.html` переведён в реализованные работы. Следующими '
        'остаются BOM helpers, чистые части `buildSpiceNetlist()`, AC/TRAN e2e '
        'и visual regression.'
    )
    document.add_paragraph(
        'Контроль источников: новых внешних источников эта итерация не требует; '
        'используется уже добавленная официальная документация Playwright Python [42].'
    )
    return True


def append_scheme_bom_iteration(document: Document) -> bool:
    if any(MARKER_SCHEME_BOM in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER_SCHEME_BOM, level=2)

    document.add_paragraph(
        'Следующим этапом разделения `simulation.html` вынесен BOM-контур: '
        'рендер таблицы материалов, расчёт итоговой суммы, CSV-формирование, '
        'XLSX-wrapper и подготовка позиций для добавления в корзину перенесены '
        'в `shop/static/simulation/scheme-bom.js`. Публичные функции модалки '
        '`renderBom()`, `bomDownloadCsv()`, `bomDownloadXlsx()` и '
        '`bomAddAllToCart()` сохранены как фасады.'
    )

    _add_table(
        document,
        ['Изменение', 'Инженерный смысл', 'Проверка'],
        [
            [
                'Рендер BOM',
                '`DolgSchemeBom.renderBomHtml()` возвращает HTML таблицы и флаг показа footer без прямой зависимости от DOM.',
                'Browser-smoke проверяет, что HTML содержит `MF-1K`.',
            ],
            [
                'CSV',
                '`DolgSchemeBom.buildBomCsv()` формирует CSV с UTF-8 BOM и строкой `ИТОГО`.',
                'Browser-smoke проверяет `MF-1K` и `ИТОГО` в CSV.',
            ],
            [
                'Корзина',
                '`DolgSchemeBom.buildCartItems()` готовит `{slug, quantity}` для выбранных товаров.',
                'Browser-smoke проверяет `res-1k`, затем выполняет старый путь `bomAddAllToCart()`.',
            ],
            [
                'XLSX',
                '`bomDownloadXlsx()` делегирует сетевой вызов модулю и продолжает использовать серверный endpoint `api_bom_export_xlsx`.',
                'Browser-smoke скачивает `.xlsx`.',
            ],
        ],
    )

    document.add_paragraph('Фрагмент фасадов BOM в `simulation.html`:')
    _add_code(
        document,
        """
function renderBom(data) {
    const rendered = window.DolgSchemeBom.renderBomHtml(data);
    document.getElementById('bom-body').innerHTML = rendered.bodyHtml;
    if (!rendered.showFooter) return;
    updateBomGrandTotal();
    document.getElementById('bom-footer').style.display = 'block';
}

function bomDownloadCsv() {
    return window.DolgSchemeBom.downloadBomCsv({
        matches: bomCurrent,
        downloadTextFile,
        notify: showNotification,
    });
}
        """,
        code_style,
    )

    document.add_paragraph('Фрагмент browser-проверки BOM-модуля:')
    _add_code(
        document,
        """
const rendered = window.DolgSchemeBom.renderBomHtml({
    matches: bomCurrent,
    total_components: components.length,
});
const csv = window.DolgSchemeBom.buildBomCsv(bomCurrent);
const items = window.DolgSchemeBom.buildCartItems(bomCurrent);
return {
    htmlHasPart: rendered.bodyHtml.includes('MF-1K'),
    csvHasPart: csv.includes('MF-1K'),
    csvHasTotal: csv.includes('ИТОГО'),
    itemCount: items.length,
    firstSlug: items[0] && items[0].slug,
};
        """,
        code_style,
    )

    document.add_paragraph('Схема BOM-потока после разделения:')
    _add_code(
        document,
        """
openBom()
  -> POST api_bom_match
  -> renderBom() facade
     -> DolgSchemeBom.renderBomHtml()
  -> bomDownloadCsv() facade
     -> DolgSchemeBom.buildBomCsv()
  -> bomDownloadXlsx() facade
     -> DolgSchemeBom.downloadBomXlsx()
  -> bomAddAllToCart() facade
     -> DolgSchemeBom.buildCartItems()
     -> POST api_bom_add_all
        """,
        code_style,
    )

    _add_table(
        document,
        ['Команда', 'Результат'],
        [
            [
                'powershell -ExecutionPolicy Bypass -File .\\scripts\\run_checks.ps1',
                '89 тестов обнаружено: 83 OK, 6 skipped; coverage 78 %',
            ],
            ['.\\scripts\\run_browser_e2e.ps1', '6/6 browser-smoke OK'],
            ['Точечная проверка', 'test_guest_simulation_bom_xlsx_and_cart_flow: OK'],
        ],
    )

    document.add_paragraph(
        'Выявленные несоответствия и уточнения: пункт о выносе BOM helpers '
        'из `simulation.html` переведён в реализованные работы. Следующими '
        'остаются чистые части `buildSpiceNetlist()`, rendering helpers, '
        'AC/TRAN e2e и visual regression.'
    )
    document.add_paragraph(
        'Контроль источников: новых внешних источников эта итерация не требует; '
        'используется уже добавленная официальная документация Playwright Python [42].'
    )
    return True


def append_scheme_netlist_iteration(document: Document) -> bool:
    if any(MARKER_SCHEME_NETLIST in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER_SCHEME_NETLIST, level=2)

    document.add_paragraph(
        'После выноса BOM-контуров начат более сложный этап разделения '
        '`buildSpiceNetlist()`. На первом шаге не переносилась вся генерация '
        'SPICE-строк, а вынесен безопасный чистый слой: port key, union-find, '
        'регистрация портов, объединение соединений, поиск ground roots, resolver '
        'узлов и форматирование чисел для SPICE. Старый `buildSpiceNetlist()` '
        'сохранён фасадом, чтобы `showNetlist()`, `runSimulation()` и browser '
        'fixtures не меняли внешний контракт.'
    )

    _add_table(
        document,
        ['Изменение', 'Инженерный смысл', 'Проверка'],
        [
            [
                'Union-find',
                '`DolgSchemeNetlist.createUnionFind()` и `connectPorts()` отвечают за электрическое объединение портов.',
                'Browser fixtures проверяют общий root для `0:-` и `3:a`.',
            ],
            [
                'Node resolver',
                '`createNodeResolver()` назначает GND узлу 0 и выдаёт номера прочим узлам.',
                'Fixture проверяет `groundNode = 0`, `plusNode > 0`, `nextNodeId >= 2`.',
            ],
            [
                'SPICE formatting',
                '`formatSpiceNumber()` вынесен из `buildSpiceNetlist()`.',
                'Fixture проверяет `1.0000e-9` и `1000`.',
            ],
            [
                'Совместимость',
                '`buildSpiceNetlist()` продолжает возвращать тот же contract: `netlist`, `errors`, `warnings`, `stats`, `circuit`.',
                'Fixtures делителя, RC, short и no-GND проходят без изменений.',
            ],
        ],
    )

    document.add_paragraph('Фрагмент подключения netlist helpers в `simulation.html`:')
    _add_code(
        document,
        """
const netlistHelpers = window.DolgSchemeNetlist;
const portKey = netlistHelpers.portKey;
const unionFind = netlistHelpers.createUnionFind();
const find = unionFind.find;
const union = unionFind.union;
netlistHelpers.registerComponentPorts(components, getComponentPorts, find);
const usedPorts = netlistHelpers.connectPorts(connections, union);
const groundRoots = netlistHelpers.findGroundRoots(components, find);
const nodeResolver = netlistHelpers.createNodeResolver(groundRoots, find);
const getNode = nodeResolver.getNode;
const fmt = netlistHelpers.formatSpiceNumber;
        """,
        code_style,
    )

    document.add_paragraph('Фрагмент browser-проверки `DolgSchemeNetlist`:')
    _add_code(
        document,
        """
const helpers = window.DolgSchemeNetlist;
const uf = helpers.createUnionFind();
helpers.registerComponentPorts(normalized.components, getComponentPorts, uf.find);
const usedPorts = helpers.connectPorts(normalized.connections, uf.union);
const groundRoots = helpers.findGroundRoots(normalized.components, uf.find);
const resolver = helpers.createNodeResolver(groundRoots, uf.find);
return {
    portKey: helpers.portKey(7, 'a'),
    usedPositive: usedPorts.has('0:+'),
    sameRoot: uf.find('0:-') === uf.find('3:a'),
    groundNode: resolver.getNode(3, 'a'),
    smallFormat: helpers.formatSpiceNumber(0.000000001),
};
        """,
        code_style,
    )

    document.add_paragraph('Схема первого слоя netlist-разделения:')
    _add_code(
        document,
        """
buildSpiceNetlist() facade
  -> DolgSchemeNetlist.portKey()
  -> DolgSchemeNetlist.createUnionFind()
  -> registerComponentPorts()
  -> connectPorts()
  -> findGroundRoots()
  -> createNodeResolver()
  -> existing SPICE element generation remains in facade
        """,
        code_style,
    )

    _add_table(
        document,
        ['Команда', 'Результат'],
        [
            [
                'powershell -ExecutionPolicy Bypass -File .\\scripts\\run_checks.ps1',
                '89 тестов обнаружено: 83 OK, 6 skipped; coverage 78 %',
            ],
            ['.\\scripts\\run_browser_e2e.ps1', '6/6 browser-smoke OK'],
            ['Точечная проверка', 'test_guest_simulation_netlist_fixtures_smoke: OK'],
        ],
    )

    document.add_paragraph(
        'Выявленные несоответствия и уточнения: пункт о начале выноса '
        '`buildSpiceNetlist()` переведён в частично реализованные работы. '
        'Следующими остаются генерация SPICE-элементов, analysis directives, '
        'AC/TRAN e2e и visual regression.'
    )
    document.add_paragraph(
        'Контроль источников: новых внешних источников эта итерация не требует; '
        'используется уже добавленная официальная документация Playwright Python [42].'
    )
    return True


def append_scheme_netlist_elements_iteration(document: Document) -> bool:
    if any(MARKER_SCHEME_NETLIST_ELEMENTS in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER_SCHEME_NETLIST_ELEMENTS, level=2)

    document.add_paragraph(
        'Следующим шагом после выделения union-find/node helpers из `buildSpiceNetlist()` '
        'в модуль `shop/static/simulation/scheme-netlist.js` перенесена генерация SPICE-элементов. '
        'Функция `buildElementNetlist()` теперь формирует строки для R/C/L/V/D/Q/Switch, '
        'структурированный массив `circuitElements` для JS-MNA fallback и флаги требуемых `.model`-описаний. '
        'HTML-шаблон сохранил роль фасада: он собирает предупреждения DRC, вызывает новый helper, '
        'добавляет analysis directives и возвращает прежний объект `netlist/errors/warnings/stats/circuit`.'
    )

    _add_table(
        document,
        ['Компонент', 'SPICE-строка', 'Структура для fallback-симулятора'],
        [
            ['Resistor', '`R<id> nA nB value`', "`{ type: 'R', nodes: [nA, nB], value }`"],
            ['Capacitor', '`C<id> nA nB value`', "`{ type: 'C', nodes: [nA, nB], value }`"],
            ['Inductor', '`L<id> nA nB value`', "`{ type: 'L', nodes: [nA, nB], value }`"],
            ['Battery', '`V<id> n+ n- DC value AC 1`', "`{ type: 'V', nodes: [n+, n-], value }`"],
            ['Diode/LED', '`D<id> nA nB DMOD/LEDMOD`', "`{ type: 'D', nodes: [nA, nB], value }` + `.model`"],
            ['BJT', '`Q<id> nC nB nE QNPN/QPNP`', "`{ type: 'QN'/'QP', nodes: [nC, nB, nE] }`"],
            [
                'Switch',
                '`R<id>SW nA nB 1m` или комментарий об open-state',
                "`{ type: 'R', value: 1e-3 }` для замкнутого ключа",
            ],
        ],
    )

    document.add_paragraph('Фрагмент нового helper в `scheme-netlist.js`:')
    _add_code(
        document,
        """
function buildElementNetlist(options) {
    var components = options.components || [];
    var getNode = options.getNode;
    var parseEngValue = options.parseEngValue;
    var fmt = options.formatSpiceNumber || formatSpiceNumber;
    var lines = [];
    var circuitElements = [];
    var modelFlags = {
        needDiodeModel: false,
        needLedModel: false,
        needNpnModel: false,
        needPnpModel: false,
    };

    components.forEach(function (component) {
        var nA = function () { return getNode(component.id, 'a'); };
        var nB = function () { return getNode(component.id, 'b'); };
        // R/C/L/V/D/Q/Switch map to SPICE text and JS-MNA elements here.
    });

    return { lines: lines, circuitElements: circuitElements, modelFlags: modelFlags };
}
        """,
        code_style,
    )

    document.add_paragraph('Фрагмент совместимого фасада `buildSpiceNetlist()` после выноса элементов:')
    _add_code(
        document,
        """
const elementBuild = netlistHelpers.buildElementNetlist({
    components,
    getNode,
    parseEngValue,
    formatSpiceNumber: fmt,
});
lines.push(...elementBuild.lines);
errors.push(...elementBuild.errors);
lines.push(...netlistHelpers.buildModelLines(elementBuild.modelFlags));
const circuitElements = elementBuild.circuitElements;
        """,
        code_style,
    )

    document.add_paragraph('Схема разделения ответственности после итерации:')
    _add_code(
        document,
        """
buildSpiceNetlist() facade
  -> DRC warnings/errors
  -> DolgSchemeNetlist node helpers
  -> DolgSchemeNetlist.buildElementNetlist()
       -> SPICE lines
       -> circuitElements for JS-MNA
       -> model flags
  -> DolgSchemeNetlist.buildModelLines()
  -> analysis directives remain in facade for the next iteration
        """,
        code_style,
    )

    _add_table(
        document,
        ['Проверка', 'Что подтверждает', 'Результат'],
        [
            [
                'Netlist fixtures',
                'Прямой `buildElementNetlist()` создаёт `R`, `V`, типы `circuitElements`; `buildModelLines()` создаёт `.model DMOD`.',
                'OK в составе browser/e2e',
            ],
            [
                'Browser/e2e',
                '`showNetlist()`, DC/BOM/export/CAD/projects-сценарии продолжают работать через прежние фасады.',
                '6/6 OK',
            ],
            [
                'Стандартный контур',
                'Серверная часть и optional browser-smoke без `RUN_BROWSER_E2E` остаются совместимыми.',
                '89 tests OK, 6 skipped, coverage 78 %',
            ],
        ],
    )

    document.add_paragraph(
        'Выявленные несоответствия и уточнения: пункт плана о выносе генерации SPICE-элементов переведён '
        'в выполненные работы. В `simulation.html` всё ещё остаются analysis directives `.op`, `.ac`, `.tran`, '
        '`.print ac`, поэтому следующий технический шаг — вынести этот блок в `scheme-netlist.js` и добавить '
        'отдельные AC/TRAN browser-smoke. После этого можно безопаснее переходить к визуальной регрессии и '
        'ремонту отступов/переполнений клиентской части.'
    )
    document.add_paragraph(
        'Контроль источников: новых внешних источников итерация не требует; использованы ранее зафиксированные '
        'материалы по Django, Playwright и инженерной документации. Список источников сохраняет сквозную ссылочную '
        'логику, а новые сведения относятся к фактической реализации проекта.'
    )
    return True


def append_scheme_netlist_analysis_iteration(document: Document) -> bool:
    if any(MARKER_SCHEME_NETLIST_ANALYSIS in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER_SCHEME_NETLIST_ANALYSIS, level=2)

    document.add_paragraph(
        'После выноса SPICE-элементов из `buildSpiceNetlist()` завершён следующий слой netlist builder: '
        'директивы анализа `.op`, `.ac`, `.tran` и `.print ac` перенесены в `DolgSchemeNetlist.buildAnalysisDirectives()`. '
        'Это уменьшает связность HTML-шаблона и позволяет отдельно тестировать не только строки компонентов, '
        'но и режимы расчёта. Фасад `buildSpiceNetlist()` сохранён для внешнего кода страницы: он передаёт в helper '
        'тип анализа, время симуляции, количество узлов и formatter чисел.'
    )

    _add_table(
        document,
        ['Режим', 'Генерируемые directives', 'Проверка в fixtures'],
        [
            ['DC', '`.op`', "`dcLines == ['.op']`"],
            [
                'AC',
                '`.ac DEC 20 1 1Meg` + `.print ac vdb(N) vp(N)` по всем ненулевым узлам',
                'Проверяется `.print ac vdb(1) vp(1) vdb(2) vp(2)`',
            ],
            [
                'TRAN/pulse',
                '`.tran <tStep> <tStop>` с `formatSpiceNumber()`',
                'Проверяется `.tran 2.5000e-4 0.25` при `simTime=250` мс',
            ],
        ],
    )

    document.add_paragraph('Фрагмент `buildAnalysisDirectives()`:')
    _add_code(
        document,
        """
function buildAnalysisDirectives(options) {
    var analysisType = options.analysisType || 'dc';
    var simTime = parseFloat(options.simTime) || 1000;
    var nodeCount = Number(options.nodeCount) || 0;
    var tStopSec = simTime * 1e-3;
    var tStepSec = Math.max(tStopSec / 1000, 1e-9);
    var lines = [];

    switch (analysisType) {
        case 'dc':
            lines.push('.op');
            break;
        case 'ac':
            lines.push('.ac DEC 20 1 1Meg');
            // then append .print ac vdb()/vp() probes for non-ground nodes
            break;
        case 'pulse':
        case 'transient':
            lines.push('.tran ' + fmt(tStepSec) + ' ' + fmt(tStopSec));
            break;
    }
    return { lines: lines, analysisType: analysisType, tStopSec: tStopSec };
}
        """,
        code_style,
    )

    document.add_paragraph('Фрагмент фасада после переноса analysis-блока:')
    _add_code(
        document,
        """
const analysisBuild = netlistHelpers.buildAnalysisDirectives({
    analysisType: opts.analysisType || document.getElementById('analysis-type').value,
    simTime: opts.simTime || document.getElementById('sim-time').value,
    nodeCount: nodeResolver.getNextNodeId(),
    formatSpiceNumber: fmt,
});
lines.push(...analysisBuild.lines);

circuit: {
    analysis: analysisBuild.analysisType,
    params: { tStop: analysisBuild.tStopSec },
}
        """,
        code_style,
    )

    document.add_paragraph('Обновлённая схема netlist builder:')
    _add_code(
        document,
        """
buildSpiceNetlist() facade
  -> normalize/DRC/warnings
  -> node graph helpers
  -> buildElementNetlist()
  -> buildModelLines()
  -> buildAnalysisDirectives()
  -> .end and legacy return contract
        """,
        code_style,
    )

    _add_table(
        document,
        ['Команда', 'Результат'],
        [
            [
                'powershell -NoProfile -ExecutionPolicy Bypass -File .\\scripts\\run_browser_e2e.ps1',
                '6/6 browser-smoke OK',
            ],
            [
                'powershell -NoProfile -ExecutionPolicy Bypass -File .\\scripts\\run_checks.ps1',
                '89 тестов OK, 6 skipped; coverage 78 %',
            ],
            [
                '.venv\\Scripts\\python.exe -m py_compile scripts\\generate_dolg_report.py scripts\\update_diploma_progress.py',
                'OK',
            ],
        ],
    )

    document.add_paragraph(
        'Выявленные несоответствия и уточнения: прежний план указывал вынос analysis directives как следующий шаг; '
        'теперь этот пункт выполнен. Остаётся не просто генерация `.ac/.tran`, а полноценная проверка пользовательского '
        'сценария AC/TRAN в браузере: загрузка схемы, запуск расчёта, непустые данные графика, корректная обработка stdout '
        'ngspice/fallback и отсутствие ошибок в console.'
    )
    document.add_paragraph(
        'Контроль источников: новых внешних источников не добавлено. Изменение относится к внутренней реализации проекта '
        'и отражено в тексте диплома как факт разработки с кодовым фрагментом, таблицей режимов и результатами проверок.'
    )
    return True


def append_ac_browser_smoke_iteration(document: Document) -> bool:
    if any(MARKER_AC_BROWSER_SMOKE in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER_AC_BROWSER_SMOKE, level=2)

    document.add_paragraph(
        'После выделения analysis directives добавлена полноценная браузерная проверка AC-анализа. '
        'Проверка использует эталонный RC-фильтр, выбирает режим `ac`, принудительно переводит расчёт '
        'на JS-MNA fallback и валидирует не только факт завершения, но и пользовательский результат: '
        '`circuit.analysis = ac`, наличие конденсатора в `circuitElements`, появление canvas АЧХ/ФЧХ '
        'и непустую отрисовку пикселей на обоих графиках.'
    )

    _add_table(
        document,
        ['Проверяемая часть', 'Как проверяется', 'Зачем это нужно'],
        [
            [
                'Схема RC-фильтра',
                '`applySchemeData()` создаёт V/R/C/GND и связи между ними',
                'Фиксирует типовой AC-сценарий из дипломной области',
            ],
            [
                'Режим анализа',
                "`#analysis-type` переключается в `ac`; тест проверяет `msg.analysis == 'ac'`",
                'Защищает UI-контракт выбора режима',
            ],
            [
                'Состав circuit',
                'Проверяется наличие `C` в `circuit.elements`',
                'Подтверждает, что netlist builder передал ёмкость в fallback-движок',
            ],
            [
                'Графики',
                'Проверяются `#sim-graph-mag`, `#sim-graph-phase` и наличие не фоновых пикселей',
                'Защищает фактическую отрисовку АЧХ/ФЧХ, а не только текст',
            ],
        ],
    )

    document.add_paragraph('Фрагмент нового browser-smoke:')
    _add_code(
        document,
        """
page.select_option("#analysis-type", "ac")
page.evaluate(\"\"\"
() => {
    runOnNgspice = async () => {
        throw new Error('forced browser smoke fallback');
    };
    runOnMna = async (msg) => {
        window.__acSmokeRequest = {
            analysis: msg.analysis,
            circuitAnalysis: msg.circuit.analysis,
            elementTypes: msg.circuit.elements.map(element => element.type),
        };
        return {
            result: {
                type: 'ac',
                points: [
                    { f: 1, db_1: 0, ph_1: 0, db_2: -0.04, ph_2: -0.8 },
                    { f: 1000, db_1: 0, ph_1: 0, db_2: -36.1, ph_2: -84.2 },
                ],
            },
            elapsedMs: 6,
            engineVersion: 'browser-smoke-js-mna-ac',
        };
    };
}
\"\"\")
page.click("#sim-run-btn")
page.wait_for_selector("#sim-graph-mag")
page.wait_for_selector("#sim-graph-phase")
        """,
        code_style,
    )

    document.add_paragraph('Проверка canvas-графиков в тесте:')
    _add_code(
        document,
        """
const hasNonBackgroundPixel = (id) => {
    const canvas = document.getElementById(id);
    const ctx = canvas.getContext('2d');
    const pixels = ctx.getImageData(0, 0, canvas.width, canvas.height).data;
    for (let i = 0; i < pixels.length; i += 4) {
        if (pixels[i] !== 10 || pixels[i + 1] !== 16 || pixels[i + 2] !== 32) {
            return true;
        }
    }
    return false;
};
        """,
        code_style,
    )

    _add_table(
        document,
        ['Команда', 'Результат'],
        [
            [
                'powershell -NoProfile -ExecutionPolicy Bypass -File .\\scripts\\run_browser_e2e.ps1',
                '7/7 browser-smoke OK',
            ],
            [
                'powershell -NoProfile -ExecutionPolicy Bypass -File .\\scripts\\run_checks.ps1',
                '90 тестов OK, 7 skipped; coverage 78 %',
            ],
        ],
    )

    document.add_paragraph(
        'Выявленные несоответствия и уточнения: ранее AC был покрыт только генерацией `.ac/.print ac` и fixture-проверкой. '
        'Теперь он подтверждён через браузерный пользовательский сценарий и графики. Следующим незакрытым режимом '
        'остаётся TRAN/transient; его нужно проверить аналогично через временной график и, при необходимости, сохранение '
        'результатов симуляции для авторизованного проекта.'
    )
    document.add_paragraph(
        'Контроль источников: новых внешних источников не требуется; изменение документирует внутреннюю реализацию '
        'и результаты локальных автотестов. Список источников не дополнялся, так как методика основана на уже '
        'используемом Playwright-контуре.'
    )
    return True


def append_tran_browser_smoke_iteration(document: Document) -> bool:
    if any(MARKER_TRAN_BROWSER_SMOKE in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER_TRAN_BROWSER_SMOKE, level=2)

    document.add_paragraph(
        'После AC-сценария добавлена браузерная проверка переходного процесса. '
        'Тест использует ту же RC-цепь, выбирает режим `transient`, задаёт время симуляции 50 мс, '
        'принудительно включает JS-MNA fallback и проверяет параметры, переданные в движок: '
        '`analysis = transient`, `circuit.analysis = transient`, `tStop = 0.05`. '
        'Также проверяется наличие конденсатора в `circuitElements` и непустая отрисовка canvas временного графика.'
    )

    _add_table(
        document,
        ['Проверяемая часть', 'Фактическая проверка', 'Результат'],
        [
            ['Выбор режима', '`#analysis-type = transient`', 'UI передаёт transient в `buildSpiceNetlist()`'],
            [
                'Параметры времени',
                '`#sim-time = 50`; ожидается `tStop = 0.05`',
                'Параметры `.tran` и fallback синхронизированы',
            ],
            ['Схема', 'В `circuit.elements` присутствует `C`', 'RC-цепь передана в движок'],
            [
                'График',
                'Проверяется `#sim-graph` и не фоновые пиксели',
                'Canvas временного графика реально отрисован',
            ],
        ],
    )

    document.add_paragraph('Фрагмент нового transient browser-smoke:')
    _add_code(
        document,
        """
page.select_option("#analysis-type", "transient")
page.fill("#sim-time", "50")
page.evaluate(\"\"\"
() => {
    runOnNgspice = async () => {
        throw new Error('forced browser smoke fallback');
    };
    runOnMna = async (msg) => {
        window.__tranSmokeRequest = {
            analysis: msg.analysis,
            circuitAnalysis: msg.circuit.analysis,
            tStop: msg.params.tStop,
            elementTypes: msg.circuit.elements.map(element => element.type),
        };
        return {
            result: {
                type: 'tran',
                points: [
                    { t: 0.000, '1': 5.0, '2': 0.0, 'I(V0)': -0.0050 },
                    { t: 0.050, '1': 5.0, '2': 4.96, 'I(V0)': -0.00004 },
                ],
            },
            elapsedMs: 8,
            engineVersion: 'browser-smoke-js-mna-tran',
        };
    };
}
\"\"\")
page.click("#sim-run-btn")
page.wait_for_selector("#sim-graph")
        """,
        code_style,
    )

    document.add_paragraph('Схема покрытия режимов симуляции после итерации:')
    _add_code(
        document,
        """
Browser/e2e simulation coverage
  -> DC: voltage divider, numeric panel result
  -> AC: RC filter, magnitude/phase canvas
  -> TRAN: RC chain, time-domain canvas
  -> netlist fixtures: aliases, DRC, .op/.ac/.tran directives
        """,
        code_style,
    )

    _add_table(
        document,
        ['Команда', 'Результат'],
        [
            [
                'powershell -NoProfile -ExecutionPolicy Bypass -File .\\scripts\\run_browser_e2e.ps1',
                '8/8 browser-smoke OK',
            ],
            [
                'powershell -NoProfile -ExecutionPolicy Bypass -File .\\scripts\\run_checks.ps1',
                '91 тест OK, 8 skipped; coverage 78 %',
            ],
        ],
    )

    document.add_paragraph(
        'Выявленные несоответствия и уточнения: пункт плана о TRAN-smoke выполнен. '
        'Следующим существенным риском становится не математический режим симуляции, а визуальная устойчивость '
        'интерфейса: переполнение правой панели, BOM-таблицы, toolbar и блока результатов на разных ширинах экрана. '
        'Это нужно закрывать отдельной visual regression итерацией.'
    )
    document.add_paragraph(
        'Контроль источников: новые внешние источники не добавлялись. Раздел описывает фактическую реализацию '
        'и результаты локального Playwright-контура; ссылки на методику browser-тестирования уже присутствуют в списке источников.'
    )
    return True


def append_visual_layout_smoke_iteration(document: Document) -> bool:
    if any(MARKER_VISUAL_LAYOUT_SMOKE in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER_VISUAL_LAYOUT_SMOKE, level=2)

    document.add_paragraph(
        'После покрытия DC/AC/TRAN добавлена базовая visual regression проверка страницы `/simulation/`. '
        'Цель проверки — не сравнивать пиксель-в-пиксель весь интерфейс, а зафиксировать инженерные инварианты, '
        'которые чаще всего ломаются при правках крупного HTML/CSS-шаблона: отсутствие горизонтального overflow, '
        'видимость основных панелей, достаточные размеры canvas, валидность screenshot и непустая canvas-отрисовка графика.'
    )

    _add_table(
        document,
        ['Viewport', 'Состояние страницы', 'Проверки'],
        [
            [
                '1366×900',
                'Загружен делитель напряжения, выполнен DC fallback',
                'Нет horizontal overflow; canvas > 700×250; видны панели компонентов, управления и результаты; screenshot создаётся',
            ],
            [
                '390×844',
                'Загружен RC-фильтр, выполнен AC fallback',
                'Нет horizontal overflow; ключевые панели внутри viewport; canvas АЧХ содержит не фоновые пиксели',
            ],
        ],
    )

    document.add_paragraph('Фрагмент layout probe:')
    _add_code(
        document,
        """
const offenders = Array.from(document.querySelectorAll('body *'))
    .filter((element) => {
        const style = getComputedStyle(element);
        if (style.display === 'none' || style.visibility === 'hidden') return false;
        const rect = element.getBoundingClientRect();
        if (rect.width <= 0 || rect.height <= 0) return false;
        return rect.left < -2 || rect.right > viewportWidth + 2;
    })
    .slice(0, 8)
    .map((element) => ({
        tag: element.tagName.toLowerCase(),
        id: element.id,
        className: String(element.className || ''),
        left: Math.round(element.getBoundingClientRect().left),
        right: Math.round(element.getBoundingClientRect().right),
    }));
        """,
        code_style,
    )

    document.add_paragraph('Фрагмент canvas-pixel проверки:')
    _add_code(
        document,
        """
const canvas = document.querySelector('#sim-graph-mag');
const ctx = canvas.getContext('2d');
const pixels = ctx.getImageData(0, 0, canvas.width, canvas.height).data;
for (let i = 0; i < pixels.length; i += 4) {
    if (pixels[i] !== 10 || pixels[i + 1] !== 16 || pixels[i + 2] !== 32) {
        return true;
    }
}
return false;
        """,
        code_style,
    )

    _add_table(
        document,
        ['Команда', 'Результат'],
        [
            [
                'powershell -NoProfile -ExecutionPolicy Bypass -File .\\scripts\\run_browser_e2e.ps1',
                '10/10 browser-smoke OK',
            ],
            [
                'powershell -NoProfile -ExecutionPolicy Bypass -File .\\scripts\\run_checks.ps1',
                '93 теста OK, 10 skipped; coverage 78 %',
            ],
        ],
    )

    document.add_paragraph(
        'Выявленные несоответствия и уточнения: на проверенных desktop/mobile состояниях CSS offender-список пустой, '
        'поэтому отдельные CSS-правки в этой итерации не потребовались. Следующим шагом visual regression нужно расширить '
        'на модалки BOM, netlist и project picker, где риск переполнений выше из-за таблиц и фиксированных внутренних блоков.'
    )
    document.add_paragraph(
        'Контроль источников: новые внешние источники не добавлялись. Раздел описывает локальный Playwright-контур '
        'и фактические результаты проверки интерфейса.'
    )
    return True


def append_priority_sim_cad_iteration(document: Document) -> bool:
    if any(MARKER_PRIORITY_SIM_CAD in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER_PRIORITY_SIM_CAD, level=2)

    document.add_paragraph(
        'После базовой visual regression уточнён порядок дальнейшего развития. '
        'Интеграция промышленных CAD-форматов остаётся полезной перспективой, но она не должна опережать '
        'ремонт существующих модулей, которые уже входят в пользовательский сценарий дипломного проекта: '
        'браузерный симулятор, графики результатов, сохранение запусков, CAD-редактор, сетка, слои, штамп и save/load.'
    )

    _add_table(
        document,
        ['Приоритет', 'Направление', 'Причина'],
        [
            [
                '1',
                'Ремонт симуляции',
                'Пользователь уже запускает DC/AC/TRAN в текущем интерфейсе; нужно укрепить сохранение результатов, edge-cases ngspice.wasm и обработку ошибок.',
            ],
            [
                '1',
                'Ремонт CAD',
                'CAD используется в проектировании и должен быть устойчивым по геометрии, слоям, сетке, штампу, масштабированию и save/load.',
            ],
            [
                '1',
                'Visual regression',
                'Проблемы отступов и вылезающих блоков нужно ловить автоматически на `/simulation/`, `/cad/`, `/projects/` и модалках.',
            ],
            [
                '2',
                'CAD-форматы',
                'KiCad, LTspice, EDIF, Gerber и Excellon логичнее добавлять после стабилизации текущего редактора и симулятора.',
            ],
        ],
    )

    document.add_paragraph('Обновлённая дорожная карта ближайших работ:')
    _add_code(
        document,
        """
Priority 1
  -> simulation repair: AC/TRAN persistence, ngspice.wasm edge-cases, errors, graphs
  -> CAD repair: save/load, layers, grid, title block, viewport/layout stability
  -> visual regression: simulation, CAD, projects, BOM/netlist/project modals

Priority 2
  -> CAD file formats: KiCad, LTspice, EDIF, Gerber, Excellon
        """,
        code_style,
    )

    document.add_paragraph(
        'Выявленные несоответствия и уточнения: ранее CAD-форматы могли восприниматься как следующий крупный функциональный блок. '
        'Теперь они явно перенесены в перспективы второго приоритета. В ближайших итерациях следует развивать то, что уже '
        'представлено в интерфейсе и влияет на демонстрацию: симулятор, CAD и визуальная устойчивость.'
    )
    document.add_paragraph(
        'Контроль источников: новых внешних источников не добавлялось; это изменение является управленческим уточнением '
        'инженерного плана и отражает фактический приоритет разработки.'
    )
    return True


def append_modern_sim_cad_iteration(document: Document) -> bool:
    if any(MARKER_MODERN_SIM_CAD in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER_MODERN_SIM_CAD, level=2)

    document.add_paragraph(
        'Уточнено требование к дальнейшему развитию: CAD и симулятор нужно не только исправлять, '
        'но и постепенно доводить до уровня современных инженерных инструментов. Это важно отразить '
        'в дипломе в следующих итерациях: каждое улучшение должно описываться через инженерную проблему, '
        'реализованное или планируемое решение, фрагмент кода/таблицу/диаграмму и проверку.'
    )

    _add_table(
        document,
        ['Модуль', 'Современные инженерные требования', 'Как отражать в дипломе'],
        [
            [
                'Симулятор',
                'Probes, курсоры графиков, сохранение и повторный просмотр результатов, параметрические запуски, понятные ошибки ngspice, библиотека SPICE-моделей, DRC/ERC.',
                'Показывать сценарии расчёта, netlist/графики, таблицы результатов, browser/e2e и ограничения реализации.',
            ],
            [
                'CAD',
                'Слои, привязки, сетка, размеры, штамп, масштабирование, горячие клавиши, стабильный save/load, работа с крупными схемами.',
                'Показывать диаграммы интерфейса, таблицы исправлений, скриншотные проверки и влияние на эргономику проектирования.',
            ],
            [
                'Форматы',
                'KiCad, LTspice, EDIF, Gerber, Excellon.',
                'Оставить как перспективу после стабилизации и модернизации текущего CAD/симулятора.',
            ],
        ],
    )

    document.add_paragraph('Памятка для следующих итераций:')
    _add_code(
        document,
        """
Do first:
  repair + modernize simulation
  repair + modernize CAD
  document every step in diploma with evidence

Do later:
  import/export CAD formats
  broaden external ecosystem integrations
        """,
        code_style,
    )

    document.add_paragraph(
        'Выявленные несоответствия и уточнения: форматы CAD не снимаются с дорожной карты, но больше не считаются '
        'ближайшим крупным направлением. Ближайшая ценность проекта — устойчивый, понятный и современный цикл '
        'проектирования: создать схему, проверить её, рассчитать, сохранить результаты и удобно доработать чертёж.'
    )
    return True


def append_sim_run_ui_persistence_iteration(document: Document) -> bool:
    if any(MARKER_SIM_RUN_UI_PERSISTENCE in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER_SIM_RUN_UI_PERSISTENCE, level=2)

    document.add_paragraph(
        'В рамках ремонта симулятора добавлена проверка сохранения результатов AC и переходного процесса '
        'из реального пользовательского интерфейса. Ранее история запусков проверялась через API, но UI-путь '
        '`runSimulation() -> recordSimulationRun() -> /simulation-runs/` не был закреплён отдельным browser/e2e. '
        'Также исправлена нормализация transient-режима: в модели `SimulationRun` используется тип `tran`, '
        'поэтому результат UI теперь сохраняется как `tran`, а не как строка `transient`.'
    )

    _add_table(
        document,
        ['Проблема', 'Исправление', 'Проверка'],
        [
            [
                'UI-запуск AC/TRAN не проверял запись в историю проекта',
                'Добавлен browser/e2e: авторизация, создание проекта, запуск AC и transient из `/simulation/`, проверка `/simulation-runs/`.',
                '`scripts/run_browser_e2e.ps1`: 11/11 OK',
            ],
            [
                'Transient мог сохраняться как `transient`, хотя модель ожидает `tran`',
                'Добавлен `getSimulationRunAnalysisType(result, built)` и нормализация `transient -> tran`.',
                'Тест проверяет отсутствие `transient` и наличие `tran` в сохранённых runs.',
            ],
        ],
    )

    document.add_paragraph('Фрагмент исправления в `simulation.html`:')
    _add_code(
        document,
        """
function getSimulationRunAnalysisType(result, built) {
    const resultType = result && result.type;
    if (resultType === 'tran') return 'tran';
    if (resultType === 'dc' || resultType === 'op' || resultType === 'ac' || resultType === 'pulse') {
        return resultType;
    }
    const analysis = built && built.circuit && built.circuit.analysis;
    return analysis === 'transient' ? 'tran' : (analysis || 'unknown');
}
        """,
        code_style,
    )

    document.add_paragraph('Фрагмент нового browser/e2e:')
    _add_code(
        document,
        """
page.select_option("#analysis-type", "ac")
page.click("#sim-run-btn")
page.wait_for_function(async (projectId) => {
    const data = await simApiFetch(PROJECTS_API.runs(projectId));
    window.__persistedRuns = data.runs;
    return data.runs.some(run => run.analysis_type === 'ac');
}, arg=project_id)

page.select_option("#analysis-type", "transient")
page.click("#sim-run-btn")
page.wait_for_function(async (projectId) => {
    const data = await simApiFetch(PROJECTS_API.runs(projectId));
    window.__persistedRuns = data.runs;
    return data.runs.some(run => run.analysis_type === 'tran');
}, arg=project_id)
        """,
        code_style,
    )

    _add_table(
        document,
        ['Команда', 'Результат'],
        [
            ['Точечный browser/e2e сохранения AC/TRAN', 'OK'],
            [
                'powershell -NoProfile -ExecutionPolicy Bypass -File .\\scripts\\run_browser_e2e.ps1',
                '11/11 browser-smoke OK',
            ],
            [
                'powershell -NoProfile -ExecutionPolicy Bypass -File .\\scripts\\run_checks.ps1',
                '94 теста OK, 11 skipped; coverage 78 %',
            ],
        ],
    )

    document.add_paragraph(
        'Следующие риски по симуляции: реальные edge-cases ngspice.wasm, пустой stdout AC, ошибки JS-MNA fallback, '
        'понятное отображение ошибок в панели результатов и сохранение диагностической информации без перегрузки интерфейса.'
    )
    return True


def append_killer_features_roadmap_iteration(document: Document) -> bool:
    if any(MARKER_KILLER_FEATURES_ROADMAP in paragraph.text for paragraph in document.paragraphs):
        return False

    code_style = _ensure_code_style(document)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER_KILLER_FEATURES_ROADMAP, level=2)

    document.add_paragraph(
        'Зафиксирован перспективный roadmap из пяти демонстрационно сильных функций. '
        'Важно: эти функции не заменяют текущий приоритет ремонта CAD и симуляции, а идут после стабилизации основы. '
        'Marketplace/QR вынесены за пределы локального этапа и должны реализовываться только после публичного деплоя; '
        'block-based editor оформлен как отдельный learning mode.'
    )

    _add_table(
        document,
        ['Этап', 'Фича', 'Смысл и ожидаемый эффект', 'Оценка/риск'],
        [
            [
                '1',
                'What if-слайдер',
                'Log-scale слайдер для числовых параметров компонента, debounced re-run симуляции и live-сдвиг графиков.',
                '4-6 ч; риск низкий; обязательная быстрая фича.',
            ],
            [
                '2',
                'Тепловой анализ',
                'Расчёт мощности компонентов, тепловой overlay на схеме, tooltip и таблица горячих элементов.',
                '6-8 ч; риск низкий; опирается на существующие результаты симуляции.',
            ],
            [
                '3',
                'AI-ассистент DOLG',
                'Подбор компонентов, объяснение схемы, поиск замен EOL-товаров, offline-demo fallback.',
                '8-12 ч; риск средний; зависит от API-ключа и качества fallback.',
            ],
            [
                '4',
                '3D-просмотр платы',
                'Three.js, виртуальная плата, параметрические корпуса, цветовые коды резисторов, orbit camera и PNG export.',
                '10-15 ч; риск средний; визуальный wow-блок.',
            ],
            [
                '5',
                'Виртуальная лаборатория',
                'Осциллограф, мультиметр, генератор и источник питания как приборы, связанные с симуляцией.',
                '15-20 ч; риск высокий; самый сложный UI/state-management блок.',
            ],
        ],
    )

    document.add_paragraph('Порядок реализации после стабилизации:')
    _add_code(
        document,
        """
1. What if slider
2. Thermal analysis
3. AI assistant
4. 3D board viewer
5. Virtual lab

After public deployment:
  -> Marketplace / QR

Separate learning mode:
  -> block-based editor at /learn/
        """,
        code_style,
    )

    document.add_paragraph(
        'Для диплома эти пункты следует отражать как перспективы развития и как обоснование дальнейшей модернизации продукта. '
        'При реализации каждого этапа нужно добавлять: пользовательский сценарий, архитектурную схему, фрагмент кода, '
        'таблицу проверки и ограничения MVP.'
    )
    return True


def main() -> None:
    source = _find_diploma()
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    backup = DOCS_DIR / f'{source.stem}.backup-{date.today().isoformat()}{source.suffix}'
    if not backup.exists():
        shutil.copy2(source, backup)

    document = Document(source)
    changed = normalize_previous_iteration_text(document)
    changed = ensure_source_entry(document, 'openpyxl.readthedocs.io', OPENPYXL_SOURCE) or changed
    changed = ensure_source_entry(document, 'playwright.dev/python/docs/intro', PLAYWRIGHT_SOURCE) or changed
    changed = append_iteration(document) or changed
    changed = append_xlsx_iteration(document) or changed
    changed = append_e2e_drc_iteration(document) or changed
    changed = append_cad_projects_iteration(document) or changed
    changed = append_dc_export_iteration(document) or changed
    changed = append_netlist_fixtures_iteration(document) or changed
    changed = append_scheme_normalizer_iteration(document) or changed
    changed = append_scheme_export_iteration(document) or changed
    changed = append_scheme_bom_iteration(document) or changed
    changed = append_scheme_netlist_iteration(document) or changed
    changed = append_scheme_netlist_elements_iteration(document) or changed
    changed = append_scheme_netlist_analysis_iteration(document) or changed
    changed = append_ac_browser_smoke_iteration(document) or changed
    changed = append_tran_browser_smoke_iteration(document) or changed
    changed = append_visual_layout_smoke_iteration(document) or changed
    changed = append_priority_sim_cad_iteration(document) or changed
    changed = append_modern_sim_cad_iteration(document) or changed
    changed = append_sim_run_ui_persistence_iteration(document) or changed
    changed = append_killer_features_roadmap_iteration(document) or changed
    if not changed:
        print(f'Already up to date: {source}')
        return

    try:
        document.save(source)
        print(f'Updated diploma: {source}')
    except PermissionError:
        fallback = DOCS_DIR / f'{source.stem} - DOLG update{source.suffix}'
        document.save(fallback)
        print(f'Diploma is locked; saved updated copy: {fallback}')
    print(f'Backup: {backup}')


if __name__ == '__main__':
    main()

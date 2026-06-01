from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from docx.table import _Cell
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SCREENSHOTS = DOCS / "diploma_assets" / "screenshots"
REFERENCE_IMAGES = DOCS / "diploma_assets" / "reference_appendix_images"


def clean(text: str) -> str:
    return " ".join((text or "").split())


def find_latest_input() -> Path:
    candidates = sorted(DOCS.glob("*010109*актуализировано.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
    if not candidates:
        raise FileNotFoundError("Не найдена актуализированная копия диплома")
    return candidates[0]


def find_para(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if clean(paragraph.text).startswith(prefix):
            return paragraph
    raise ValueError(f"Не найден абзац: {prefix}")


def remove_after(paragraph: Paragraph) -> None:
    node = paragraph._p.getnext()
    while node is not None:
        next_node = node.getnext()
        node.getparent().remove(node)
        node = next_node


def insert_paragraph_after(anchor: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def set_cell_text(cell: _Cell, text: str, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell: _Cell, fill: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell: _Cell, top: int = 80, start: int = 80, bottom: int = 80, end: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell: _Cell, width_cm: float) -> None:
    width_twips = int(width_cm * 567)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_table_width_pct(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")


def format_table(table, font_size: int = 9) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_table_width_pct(table)

    col_count = len(table.columns)
    width_map = {
        2: [5.2, 11.2],
        3: [4.2, 5.3, 7.0],
        4: [2.4, 4.3, 5.8, 4.0],
    }
    widths = width_map.get(col_count, [16.5 / max(col_count, 1)] * col_count)

    if table.rows:
        set_repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(table.rows):
        set_row_cant_split(row)
        for col_idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(col_idx, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                shade_cell(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(font_size)
                    if row_idx == 0:
                        run.bold = True


def insert_table_after(anchor: Paragraph, caption: str, headers: list[str], rows: list[list[str]], size: int = 10) -> Paragraph:
    caption_para = insert_paragraph_after(anchor, caption, "Normal")
    caption_para.paragraph_format.first_line_indent = Cm(0)
    caption_para.paragraph_format.space_before = Pt(6)
    caption_para.paragraph_format.space_after = Pt(3)

    table = anchor._parent.add_table(rows=1, cols=len(headers), width=Inches(6.3))
    try:
        table.style = "Table Grid"
    except KeyError:
        try:
            table.style = "Normal Table"
        except KeyError:
            pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=size)
        shade_cell(table.rows[0].cells[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=size)
    format_table(table, font_size=min(size, 9))

    table_xml = table._tbl
    table_xml.getparent().remove(table_xml)
    caption_para._p.addnext(table_xml)
    tail = OxmlElement("w:p")
    table_xml.addnext(tail)
    return Paragraph(tail, caption_para._parent)


def insert_code_after(doc: Document, anchor: Paragraph, code: str) -> Paragraph:
    style = "DOLG Code" if "DOLG Code" in [s.name for s in doc.styles] else "Normal"
    for line in code.strip("\n").splitlines():
        paragraph = insert_paragraph_after(anchor, line.rstrip(), style)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(8)
        anchor = paragraph
    return anchor


def insert_picture_after(anchor: Paragraph, image_path: Path, caption: str, width_inches: float = 5.9) -> Paragraph:
    picture_para = insert_paragraph_after(anchor, "", "Normal")
    picture_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_para.paragraph_format.first_line_indent = Cm(0)
    picture_para.paragraph_format.space_before = Pt(6)
    picture_para.paragraph_format.space_after = Pt(3)
    picture_para.add_run().add_picture(str(image_path), width=Inches(width_inches))

    caption_para = insert_paragraph_after(picture_para, caption, "Normal")
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.paragraph_format.first_line_indent = Cm(0)
    caption_para.paragraph_format.space_before = Pt(0)
    caption_para.paragraph_format.space_after = Pt(6)
    return caption_para


def page_break_before(paragraph: Paragraph) -> None:
    run = paragraph.insert_paragraph_before().add_run()
    run.add_break(WD_BREAK.PAGE)


def set_section_a4_portrait(section) -> None:
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)


def set_section_a3_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(42)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)


def is_a3_landscape(section) -> bool:
    if section.page_width is None or section.page_height is None:
        return False
    return section.page_width > section.page_height and abs(section.page_width - Cm(42)) < Cm(0.5)


def apply_document_layout(doc: Document) -> None:
    # Reference-style VКР layout: portrait A4 for text, landscape A3 for readable UI screenshots.
    for section in doc.sections:
        if is_a3_landscape(section):
            set_section_a3_landscape(section)
        else:
            set_section_a4_portrait(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    heading1 = styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1.paragraph_format.first_line_indent = Cm(0)
    heading1.paragraph_format.space_before = Pt(0)
    heading1.paragraph_format.space_after = Pt(12)
    heading1.paragraph_format.line_spacing = 1.5
    heading1.paragraph_format.page_break_before = True

    heading2 = styles["Heading 2"]
    heading2.font.name = "Times New Roman"
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    heading2.paragraph_format.first_line_indent = Cm(0)
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(6)
    heading2.paragraph_format.line_spacing = 1.5

    if "DOLG Code" in [s.name for s in styles]:
        code_style = styles["DOLG Code"]
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(8)
        code_style.paragraph_format.first_line_indent = Cm(0)
        code_style.paragraph_format.line_spacing = 1.0
        code_style.paragraph_format.space_after = Pt(1)

    for paragraph in doc.paragraphs:
        text = clean(paragraph.text)
        if not text:
            continue
        if paragraph.style.name == "Heading 1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(16)
                run.bold = True
        elif paragraph.style.name == "Heading 2":
            paragraph.paragraph_format.first_line_indent = Cm(0)
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
                run.bold = True
        elif paragraph.style.name == "DOLG Code":
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(8)
        elif text.startswith(("Таблица", "Рисунок", "Приложение")):
            paragraph.paragraph_format.first_line_indent = Cm(0)
        else:
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                if run.font.size is None:
                    run.font.size = Pt(14)

    for table in doc.tables:
        format_table(table, font_size=9)


def rebuild_appendices(doc: Document) -> None:
    appendix = find_para(doc, "ПРИЛОЖЕНИЯ")
    remove_after(appendix)
    appendix.style = "Heading 1"

    anchor = appendix
    anchor = insert_paragraph_after(anchor, "Приложение А. Сводные сведения о реализации проекта", "Heading 2")
    anchor = insert_table_after(
        anchor,
        "Таблица А.1 – Состав программных модулей DOLG",
        ["Модуль", "Назначение", "Основные файлы"],
        [
            ["shop", "Каталог, полки категорий, карточки товаров, корзина, BOM, сравнение", "shop/models.py; shop/views.py; shop/templates/shop/includes/product_card.html"],
            ["accounts", "Профиль, регистрация, адреса", "accounts/models.py; accounts/views.py"],
            ["orders", "Заказы, повтор заказа, платежные сценарии", "orders/models.py; orders/views.py; orders/payment_views.py"],
            ["knowledge", "Энциклопедия, статьи, внутренние ссылки и вложенные материалы", "knowledge/models.py; knowledge/views.py; knowledge/templates/knowledge/*"],
            ["Dolg_APP", "Схемы, CAD, симуляция, AI, 3D, sharing", "Dolg_APP/models.py; Dolg_APP/views.py; templates/tools/*"],
        ],
    )
    anchor = insert_table_after(
        anchor,
        "Таблица А.2 – Реализованные инженерные функции",
        ["Функция", "Результат", "Статус"],
        [
            ["Редактор схем", "УГО-режим, шаг хода, маршрутизация проводов, BOM", "реализовано"],
            ["Симуляция", "DC, AC, TRAN, ngspice.wasm, JS-MNA fallback", "реализовано"],
            ["Тепловой анализ", "Расчет мощности и TDP-таблица", "реализовано"],
            ["AI-ассистент", "Подбор, объяснение схемы, замена EOL", "реализовано с demo fallback"],
            ["3D-просмотр", "Three.js-модель виртуальной платы", "реализовано"],
            ["Virtual lab", "Осциллограф, мультиметр, генератор сигналов", "реализовано"],
            ["Каталог", "89 товаров, 43 РЭБ-компонента, фото товаров, полки категорий", "реализовано"],
            ["Энциклопедия", "21 статья, 22 материала: ссылки, фото, gif, видео и файлы", "реализовано"],
        ],
    )
    anchor = insert_table_after(
        anchor,
        "Таблица А.3 – Интеграция материалов предыдущих приложений",
        ["Материал предыдущей версии", "Как интегрирован в финальную редакцию", "Раздел"],
        [
            ["Связь редактора схем с каталогом", "Сжато до архитектурного описания BOM, SKU-привязки и поиска компонентов", "2.5, приложение А"],
            ["Серверный XLSX-экспорт BOM", "Оставлен как доказательство связки проектирования и закупки", "2.5, приложение Г"],
            ["Browser-smoke, DRC/BOM и аудит UI", "Описание сокращено до таблицы контроля качества", "2.6, приложение Е"],
            ["CAD, проекты и normalizeSchemeData", "Включено в описание устойчивого формата JSON и CAD-редактора", "2.2, 2.5"],
            ["DC/AC/TRAN, SVG/PDF и netlist builder", "Сведено в раздел симуляции и кодовые листинги", "2.6, приложение Д"],
            ["Приоритет ремонта CAD и симуляции", "Вынесено в план развития без дневниковых формулировок", "Заключение, приложение Е"],
        ],
    )
    anchor = insert_table_after(
        anchor,
        "Таблица А.4 – Актуализация пользовательского контура от 12.05.2026",
        ["Направление", "Что изменено", "Эффект для демонстрации"],
        [
            ["Ассортимент", "Каталог расширен до 89 товаров, включая 43 РЭБ-компонента; у всех позиций заполнены фотографии", "Карточки выглядят как реальные товарные позиции, а не как набор заглушек"],
            ["Главная страница", "Вместо неоднозначного блока «Все товары» добавлены полки категорий и блок «Новые и актуальные позиции»", "Каталог читается как marketplace и масштабируется за пределы первой страницы"],
            ["Карточки товаров", "Общий шаблон карточки используется в каталоге, похожих товарах и недавно просмотренных позициях", "Единый внешний вид снижает визуальные расхождения между разделами"],
            ["Энциклопедия", "Добавлены внутренние ссылки и модель ArticleMaterial для изображений, gif-анимаций, видео, файлов и внешних материалов", "Статьи стали не только текстовыми, но и учебно-демонстрационными"],
            ["Наполнение", "Опубликована 21 статья по 6 категориям и 22 дополнительных материала", "На защите можно показать предметное содержание по каждой категории"],
        ],
    )

    set_section_a3_landscape(doc.add_section(WD_SECTION_START.NEW_PAGE))
    anchor = doc.add_paragraph("Приложение Б. Иллюстрации пользовательского интерфейса", style="Heading 2")
    images = [
        ("01_home_catalog.png", "Рисунок Б.1 – Главная страница и верхний блок каталога", 14.1),
        ("05_catalog_shelves.png", "Рисунок Б.2 – Полки категорий на главной странице каталога", 14.1),
        ("06_product_detail_photo.png", "Рисунок Б.3 – Карточка товара с фотографией и техническими параметрами", 13.2),
        ("07_product_related_cards.png", "Рисунок Б.4 – Блок похожих товаров с единым шаблоном карточек", 14.1),
        ("08_knowledge_article_materials.png", "Рисунок Б.5 – Статья энциклопедии с внутренними ссылками и вложенными материалами", 13.2),
        ("02_simulation_ac_graph.png", "Рисунок Б.6 – Редактор схем и панель симуляции", 14.1),
        ("03_ac_graph_panel.png", "Рисунок Б.7 – График амплитудно-частотной характеристики", 11.2),
        ("03b_ac_phase_graph.png", "Рисунок Б.8 – График фазочастотной характеристики", 11.2),
    ]
    for index, (name, caption, width) in enumerate(images):
        path = SCREENSHOTS / name
        if path.exists():
            if index:
                breaker = insert_paragraph_after(anchor, "", "Normal")
                breaker.add_run().add_break(WD_BREAK.PAGE)
                anchor = breaker
            anchor = insert_picture_after(anchor, path, caption, width)

    set_section_a3_landscape(doc.add_section(WD_SECTION_START.NEW_PAGE))
    anchor = doc.add_paragraph("Приложение В. Диаграммы из референсной версии проекта", style="Heading 2")
    reference_images = [
        ("reference_01.png", "Рисунок В.1 – Диаграмма последовательности оформления заказа", 14.1),
        ("reference_02.png", "Рисунок В.2 – Диаграмма классов коммерческого контура", 14.1),
        ("reference_03.png", "Рисунок В.3 – Диаграмма состояний заказа", 7.2),
        ("reference_04.png", "Рисунок В.4 – Диаграмма процесса покупки компонента", 5.6),
        ("reference_05.png", "Рисунок В.5 – Детализированная диаграмма последовательности", 13.2),
    ]
    for index, (name, caption, width) in enumerate(reference_images):
        path = REFERENCE_IMAGES / name
        if path.exists():
            if index:
                breaker = insert_paragraph_after(anchor, "", "Normal")
                breaker.add_run().add_break(WD_BREAK.PAGE)
                anchor = breaker
            anchor = insert_picture_after(anchor, path, caption, width)

    set_section_a4_portrait(doc.add_section(WD_SECTION_START.NEW_PAGE))
    anchor = doc.add_paragraph("Приложение Г. Фрагменты исходного кода", style="Heading 2")
    anchor = insert_paragraph_after(anchor, "Листинг В.1 – Кеширование стабильного AI-префикса", "Normal")
    anchor.paragraph_format.first_line_indent = Cm(0)
    anchor = insert_code_after(
        doc,
        anchor,
        """
PROMPT_CACHE_MIN_CHARS = 4500
if len(stable_text) >= PROMPT_CACHE_MIN_CHARS:
    stable_block["cache_control"] = {"type": "ephemeral"}
        """,
    )
    anchor = insert_paragraph_after(anchor, "Листинг В.2 – Преобразование сигнала лаборатории в SPICE-источник", "Normal")
    anchor.paragraph_format.first_line_indent = Cm(0)
    anchor = insert_code_after(
        doc,
        anchor,
        """
if (sig.wave === "sine") {
    value = "SIN(" + offset + " " + amplitude + " " + frequency + ")";
} else if (sig.wave === "square") {
    value = "PULSE(" + low + " " + high + " 0 1n 1n " + halfPeriod + " " + period + ")";
} else if (sig.wave === "triangle") {
    value = "PWL(0 " + offset + " " + q1 + " " + high + " " + q3 + " " + low + " " + period + " " + offset + ")";
}
        """,
    )
    anchor = insert_paragraph_after(anchor, "Листинг В.3 – Защита shared-материалов Three.js", "Normal")
    anchor.paragraph_format.first_line_indent = Cm(0)
    anchor = insert_code_after(
        doc,
        anchor,
        """
const _sharedMat = (opts) => {
    const material = new THREE.MeshStandardMaterial(opts);
    material.userData._shared = true;
    return material;
};
if (!material.userData || !material.userData._shared) {
    material.dispose();
}
        """,
    )

    anchor = insert_paragraph_after(anchor, "Приложение Д. Контроль качества и план развития", "Heading 2")
    anchor = insert_table_after(
        anchor,
        "Таблица Г.1 – Сводка проверок",
        ["Проверка", "Результат", "Назначение"],
        [
            ["manage.py check", "0 замечаний", "Проверка конфигурации Django"],
            ["Django tests", "около 130 тестов; browser-smoke пропускается без RUN_BROWSER_E2E", "Регрессия backend и доменной логики"],
            ["Browser e2e", "16/16 OK при отдельном запуске", "Проверка canvas, CAD, модальных окон и графиков"],
            ["AI+share suite", "31/31 OK в FAST_TESTS-режиме", "Проверка AI endpoint, sharing и catalog cache"],
            ["Аудит", "P0 отсутствуют; P1 закрыты; P2 задокументированы", "Фиксация остаточных рисков"],
        ],
    )
    anchor = insert_table_after(
        anchor,
        "Таблица Г.2 – Приоритеты дальнейшего развития",
        ["Приоритет", "Направление", "Ожидаемый результат"],
        [
            ["1", "Ремонт и модернизация симуляции", "Стабильные расчеты, probes, курсоры, сохранение результатов"],
            ["1", "Ремонт и модернизация CAD", "Чертежный вид, слои, привязки, размеры, устойчивый layout"],
            ["2", "Screenshot-регрессия", "Автоматический контроль вылезающих блоков и кривых отступов"],
            ["3", "CAD-форматы", "KiCad, LTspice, Gerber, Excellon после стабилизации ядра"],
            ["4", "Публичный деплой", "PostgreSQL, мониторинг, резервное копирование, HTTPS"],
        ],
    )


def add_page_breaks(doc: Document) -> None:
    break_heads = {
        "ГЛАВА 1. АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ",
        "ГЛАВА 2. ПРОЕКТИРОВАНИЕ СИСТЕМЫ",
        "ЗАКЛЮЧЕНИЕ",
        "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ И ИСТОЧНИКОВ",
        "ПРИЛОЖЕНИЯ",
        "Приложение А. Сводные сведения о реализации проекта",
        "Приложение Г. Фрагменты исходного кода",
        "Приложение Д. Контроль качества и план развития",
    }
    for paragraph in list(doc.paragraphs):
        if clean(paragraph.text) in break_heads:
            previous = paragraph._p.getprevious()
            if previous is not None:
                page_break_before(paragraph)


def final_text_polish(doc: Document) -> None:
    replacements = {
        "Web-сайта": "web-сайта",
        "web-сайта": "веб-сайта",
        "web-приложение": "веб-приложение",
        "Frontend": "frontend",
        "Backend": "backend",
        "production-план": "план промышленного развертывания",
        "read-only": "режим просмотра",
        "sharing": "публичный просмотр",
        "fallback": "резервный режим",
        "smoke-тест": "smoke-тест",
        "72 товара": "89 товаров",
        "72 сейчас": "89 сейчас",
        "26 РЭБ-компонентов": "43 РЭБ-компонента",
        "12 демо-схем": "12 демонстрационных схем",
    }
    for paragraph in iter_all_paragraphs(doc):
        text = paragraph.text
        updated = text
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated != text:
            paragraph.clear()
            paragraph.add_run(updated)


def iter_all_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def validate(path: Path) -> dict[str, object]:
    with ZipFile(path) as archive:
        bad = archive.testzip()
    doc = Document(path)
    headings = [clean(p.text) for p in doc.paragraphs if p.style.name.startswith("Heading") and clean(p.text)]
    return {
        "zip_bad_member": bad,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "headings": headings,
        "has_dopolnenie": any("Дополнение от" in p.text for p in doc.paragraphs),
        "chapter3": any(clean(p.text).startswith("ГЛАВА 3") for p in doc.paragraphs),
        "images": len(doc.inline_shapes),
    }


def main() -> None:
    source = find_latest_input()
    output = DOCS / "Диплом_DOLG_финальная_редакция_20260513_v4.docx"
    doc = Document(source)
    rebuild_appendices(doc)
    final_text_polish(doc)
    apply_document_layout(doc)
    add_page_breaks(doc)
    doc.save(output)
    stats = validate(output)
    print(f"SOURCE={source}")
    print(f"OUTPUT={output}")
    for key, value in stats.items():
        if key != "headings":
            print(f"{key}={value}")
    print("headings:")
    for heading in stats["headings"]:
        print(f"- {heading}")


if __name__ == "__main__":
    main()

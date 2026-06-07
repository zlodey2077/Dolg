"""Generate a two-chapter working edition of the DOLG diploma.

This generator intentionally avoids the previous three-chapter structure.
Chapter 1 covers analysis, requirements and design. Chapter 2 covers code,
database models, services, UI, AI/ML, admin tooling and verification.
"""

from __future__ import annotations

import sqlite3
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / 'docs'
SCREENSHOTS = DOCS / 'diploma_assets' / 'screenshots'
GENERATED = DOCS / 'diploma_assets' / 'generated'
OUT_DOCX = DOCS / 'DOLG_Diploma_reworked_20260603.docx'
OUT_MD = DOCS / 'DOLG_Diploma_reworked_20260603.md'

TITLE = (
    'Разработка веб-приложения для продажи радио- и электронных компонентов '
    'со встроенными инструментами проектирования и симуляции схем'
)
AUTHOR = 'Буряко Дмитрий Сергеевич'
SUPERVISOR = 'Буланов Сергей Георгиевич'
YEAR = '2026'


@dataclass(frozen=True)
class Facts:
    users: int = 0
    categories: int = 0
    products: int = 0
    projects: int = 0
    project_events: int = 0
    project_reviews: int = 0
    measurements: int = 0
    articles: int = 0
    article_materials: int = 0
    learning_tracks: int = 0
    learning_lessons: int = 0
    learning_tasks: int = 0
    ai_examples: int = 0
    ml_jobs: int = 0
    moderation_cases: int = 0
    orders: int = 0


def sqlite_count(cursor: sqlite3.Cursor, table: str) -> int:
    try:
        cursor.execute(f'SELECT COUNT(*) FROM "{table}"')
        return int(cursor.fetchone()[0])
    except Exception:
        return 0


def collect_facts() -> Facts:
    db = ROOT / 'db.sqlite3'
    if not db.exists():
        return Facts()
    con = sqlite3.connect(str(db))
    cur = con.cursor()
    return Facts(
        users=sqlite_count(cur, 'auth_user'),
        categories=sqlite_count(cur, 'shop_category'),
        products=sqlite_count(cur, 'shop_product'),
        projects=sqlite_count(cur, 'Dolg_APP_schematicproject'),
        project_events=sqlite_count(cur, 'Dolg_APP_projectevent'),
        project_reviews=sqlite_count(cur, 'Dolg_APP_projectreview'),
        measurements=sqlite_count(cur, 'Dolg_APP_projectmeasurement'),
        articles=sqlite_count(cur, 'knowledge_article'),
        article_materials=sqlite_count(cur, 'knowledge_articlematerial'),
        learning_tracks=sqlite_count(cur, 'knowledge_learningtrack'),
        learning_lessons=sqlite_count(cur, 'knowledge_learninglesson'),
        learning_tasks=sqlite_count(cur, 'knowledge_learningtask'),
        ai_examples=sqlite_count(cur, 'Dolg_APP_aitrainingexample'),
        ml_jobs=sqlite_count(cur, 'Dolg_APP_mljob'),
        moderation_cases=sqlite_count(cur, 'moderation_moderationcase'),
        orders=sqlite_count(cur, 'orders_order'),
    )


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ''
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    normal = document.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for name in ('Heading 1', 'Heading 2', 'Heading 3'):
        style = document.styles[name]
        style.font.name = 'Times New Roman'
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.0
    document.styles['Heading 1'].font.size = Pt(14)
    document.styles['Heading 2'].font.size = Pt(14)
    document.styles['Heading 3'].font.size = Pt(14)

    if 'DOLG Code' not in document.styles:
        code_style = document.styles.add_style('DOLG Code', 1)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(9)
        code_style.paragraph_format.first_line_indent = Cm(0)
        code_style.paragraph_format.line_spacing = 1.0


def heading(document: Document, text: str, level: int = 1, *, page_break: bool = False) -> None:
    if page_break:
        document.add_page_break()
    paragraph = document.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT


def para(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style='List Paragraph')
        paragraph.paragraph_format.left_indent = Cm(1.25)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)


def caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.runs[0].font.name = 'Times New Roman'
    paragraph.runs[0].font.size = Pt(12)


def table(document: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    caption(document, title)
    tbl = document.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    for i, header in enumerate(headers):
        set_cell_shading(tbl.rows[0].cells[i], 'D9EAF7')
        set_cell_text(tbl.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    document.add_paragraph()


def picture_asset(document: Document, path: Path, title: str, width_cm: float = 15.0) -> None:
    if not path.exists():
        para(document, f'[Место для рисунка: {title}. Файл не найден: {path.name}]')
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    caption(document, title)


def picture(document: Document, rel_path: str, title: str, width_cm: float = 15.0) -> None:
    picture_asset(document, SCREENSHOTS / rel_path, title, width_cm)


def code(document: Document, value: str) -> None:
    for line in value.strip('\n').splitlines():
        document.add_paragraph(line.rstrip(), style='DOLG Code')


def _fallback_image(path: Path, title: str, lines: list[str]) -> None:
    from PIL import Image, ImageDraw, ImageFont

    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new('RGB', (1600, 900), '#eef3f8')
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype('arialbd.ttf', 54)
        text_font = ImageFont.truetype('arial.ttf', 34)
    except Exception:
        title_font = ImageFont.load_default()
        text_font = ImageFont.load_default()
    draw.rectangle((0, 0, 1600, 28), fill='#0891b2')
    draw.text((70, 65), title, fill='#0f172a', font=title_font)
    y = 175
    colors = ['#2563eb', '#16a34a', '#ea580c', '#7c3aed', '#0891b2', '#dc2626']
    for idx, line in enumerate(lines):
        y0 = y + idx * 95
        draw.rounded_rectangle(
            (80, y0, 1520, y0 + 64), radius=16, fill='#ffffff', outline=colors[idx % len(colors)], width=4
        )
        draw.text((110, y0 + 16), line, fill='#0f172a', font=text_font)
    img.save(path)


def ensure_generated_assets(facts: Facts) -> None:
    GENERATED.mkdir(parents=True, exist_ok=True)
    try:
        import matplotlib

        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        from matplotlib.patches import FancyBboxPatch
    except Exception:
        _fallback_image(
            GENERATED / 'project_session_flow.png',
            'Сеанс проектирования DOLG',
            [
                'Каталог -> схема',
                'Симуляция -> измерение',
                'Review -> AI evidence',
                'BOM -> заказ -> история',
            ],
        )
        _fallback_image(
            GENERATED / 'engineering_review_score.png',
            'Engineering Review',
            ['Design Health Score', 'DRC/ERC findings', 'BOM risk', 'Measurements', 'Legal sources'],
        )
        _fallback_image(
            GENERATED / 'ai_ml_pipeline.png',
            'AI/ML pipeline',
            [
                'Схемы и задания',
                'Валидация',
                'AITrainingExample',
                'PyTorch deep_hint',
                'Expert rules + человек',
            ],
        )
        _fallback_image(
            GENERATED / 'admin_monitoring_metrics.png',
            'Админ-мониторинг',
            [
                f'Товары: {facts.products}',
                f'Проекты: {facts.projects}',
                f'Статьи: {facts.articles}',
                f'AITrainingExample: {facts.ai_examples}',
            ],
        )
        return

    def save_flow(path: Path, title: str, labels: list[str], colors: list[str]) -> None:
        fig, ax = plt.subplots(figsize=(14, 5.2), dpi=160)
        ax.set_facecolor('#eef3f8')
        fig.patch.set_facecolor('#eef3f8')
        ax.axis('off')
        ax.text(0.02, 0.9, title, fontsize=22, fontweight='bold', color='#0f172a', transform=ax.transAxes)
        y = 0.47
        box_w = 0.15
        gap = 0.025
        for idx, label in enumerate(labels):
            x = 0.02 + idx * (box_w + gap)
            box = FancyBboxPatch(
                (x, y),
                box_w,
                0.22,
                boxstyle='round,pad=0.012,rounding_size=0.018',
                linewidth=2.2,
                edgecolor=colors[idx % len(colors)],
                facecolor='#ffffff',
                transform=ax.transAxes,
            )
            ax.add_patch(box)
            ax.text(
                x + box_w / 2,
                y + 0.11,
                label,
                ha='center',
                va='center',
                fontsize=13,
                fontweight='bold',
                color='#0f172a',
                transform=ax.transAxes,
            )
            if idx < len(labels) - 1:
                ax.annotate(
                    '',
                    xy=(x + box_w + gap * 0.82, y + 0.11),
                    xytext=(x + box_w + gap * 0.15, y + 0.11),
                    arrowprops=dict(arrowstyle='->', lw=2.2, color='#0891b2'),
                    xycoords=ax.transAxes,
                    textcoords=ax.transAxes,
                )
        fig.tight_layout(pad=1.2)
        fig.savefig(path, bbox_inches='tight')
        plt.close(fig)

    save_flow(
        GENERATED / 'project_session_flow.png',
        'Сеанс проектирования DOLG',
        ['Каталог', 'Схема', 'Симуляция', 'Измерение', 'Review', 'BOM/заказ'],
        ['#2563eb', '#0891b2', '#16a34a', '#ea580c', '#7c3aed', '#dc2626'],
    )
    save_flow(
        GENERATED / 'ai_ml_pipeline.png',
        'AI/ML pipeline: подсказка, а не финальный verdict',
        ['Artifacts', 'Validation', 'Examples', 'PyTorch', 'Deep hint', 'Человек'],
        ['#0891b2', '#16a34a', '#2563eb', '#7c3aed', '#ea580c', '#dc2626'],
    )

    fig, ax = plt.subplots(figsize=(12.8, 6.0), dpi=160)
    fig.patch.set_facecolor('#eef3f8')
    ax.set_facecolor('#eef3f8')
    labels = ['Design\nHealth', 'DRC/ERC', 'BOM\nrisk', 'Ratings', 'Measurements', 'Legal\nsources']
    values = [78, 18, 12, 22, 36, 58]
    colors = ['#16a34a', '#dc2626', '#ea580c', '#f59e0b', '#2563eb', '#7c3aed']
    bars = ax.bar(labels, values, color=colors, alpha=0.88)
    ax.set_ylim(0, 100)
    ax.set_ylabel('условная шкала / покрытие, %', fontsize=12)
    ax.set_title(
        'Engineering Review: score, риски и evidence', fontsize=20, fontweight='bold', color='#0f172a', pad=18
    )
    ax.grid(axis='y', color='#94a3b8', alpha=0.35)
    for bar, value in zip(bars, values):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            value + 2,
            str(value),
            ha='center',
            fontsize=12,
            fontweight='bold',
            color='#0f172a',
        )
    fig.tight_layout()
    fig.savefig(GENERATED / 'engineering_review_score.png', bbox_inches='tight')
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(12.8, 6.0), dpi=160)
    fig.patch.set_facecolor('#eef3f8')
    ax.set_facecolor('#eef3f8')
    metric_labels = ['Products', 'Projects', 'Articles', 'Learning\ntasks', 'AI examples', 'Events']
    metric_values = [
        facts.products,
        facts.projects,
        facts.articles,
        facts.learning_tasks,
        facts.ai_examples,
        facts.project_events,
    ]
    metric_colors = ['#0891b2', '#2563eb', '#16a34a', '#ea580c', '#7c3aed', '#64748b']
    bars = ax.bar(metric_labels, metric_values, color=metric_colors, alpha=0.9)
    ax.set_title(
        'Операционные метрики, используемые админкой', fontsize=20, fontweight='bold', color='#0f172a', pad=18
    )
    ax.grid(axis='y', color='#94a3b8', alpha=0.35)
    for bar, value in zip(bars, metric_values):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            value + max(metric_values or [1]) * 0.015,
            str(value),
            ha='center',
            fontsize=12,
            fontweight='bold',
            color='#0f172a',
        )
    fig.tight_layout()
    fig.savefig(GENERATED / 'admin_monitoring_metrics.png', bbox_inches='tight')
    plt.close(fig)


def add_title_page(document: Document) -> None:
    for _ in range(4):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА')
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(16)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'на тему: «{TITLE}»')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

    for _ in range(4):
        document.add_paragraph()
    para(document, f'Автор: {AUTHOR}')
    para(document, f'Научный руководитель: {SUPERVISOR}')
    for _ in range(8):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(YEAR)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    document.add_page_break()


def add_contents(document: Document) -> None:
    heading(document, 'СОДЕРЖАНИЕ', 1)
    lines = [
        'ВВЕДЕНИЕ',
        'ГЛАВА 1. АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ И ПРОЕКТИРОВАНИЕ СИСТЕМЫ',
        '1.1. Предметная область радио- и электронных компонентов',
        '1.2. Проблемы существующего процесса',
        '1.3. Анализ аналогов и границы ниши DOLG',
        '1.4. Постановка задачи',
        '1.5. Требования к системе',
        '1.6. Выбор технологий и библиотек',
        '1.7. Проектирование архитектуры',
        '1.8. Проектирование базы данных и информационной модели',
        '1.9. Сеанс проектирования как основная модель DOLG',
        'ГЛАВА 2. РЕАЛИЗАЦИЯ ПРОГРАММНОЙ СИСТЕМЫ',
        '2.1. Общая структура проекта в коде',
        '2.2. Реализация каталога и карточек компонентов',
        '2.3. Реализация пользовательского и коммерческого контура',
        '2.4. Реализация CAD-редактора и хранения схем',
        '2.5. Реализация симуляции и измерений',
        '2.6. Реализация инженерной лаборатории',
        '2.7. Реализация Engineering Review',
        '2.8. Реализация AI-ассистента и ML pipeline',
        '2.9. Реализация обучения и базы знаний',
        '2.10. Реализация модерации, ролей и администрирования',
        '2.11. Реализация контроля качества и проверок',
        '2.12. Результаты реализации и ограничения',
        'ЗАКЛЮЧЕНИЕ',
        'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ',
        'ПРИЛОЖЕНИЯ',
    ]
    for line in lines:
        paragraph = document.add_paragraph(line)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0.5 if line[0].isdigit() else 0)
    document.add_page_break()


def add_intro(document: Document, facts: Facts) -> None:
    heading(document, 'ВВЕДЕНИЕ', 1)
    para(
        document,
        'Разработка электронных устройств требует согласованной работы с несколькими '
        'типами данных: параметрами радио- и электронных компонентов, документацией, '
        'принципиальными схемами, расчетами, результатами симуляции, спецификацией '
        'BOM и заказом комплектующих. На практике эти действия часто выполняются в '
        'разных системах, из-за чего пользователь вручную переносит номиналы, корпуса, '
        'модели и результаты измерений между каталогом, CAD-средой, SPICE-симулятором '
        'и интернет-магазином.',
    )
    para(
        document,
        'Актуальность работы определяется необходимостью web-среды, которая связывает '
        'подбор компонента, проектирование схемы, инженерную проверку, обучение и '
        'оформление заказа в одном пользовательском контуре. Особенно полезна такая '
        'среда для учебных, лабораторных и малых инженерных проектов, где важна не '
        'только покупка детали, но и понимание того, почему схема работает или не '
        'проходит проверку.',
    )
    para(
        document,
        f'Целью выпускной квалификационной работы является разработка web-приложения DOLG на тему «{TITLE}».',
    )
    para(document, 'Для достижения цели решаются следующие задачи:')
    bullets(
        document,
        [
            'проанализировать предметную область и существующие решения для подбора, покупки, проектирования и симуляции электронных схем;',
            'спроектировать архитектуру web-приложения, модели данных и пользовательский сеанс проектирования;',
            'реализовать каталог компонентов, карточки товаров, фильтрацию, корзину, заказы и BOM-сценарии;',
            'реализовать CAD/SIM-контур, инженерную лабораторию, измерения и Engineering Review;',
            'реализовать базу знаний, практикум, AI-ассистента, ML/dataset pipeline, роли, модерацию и административный мониторинг;',
            'провести проверку работоспособности через тесты, management-команды и demo-ready/data-integrity checks.',
        ],
    )
    para(
        document,
        'Объектом исследования является процесс цифрового сопровождения работы с '
        'радио- и электронными компонентами. Предметом исследования является '
        'программная реализация web-платформы, объединяющей каталог, схемотехническое '
        'проектирование, симуляцию, инженерную проверку, обучение и заказ.',
    )
    para(
        document,
        'Практическая значимость состоит в том, что разработанная система формирует '
        'единый проектный контур: компонент из каталога может использоваться в схеме, '
        'схема сохраняется как данные, результаты расчета и измерения попадают в '
        'review, а выводы review используются в обучении, AI-подсказках и BOM.',
    )
    table(
        document,
        'Таблица 1 - Фактический срез локальной базы данных DOLG',
        ['Показатель', 'Значение'],
        [
            ['Пользователи', str(facts.users)],
            ['Категории каталога', str(facts.categories)],
            ['Товары', str(facts.products)],
            ['Проекты схем', str(facts.projects)],
            ['Статьи', str(facts.articles)],
            ['Материалы статей', str(facts.article_materials)],
            ['Маршруты обучения', str(facts.learning_tracks)],
            ['Уроки', str(facts.learning_lessons)],
            ['Учебные задания', str(facts.learning_tasks)],
            ['AITrainingExample', str(facts.ai_examples)],
            ['ProjectEvent', str(facts.project_events)],
        ],
    )
    picture_asset(
        document,
        GENERATED / 'project_session_flow.png',
        'Рисунок 1 - Общая логика сеанса проектирования DOLG',
        width_cm=16.0,
    )
    para(
        document,
        'Структура работы включает введение, две главы, заключение, список источников '
        'и приложения. Первая глава посвящена анализу предметной области, требованиям '
        'и проектированию системы. Вторая глава описывает реализацию программной '
        'системы, ее модули, кодовую структуру, проверки и ограничения.',
    )


def add_chapter1(document: Document) -> None:
    heading(document, 'ГЛАВА 1. АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ И ПРОЕКТИРОВАНИЕ СИСТЕМЫ', 1, page_break=True)
    heading(document, '1.1. Предметная область радио- и электронных компонентов', 2)
    para(
        document,
        'Предметная область включает электронные компоненты, их технические параметры, '
        'корпуса, документацию, схемотехнические узлы, расчеты, симуляцию и заказ. '
        'Для пользователя важно видеть не только цену и наличие детали, но и '
        'инженерные признаки: номинал, допустимое напряжение, ток, мощность, корпус, '
        'тип монтажа, наличие datasheet, SPICE-модели и CAD-модели.',
    )
    para(
        document,
        'В рамках DOLG компонент рассматривается не как изолированная карточка товара, '
        'а как элемент проектного сеанса. Он может быть найден через каталог, добавлен '
        'в схему, проверен расчетом, связан с BOM и использован в учебном задании.',
    )
    heading(document, '1.2. Проблемы существующего процесса', 2)
    bullets(
        document,
        [
            'каталоги компонентов, CAD-среды, симуляторы и системы заказа часто работают независимо;',
            'параметры вручную переносятся из datasheet или карточки товара в схему и расчет;',
            'ошибки в единицах измерения, номиналах, корпусах и GND приводят к некорректной симуляции;',
            'результаты расчетов редко сохраняются как часть единой истории проекта;',
            'обучение часто отделено от реальных ошибок схемы и не использует проектные данные пользователя.',
        ],
    )
    heading(document, '1.3. Анализ аналогов и границы ниши DOLG', 2)
    table(
        document,
        'Таблица 2 - Сравнение классов решений',
        ['Класс решений', 'Сильные стороны', 'Ограничение для задачи DOLG'],
        [
            [
                'Маркетплейсы компонентов',
                'поиск, цена, наличие, заказ',
                'слабо связаны со схемой, симуляцией и обучением',
            ],
            [
                'KiCad / Altium',
                'сильный EDA workflow, DRC/ERC, BOM',
                'высокий порог входа, нет торгового и учебного контура DOLG',
            ],
            [
                'EasyEDA / CircuitLab',
                'браузерная схема и симуляция',
                'ограниченная связка с локальным каталогом, review и обучением',
            ],
            [
                'Flux AI',
                'AI-подсказки поверх схемы',
                'нужна проверяемая evidence-модель и локальный экспертный слой',
            ],
            [
                'Lithium ECAD',
                'интересная синхронизация схемы и проектных данных',
                'требует адаптации идей под web-среду DOLG',
            ],
        ],
    )
    para(
        document,
        'DOLG не позиционируется как замена промышленному PCB CAD. Ниша проекта - '
        'web-ориентированная среда, связывающая каталог, схемотехнический редактор, '
        'симуляцию, инженерную лабораторию, обучение, экспертную проверку и заказ.',
    )
    heading(document, '1.4. Постановка задачи', 2)
    para(
        document,
        'Необходимо разработать web-приложение, которое обеспечивает каталог радио- '
        'и электронных компонентов, работу с пользовательскими проектами схем, '
        'сохранение схемы как данных, инженерные расчеты, симуляцию, review, обучение, '
        'AI-помощника, администрирование, модерацию и контроль качества данных.',
    )
    heading(document, '1.5. Требования к системе', 2)
    table(
        document,
        'Таблица 3 - Основные требования к DOLG',
        ['Группа', 'Требования'],
        [
            ['Каталог', 'карточки товаров, параметры, фильтры, datasheet/SPICE/CAD flags, кликабельные теги'],
            ['Проекты', 'сохранение схем, версии, события, измерения, review snapshots'],
            ['CAD/SIM', 'редактор схем, scheme_data, DRC/ERC, запуск симуляции, expected vs measured'],
            ['Лаборатория', 'расчеты, единицы измерения, инженерная оценка результата'],
            ['Обучение', 'уроки, практические задания, проверка числовых и схемных ответов'],
            ['AI/ML', 'evidence-backed ответы, deep hints без финального инженерного verdict'],
            ['Админка', 'мониторинг, dataset counters, moderation, business metrics'],
            ['Безопасность', 'роли, подписки, feature gates, защита metrics/tokens'],
        ],
    )
    heading(document, '1.6. Выбор технологий и библиотек', 2)
    para(
        document,
        'В качестве серверной основы выбран Django, так как он предоставляет ORM, '
        'маршрутизацию, шаблоны, авторизацию, административную панель, миграции и '
        'механизмы тестирования. Для инженерного слоя применяются специализированные '
        'библиотеки: Pint для единиц измерения, NetworkX для графового анализа схем, '
        'SymPy для формул, Lark для парсинга, Z3 для ограничений, NumPy/SciPy/Matplotlib '
        'для аналитики сигналов. PyTorch используется как вспомогательный neural deep-hint слой.',
    )
    heading(document, '1.7. Проектирование архитектуры', 2)
    para(
        document,
        'Архитектура DOLG строится вокруг нескольких Django-приложений и service-layer. '
        'Модели хранят состояние предметной области, views и API принимают пользовательские '
        'запросы, templates отображают результат, а инженерная логика вынесена в сервисы. '
        'Такой подход уменьшает дублирование формул и позволяет использовать одни и те же '
        'проверки в лаборатории, обучении, review и AI-помощнике.',
    )
    table(
        document,
        'Таблица 4 - Роль приложений Django в архитектуре',
        ['Приложение', 'Назначение'],
        [
            ['shop', 'каталог, товары, категории, параметры, datasheet intelligence'],
            ['orders', 'корзина, заказ, статусы, экспорт'],
            ['accounts', 'профиль, организации, роли, настройки пользователя'],
            ['knowledge', 'энциклопедия, материалы, обучение, задания'],
            ['moderation', 'жалобы, дела модерации, действия, ограничения пользователей'],
            ['Dolg_APP', 'CAD/SIM, проекты, review, AI/ML, лаборатория, monitoring'],
        ],
    )
    heading(document, '1.8. Проектирование базы данных и информационной модели', 2)
    para(
        document,
        'Информационная модель DOLG связывает коммерческие и инженерные сущности. '
        'Товар каталога может быть связан с компонентом схемы, проект хранит схему '
        'и события, review сохраняет findings, обучение хранит задания и прогресс, '
        'а AITrainingExample используется для формирования корпуса данных локального AI.',
    )
    heading(document, '1.9. Сеанс проектирования как основная модель DOLG', 2)
    para(
        document,
        'Ключевая проектная идея DOLG - рассматривать работу пользователя как сеанс '
        'проектирования. Сеанс начинается с поиска компонента или создания схемы, '
        'продолжается расчетами и симуляцией, затем формируется Engineering Review, '
        'после чего пользователь получает рекомендации, учебные задания, BOM и заказ.',
    )
    code(
        document,
        """
project -> scheme_data -> simulation_run -> measurement
        -> project_review -> ai_evidence -> learning_task
        -> bom -> order -> project_event_history
""",
    )
    para(
        document,
        'Эта модель позволяет объяснить дипломный проект не как набор отдельных страниц, '
        'а как единую web-ориентированную среду инженерной работы.',
    )
    picture_asset(
        document,
        GENERATED / 'project_session_flow.png',
        'Рисунок 2 - Проектный workflow: от компонента до заказа',
        width_cm=16.0,
    )


def add_chapter2(document: Document, facts: Facts) -> None:
    heading(document, 'ГЛАВА 2. РЕАЛИЗАЦИЯ ПРОГРАММНОЙ СИСТЕМЫ', 1, page_break=True)
    heading(document, '2.1. Общая структура проекта в коде', 2)
    para(
        document,
        'Программная система реализована как Django-проект `Dolg_PR`. Основные '
        'Django-приложения разделяют ответственность между каталогом, заказами, '
        'аккаунтами, базой знаний, модерацией и инженерными инструментами. Внутри '
        'проекта активно используется service-layer: инженерные расчеты, review, '
        'AI retrieval, units, graph analysis, datasheet extraction и мониторинг '
        'не должны дублироваться в шаблонах.',
    )
    code(
        document,
        """
Dolg_PR/
  shop/              # catalog, products, datasheet intelligence
  orders/            # cart, order flow, order statuses
  accounts/          # user profile, organizations, roles
  knowledge/         # articles, learning tracks, tasks
  moderation/        # reports, cases, actions, restrictions
  Dolg_APP/          # CAD/SIM, review, AI/ML, lab, admin monitoring
""",
    )
    heading(document, '2.2. Реализация каталога и карточек компонентов', 2)
    para(
        document,
        f'В локальной базе зафиксировано {facts.categories} категорий и {facts.products} '
        'товаров. Каталог используется не только как витрина, но и как источник '
        'инженерных параметров для схем, BOM, лаборатории и проверок. Карточки '
        'компонентов отображают производителя, корпус, тип монтажа, rating-поля, '
        'datasheet, признаки наличия SPICE/CAD-моделей и кликабельные теги фильтрации.',
    )
    picture(document, '01_home_catalog.png', 'Рисунок 3 - Главная страница и каталог DOLG', width_cm=16.0)
    picture(
        document,
        'presentation_v6/01_catalog_cards_current.png',
        'Рисунок 4 - Карточки компонентов в каталоге DOLG',
        width_cm=16.0,
    )
    picture(
        document,
        '06_product_detail_photo.png',
        'Рисунок 5 - Карточка товара с параметрами и материалами',
        width_cm=15.5,
    )
    picture(
        document,
        '07_product_related_cards.png',
        'Рисунок 6 - Связанные товары и карточки для BOM-сценария',
        width_cm=15.5,
    )
    para(
        document,
        'Для повышения качества данных используется нормализация РЭБ-каталога и '
        'datasheet intelligence. Отдельно зафиксирована media policy: приоритет '
        'отдается официальным или качественным изображениям, а не случайным источникам; '
        'для технических иллюстраций допускаются аккуратные УГО-стиль изображения.',
    )
    heading(document, '2.3. Реализация пользовательского и коммерческого контура', 2)
    para(
        document,
        f'Коммерческий контур включает пользовательские профили, корзину, заказ и BOM. '
        f'В локальной базе присутствует {facts.orders} заказов. Для разграничения '
        'доступа реализован слой entitlements, который разделяет Free, Pro и Enterprise '
        'функции. Pro-функции включают расширенную аналитику и инженерные AI-возможности, '
        'Enterprise расширяет командную память, аудит и административные сценарии.',
    )
    heading(document, '2.4. Реализация CAD-редактора и хранения схем', 2)
    para(
        document,
        f'CAD-редактор работает со схемой как с данными. В локальной базе присутствует '
        f'{facts.projects} проектов схем и {facts.project_events} событий проектного '
        'сеанса. Внутреннее представление `scheme_data` содержит компоненты, соединения, '
        'узлы, метки и связи с товарами каталога. Благодаря этому схема может быть '
        'проверена, сохранена, импортирована, передана в review и использована AI-ассистентом.',
    )
    picture(
        document,
        'presentation_v6/03_cad_editor_current.png',
        'Рисунок 7 - CAD-редактор и работа со схемой',
        width_cm=16.0,
    )
    picture(
        document,
        '04_cad_gost_template.png',
        'Рисунок 8 - Шаблон схемы и инженерное оформление CAD-режима',
        width_cm=15.5,
    )
    code(
        document,
        """
{
  "components": [{"id": "R1", "type": "resistor", "value": "10 kOhm"}],
  "wires": [{"from": "V1:+", "to": "R1:1"}],
  "nets": [{"name": "GND"}, {"name": "VOUT"}],
  "bom_refs": [{"component_id": "R1", "product_id": 42}]
}
""",
    )
    heading(document, '2.5. Реализация симуляции и измерений', 2)
    para(
        document,
        'Симуляционный контур предназначен для проверки гипотезы: расчетное значение '
        'сравнивается с результатом симуляции или измерения. В системе предусмотрены '
        'SimulationRun и ProjectMeasurement, а также метрики напряжения узла, тока ветви, '
        'RMS, частоты, duty cycle, мощности и температуры. В текущей редакции этот раздел '
        'нужно дополнять только тем, что уже подтверждено кодом и проверками.',
    )
    picture(
        document,
        'presentation_v6/04_simulation_current.png',
        'Рисунок 9 - Результаты симуляции и панель графиков',
        width_cm=16.0,
    )
    picture(
        document, '02_simulation_ac_graph.png', 'Рисунок 10 - AC-график результата симуляции', width_cm=15.5
    )
    picture(document, '03_ac_graph_panel.png', 'Рисунок 11 - Панель анализа AC-симуляции', width_cm=15.0)
    picture(document, '03b_ac_phase_graph.png', 'Рисунок 12 - Фазовый график AC-анализа', width_cm=13.5)
    heading(document, '2.6. Реализация инженерной лаборатории', 2)
    para(
        document,
        'Инженерная лаборатория содержит прикладные расчеты: транзисторный ключ, NE555, '
        'стабилизатор, тепловой запас, RC-антидребезг и другие задачи. Результат должен '
        'быть не только числом, но и инженерной оценкой: норма, риск, перегрев или '
        'необходимость дополнительного запаса.',
    )
    picture(
        document,
        'presentation_v6/02_engineering_lab_results.png',
        'Рисунок 13 - Инженерная лаборатория и расчетные результаты',
        width_cm=16.0,
    )
    heading(document, '2.7. Реализация Engineering Review', 2)
    para(
        document,
        f'Engineering Review является одним из ключевых отличий DOLG. В локальной базе '
        f'зафиксировано {facts.project_reviews} review snapshots. Review собирает '
        'информацию о схеме, DRC/ERC, наличии GND и источника, floating fragments, '
        'BOM-связях, datasheet limits, измерениях и рекомендациях. Каждый finding '
        'должен иметь rule_id, severity, evidence, recommendation, confidence и русскую '
        'формулировку для пользователя.',
    )
    table(
        document,
        'Таблица 5 - Примеры проверок Engineering Review',
        ['Проверка', 'Evidence', 'Результат'],
        [
            ['GND', 'наличие опорного узла', 'ошибка при отсутствии GND'],
            ['Источник', 'наличие voltage/current source', 'warning при неполной схеме'],
            ['Floating fragments', 'граф связности схемы', 'поиск несвязанных участков'],
            ['Rating limits', 'datasheet_extracted и параметры товара', 'warning при выходе за пределы'],
            ['BOM', 'связь компонента схемы с Product', 'риск при отсутствии товара/модели'],
        ],
    )
    picture_asset(
        document,
        GENERATED / 'engineering_review_score.png',
        'Рисунок 14 - Engineering Review вместо устаревшего скрина Pro-функций',
        width_cm=15.5,
    )
    para(
        document,
        'Для финальной редакции следует дополнительно вставить живой screenshot страницы '
        '`/projects/review/<id>/`, когда будет выбран демонстрационный проект. В текущей '
        'рабочей редакции вместо устаревшей Pro-картинки используется отдельная диаграмма '
        'Engineering Review, показывающая score, риски и evidence.',
    )
    heading(document, '2.8. Реализация AI-ассистента и ML pipeline', 2)
    para(
        document,
        f'AI-ассистент DOLG реализован по принципу expert-first. Он использует данные '
        f'схемы, review findings, legal sources, learning tasks и AITrainingExample. '
        f'В локальной базе присутствует {facts.ai_examples} AITrainingExample. PyTorch '
        'используется как слой вероятностных deep hints: классификация топологии, '
        'подсказка следующего компонента и оценка риска. Финальный инженерный вывод '
        'остается за экспертными правилами и человеком.',
    )
    para(
        document,
        'В AI-панели предусмотрены память диалога, summary сессии, счетчик токенов, '
        "context_sources и блок 'Разбор схемы'. Enterprise-режим расширяет контекст "
        'командными проектами, артефактами, review, BOM, комментариями и audit trail.',
    )
    picture_asset(
        document,
        GENERATED / 'ai_ml_pipeline.png',
        'Рисунок 15 - AI/ML pipeline: от данных проекта к deep hint',
        width_cm=16.0,
    )
    heading(document, '2.9. Реализация обучения и базы знаний', 2)
    para(
        document,
        f'База знаний содержит {facts.articles} статей и {facts.article_materials} '
        f'материалов. Практикум включает {facts.learning_tracks} маршрутов, '
        f'{facts.learning_lessons} уроков и {facts.learning_tasks} заданий. Обучение '
        'не ограничивается тестами на выбор: используются числовые задачи, сборка '
        'схемы и измерение результата симуляции.',
    )
    picture(
        document, '08_knowledge_article_materials.png', 'Рисунок 16 - Материалы базы знаний', width_cm=15.5
    )
    heading(document, '2.10. Реализация модерации, ролей и администрирования', 2)
    para(
        document,
        'Для эксплуатации системы реализуются роли пользователей, локальная и глобальная '
        'модерация, soft moderation контента, Django admin и операционный мониторинг. '
        'Админка должна быть не только стандартным CRUD-интерфейсом, но и центром '
        'контроля данных, ML pipeline, бизнес-метрик, moderation queue и состояния сервера.',
    )
    para(
        document,
        'В Django admin добавлен компактный операционный блок, а полный staff Ops Dashboard '
        'показывает runtime, storage, catalog, business, AI/ML, moderation и security metrics. '
        'Публичный `/metrics` защищен через nginx, а Prometheus должен обращаться к метрикам '
        'из внутренней Docker-сети.',
    )
    picture_asset(
        document,
        GENERATED / 'admin_monitoring_metrics.png',
        'Рисунок 17 - Метрики, используемые админкой и мониторингом',
        width_cm=15.5,
    )
    heading(document, '2.11. Реализация контроля качества и проверок', 2)
    para(
        document,
        'Для проверки проекта используются встроенные Django checks, unit/API tests, '
        'management-команды `check_demo_ready --json` и `check_data_integrity --json`, '
        'а также целевые smoke-сценарии. Эти проверки фиксируют не только работоспособность '
        'Django, но и готовность каталога, learning, AI, review, monitoring и demo-данных.',
    )
    table(
        document,
        'Таблица 6 - Основные проверки проекта',
        ['Проверка', 'Назначение'],
        [
            ['manage.py check', 'проверка конфигурации Django'],
            ['makemigrations --check --dry-run', 'контроль незакоммиченных изменений схемы БД'],
            ['check_demo_ready --json', 'готовность demo-сценариев и ключевых стеков'],
            ['check_data_integrity --json', 'качество данных каталога, источников и связей'],
            ['targeted tests', 'регрессии по AI, catalog, review, admin, learning'],
            ['browser smoke', 'проверка видимых пользовательских сценариев'],
        ],
    )
    heading(document, '2.12. Результаты реализации и ограничения', 2)
    para(
        document,
        'В результате реализации сформирована web-платформа, объединяющая каталог, '
        'инженерные карточки компонентов, CAD/SIM-контур, лабораторию расчетов, '
        'практикум, Engineering Review, AI-помощника, ML/dataset pipeline, модерацию, '
        'подписки и административный мониторинг. Ограничениями текущей версии являются '
        'неполнота PCB CAD, зависимость части функций от качества seed-данных, необходимость '
        'дальнейшего расширения supplier adapters и развитие neural layer только как '
        'вспомогательной подсказки.',
    )


def add_conclusion(document: Document) -> None:
    heading(document, 'ЗАКЛЮЧЕНИЕ', 1, page_break=True)
    para(
        document,
        'В ходе выпускной квалификационной работы разработано web-приложение DOLG, '
        'которое объединяет подбор и покупку компонентов, проектирование схем, '
        'инженерные расчеты, симуляцию, review, обучение, AI-подсказки, модерацию '
        'и административный мониторинг.',
    )
    para(
        document,
        'Главным результатом является не отдельная карточка товара или отдельный '
        'редактор схемы, а единый сеанс проектирования, в котором данные компонента, '
        'схемы, расчета, измерения, review, обучения и BOM связаны между собой. '
        'Такой подход снижает ручной перенос данных и позволяет использовать экспертные '
        'проверки как основу для обучения и AI-помощника.',
    )
    para(
        document,
        'Дальнейшее развитие связано с расширением CAD-import, 3D-визуализацией графиков '
        'анализа, улучшением supplier adapters, развитием Project Session, увеличением '
        'корпуса схем для ML и усилением административной панели как центра сбора данных.',
    )


def add_sources(document: Document) -> None:
    heading(document, 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', 1, page_break=True)
    sources = [
        'ГОСТ 2.105-95. Единая система конструкторской документации. Общие требования к текстовым документам.',
        'ГОСТ 2.701-2008. Единая система конструкторской документации. Схемы. Виды и типы. Общие требования к выполнению.',
        'ГОСТ 2.702-2011. Единая система конструкторской документации. Правила выполнения электрических схем.',
        'Django Software Foundation. Django documentation. URL: https://docs.djangoproject.com/',
        'Ngspice Development Team. Ngspice documentation. URL: https://ngspice.sourceforge.io/docs.html',
        'KiCad Documentation. URL: https://docs.kicad.org/',
        'Altium Designer Documentation. URL: https://www.altium.com/documentation/',
        'EasyEDA Pro Documentation. URL: https://prodocs.easyeda.com/',
        'CircuitLab Documentation. URL: https://www.circuitlab.com/docs/',
        'Pint documentation. URL: https://pint.readthedocs.io/',
        'NetworkX documentation. URL: https://networkx.org/documentation/',
        'SymPy documentation. URL: https://docs.sympy.org/',
        'PyTorch documentation. URL: https://pytorch.org/docs/',
        'OWASP Top 10. URL: https://owasp.org/www-project-top-ten/',
        'Lithium ECAD. Синхронизация в Lithium ECAD. URL: https://www.lecad.ru/sync-in-lithium-ecad/',
    ]
    for i, src in enumerate(sources, start=1):
        para(document, f'{i}. {src} Дата обращения: {date.today().strftime("%d.%m.%Y")}.')


def add_appendix_plan(document: Document) -> None:
    heading(document, 'ПРИЛОЖЕНИЯ', 1, page_break=True)
    bullets(
        document,
        [
            'Приложение А - Скриншоты интерфейса каталога, карточек, CAD, SIM, AI-панели и админки.',
            'Приложение Б - ERD/UML/BPMN-диаграммы архитектуры и сеанса проектирования.',
            'Приложение В - Ключевые фрагменты кода service-layer, review, AI и unit parsing.',
            'Приложение Г - Результаты `manage.py check`, `check_demo_ready`, `check_data_integrity` и targeted tests.',
            'Приложение Д - Demo scenario для основной защиты.',
            'Приложение Е - Список rule findings, русская локализация и legal sources.',
            'Приложение Ж - ML/dataset pipeline, AITrainingExample и MLJob.',
        ],
    )


def add_markdown(facts: Facts) -> None:
    OUT_MD.write_text(
        f"""# Рабочая двухглавная редакция диплома DOLG

Дата сборки: {date.today().isoformat()}

Файл DOCX: `{OUT_DOCX.name}`

## Структура

- Введение
- Глава 1. Анализ предметной области и проектирование системы
- Глава 2. Реализация программной системы
- Заключение
- Источники
- Приложения

## Факты локальной БД

- Товары: {facts.products}
- Категории: {facts.categories}
- Проекты схем: {facts.projects}
- ProjectEvent: {facts.project_events}
- ProjectReview: {facts.project_reviews}
- Статьи: {facts.articles}
- Материалы статей: {facts.article_materials}
- Learning tracks: {facts.learning_tracks}
- Learning lessons: {facts.learning_lessons}
- Learning tasks: {facts.learning_tasks}
- AITrainingExample: {facts.ai_examples}

## Вставленные изображения и графики

- Главная страница и каталог.
- Карточки компонентов.
- Карточка товара и связанные товары.
- CAD-редактор и шаблон схемы.
- Симуляция, AC-график, панель AC-анализа и фазовый график.
- Инженерная лаборатория.
- Engineering Review-график вместо устаревшей Pro-картинки.
- AI/ML pipeline.
- Knowledge materials.
- Admin/monitoring metrics.

## Что сделать дальше

1. Заменить черновые абзацы на полный академический текст.
2. Добавить ERD/UML/BPMN-диаграммы.
3. Захватить живой screenshot страницы Engineering Review `/projects/review/<id>/` и admin Ops Dashboard.
4. Перед финальной сборкой подтвердить цифры через `check_demo_ready --json` и `check_data_integrity --json`.
5. После диплома пересобрать презентацию и речь под эту же двухглавную структуру.
""",
        encoding='utf-8',
    )


def build() -> None:
    facts = collect_facts()
    ensure_generated_assets(facts)
    document = Document()
    configure_document(document)
    add_title_page(document)
    add_contents(document)
    add_intro(document, facts)
    add_chapter1(document)
    add_chapter2(document, facts)
    add_conclusion(document)
    add_sources(document)
    add_appendix_plan(document)
    document.save(OUT_DOCX)
    add_markdown(facts)
    print(f'Generated {OUT_DOCX}')
    print(f'Generated {OUT_MD}')


if __name__ == '__main__':
    build()

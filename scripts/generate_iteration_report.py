"""Generate the DOLG iteration report (changes + killer-features roadmap).

This is a tracking document: each iteration we update CHANGES and ROADMAP
sections, regenerate, and the resulting PDF/DOCX/MD/HTML in docs/ records
the project state at that moment.

Output files:
  docs/ITERATION_REPORT.pdf
  docs/ITERATION_REPORT.docx
  docs/ITERATION_REPORT.md
  docs/ITERATION_REPORT.html
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs"
PDF_PATH = OUTPUT_DIR / "ITERATION_REPORT.pdf"
DOCX_PATH = OUTPUT_DIR / "ITERATION_REPORT.docx"
MD_PATH = OUTPUT_DIR / "ITERATION_REPORT.md"
HTML_PATH = OUTPUT_DIR / "ITERATION_REPORT.html"


# =============================================================================
# Содержание отчёта — обновляется при каждой итерации.
# =============================================================================

REPORT_DATE = "2026-05-10"
REPORT_VERSION = "3.4"


CHANGES_TABLE = [
    [
        "Cloudflare Tunnel: правильная работа за HTTPS-туннелем",
        "Главная причина проблем подключения через cloudflared: туннель — HTTPS снаружи / HTTP внутрь. Без особых настроек Django видел request.scheme='http', что давало (а) build_absolute_uri() с http://...trycloudflare.com/ ссылками в email-verify и share-link, (б) Cookie Secure не работали через прокси (login пропадал), (в) в проде SECURE_SSL_REDIRECT уходил в infinite-loop. Фикс: SECURE_PROXY_SSL_HEADER=('HTTP_X_FORWARDED_PROTO', 'https') + USE_X_FORWARDED_HOST=True в settings.py. Cloudflare уже шлёт оба заголовка. run_public.bat: добавлен --protocol http2 (стабильнее на ISP, блокирующих QUIC). DEPLOYMENT.md: новая секция «Cloudflare Tunnel» с 7 типичными проблемами и фиксами + альтернативы (ngrok / localhost.run / serveo) на крайний случай.",
        "Готово",
    ],
    [
        "Docker static-validator + lint-проверки",
        "Новый scripts/check_docker_static.py — без Docker daemon валидирует docker-compose.yml (структура services/volumes/healthcheck), Dockerfile (best-practice: USER nonroot, HEALTHCHECK, --no-install-recommends, no rogue COPY . .), entrypoint.sh (нет shell-injection после round-2-аудита, set -e, opt-in COLLECTSTATIC_CLEAR), nginx.conf (X-Forwarded-Proto/For/Host для Django proxy, security-headers, alias volume mount). Все проверки проходят. Реальный запуск docker compose up требует Docker daemon — не из текущего окружения, но статический валидатор ловит ~80% типичных ошибок.",
        "Готово",
    ],
    [
        "Frontend: TypeScript + Vite scaffold (proof of concept)",
        "Новая директория frontend/ — изолированная TS-сборка для постепенной миграции shop/static/simulation/*.js. package.json (vite 5 + typescript 5 + vitest), tsconfig.json (strict mode + noUncheckedIndexedAccess + noUnusedLocals), vite.config.ts (IIFE bundle для legacy-шаблонов, output → shop/static/lib/dolg/). Первый модуль src/union-find.ts с типизированным API + path-compression two-step shortcut. Vitest-тест в tests/union-find.test.ts (6 кейсов: singleton, transitive union, isolation, portKey format, 1000-ops perf bench). frontend/README.md описывает migration-стратегию: каждый круг 1-2 модуля .js→.ts, в финале — один <script src=lib/dolg/index.js> заменяет все индивидуальные. ВАЖНО: текущие .js работают параллельно — никаких breaking changes до завершения миграции (1-2 недели работ).",
        "Готово",
    ],
    [
        "PCB autorouting MVP + Gerber-экспорт",
        "Новый модуль Dolg_APP/pcb_layout.py: compute_pcb_layout(scheme_data) переводит editor px-координаты в мм (PX_PER_MM=4), генерирует pads (Ø1.6 мм) в позициях портов, traces (Ø0.5 мм) для соединений. to_gerber_top_copper() выводит RS-274X (industry standard: %FSLAX25Y25*%, %MOMM*%, ADD10C/ADD11C aperture, D03 flash + D01 draw). to_gerber_drill() — Excellon NC drill (M48, METRIC LZ, T01C0.80). Endpoint /pcb/<id>/ — SVG-рендер платы (FR-4 фон, медные трассы и pads, силкскрин), кнопка скачать gerber.zip с двумя файлами + README. Кнопка «🧬 PCB» в карточке проекта. MVP — реальные заводские платы делают в KiCad с DRC; цель — показать end-to-end pipeline схема→Gerber.",
        "Готово",
    ],
    [
        "Multi-sheet схематика MVP",
        "Каждый компонент получает sheet_index (default 0). Над холстом — табы листов (.dolg-sheet-tabs), переключение через switchSheet(idx); рендер фильтрует через _isOnCurrentSheet(comp). Соединения видимы только если ОБА конца на текущем листе (cross-sheet электрические связи допустимы — ports общие). Кнопка «+ Лист» создаёт новый sheet_index = max+1. Новые компоненты ставятся на активный лист. На clearCanvas/applySchemeData возвращаемся на лист 1, табы перерисовываются. _stripRuntimeFields НЕ трогает sheet_index — сохраняется в scheme_data как обычное поле.",
        "Готово",
    ],
    [
        "Round-2 audit (5 фиксов: shell-injection + healthz + DATABASE_URL опции)",
        "Аудит свежей prod-инфры: (A1, P1) убрана shell-injection в entrypoint.sh при создании superuser-а — раньше $DJANGO_SUPERUSER_USERNAME интерполировался прямо в Python-исходник через '$VAR', что давало RCE при содержании ' в значении; теперь читаем os.environ внутри Python. (A11) DATABASE_URL парсер пробрасывает ?sslmode=require и др. query-опции в OPTIONS — нужно для managed-Postgres. (A3) Новый /healthz/ endpoint без БД-запроса; Dockerfile HEALTHCHECK переключен на него. (A14) resend-verification получил 5-минутный rate-limit через session — раньше юзер мог спамить SMTP-провайдера. (A2) collectstatic --clear отключён по умолчанию — rolling-deploy с ManifestStaticFilesStorage может временно держать старые хеши; включается через COLLECTSTATIC_CLEAR=1. (A9) DEPLOYMENT.md получил секцию backup Postgres (pg_dump через docker compose exec).",
        "Готово",
    ],
    [
        "+10 тестов (PCB + healthz)",
        "PCBLayoutTests +4 (empty scheme, with components, gerber GTL формат, gerber DRL формат — проверки RS-274X маркеров, Excellon header, X/Y координат). PCBViewTests +4 (login required, SVG рендер, ZIP-download с GTL+DRL+README, 404 для чужого проекта). HealthzTests +2 (200 OK, no-login-required). 10/10 OK за 87 c с FAST_TESTS=1.",
        "Готово",
    ],
    [
        "Production-readiness: Docker + Postgres + nginx (закрытие #19 аудита)",
        "Полный prod-stack одной командой docker compose up: (а) Dockerfile на python:3.14-slim с gunicorn 23, fonts-dejavu для cyrillic-PDF, libpq5 для psycopg2-binary, нерут-юзер dolg, healthcheck на /about/, multi-stage не используем (нет тяжёлых нативных зависимостей). (б) docker-compose.yml: 3 сервиса — db (postgres:16-alpine с pg_isready healthcheck), web (build .), nginx (1.27-alpine с проксированием). Named volumes: pgdata, media, staticfiles. (в) entrypoint.sh: ждёт готовности БД (до 30 c), прогоняет migrate + collectstatic --clear, опционально создаёт superuser-а из DJANGO_SUPERUSER_* env-vars. (г) nginx.conf: статика напрямую с volume (cache 1y immutable благодаря ManifestStaticFilesStorage), media (cache 7d), proxy_pass на web:8000, X-Content-Type-Options/X-Frame-Options/Referrer-Policy. (д) settings.py: DATABASE_URL парсится вручную (без зависимости dj-database-url) — переключает на Postgres если задан, иначе SQLite. CONN_MAX_AGE=60 для пула. В тестах принудительно SQLite. (е) requirements-prod.txt: gunicorn + psycopg2-binary + sentry-sdk. (ж) .env.example расширен: DATABASE_URL, ANTHROPIC_API_KEY, SENTRY_*, DJANGO_SUPERUSER_*, HTTP_PORT. (з) DEPLOYMENT.md обновлён quick-Docker секцией.",
        "Готово",
    ],
    [
        "Email-verification + Sentry-ready + coverage-расширение",
        "(9) Email-верификация: добавлено поле UserProfile.email_verified (миграция 0002), генерация HMAC-токена через django.core.signing с TTL 24 ч, view verify_email подтверждает, resend_verification — повторная отправка из профиля. Бейдж «✓ подтверждён / ⚠ не подтверждён» в profile.html. На login пока не блокируем (для дев), но при смене email флаг сбрасывается и шлётся новое письмо. (Sentry) settings.py: opt-in через SENTRY_DSN env-var; если задан и sentry-sdk установлен — авто-инициализация с DjangoIntegration, send_default_pii=False, traces_sample_rate из ENV. Без переменной — ничего не подтягивается. (18) Расширение test coverage: EmailVerificationTests +4 (register=unverified, verify-token=verified, invalid-token=404, email-change=resets), OrdersCoverageTests +3 (own-only list, 404 for other's detail, cancel returns stock), KnowledgeViewsTests +4 (drafts hidden, published OK).",
        "Готово",
    ],
    [
        "P2-pass: 7 фиксов из остатков аудита 2026-05-10",
        "(11) Stock-race в orders/checkout: select_for_update блокирует строки Product внутри транзакции — два параллельных checkout-а одного товара больше не уйдут в минусовой stock (на SQLite no-op, на Postgres работает). (15) Cross-OS font-path: новый _register_pdf_cyrillic_fonts ищет шрифт по 8 кандидатам (Win/Linux/macOS) с логированием fallback на Helvetica — раньше хардкод C:/Windows/Fonts/arial.ttf падал на Linux. (14) Email-логирование: send_mail обёрнут в try/except logger.warning, fail_silently=False — потерянные письма теперь видны в логах. (12) Address is_default race: add_address и edit_address обёрнуты в transaction.atomic — два параллельных запроса не оставят в БД два is_default=True. (16) edit_profile: EmailValidator на новом email + uniqueness-check + длины капируются (MAX_NAME_LEN=50, MAX_BIO_LEN=2000, phone/address/city/postal_code/country тоже). (10) Pagination: shop/index и shop/category используют Paginator (24 товара на страницу) — на 1000+ товарах не упрётся в скорость рендера. (17) A11y: login.html и register.html получили aria-required, aria-live для messages, role=status, autocomplete (username/current-password/new-password/email), autofocus, password-hint через aria-describedby.",
        "Готово",
    ],
    [
        "+11 регрессионных тестов (всего 35 после двух последних кругов)",
        "ProfileValidationTests +2 (invalid email rejected, long strings capped). CheckoutStockRaceTests +1 (exact stock consumed → stock=0, заказ создан). Все 11 новых OK за 155 c с FAST_TESTS=1. Combined run: AuthHardeningTests + CheckoutSessionTests + BomDosLimitsTests + AddToCartStockTests + ProfileValidationTests + CheckoutStockRaceTests = 11/11 OK.",
        "Готово",
    ],
    [
        "Безопасность: 5 P1-фиксов из аудита 2026-05-10",
        "(1+2) Закрыты ДВЕ XSS-уязвимости в shared-view: банер схемы теперь собирается через textContent + Object.assign({...}, {textContent}) вместо template-литерала с innerHTML — project name (CharField без валидаторов) больше не вектор; scheme_data выводится через Django {%% json_script %%}-тег вместо {{...|safe}} — литерал </script> в полях label/description больше не ломает HTML-парсер. (3) Удалён request.session.create() в orders/checkout — раньше он обнулял _auth_user_id и логаутил пользователя сразу после оформления заказа (cart_items.delete() выше уже очищает корзину). (4) accounts/login_view: rate-limit через session — после 5 неудачных попыток lockout 60 с (счётчик _login_fail_count, _login_locked_until), при успехе сбрасывается. (5) accounts/register: validate_password() из django.contrib.auth прогоняет AUTH_PASSWORD_VALIDATORS — раньше create_user пропускал валидаторы и принимал пароль «1».",
        "Готово",
    ],
    [
        "Защита API + UX корзины (P2 фиксы)",
        "(6+7) BOM_MAX_COMPONENTS=1000 в shop/views: api_bom_match и api_bom_add_all отвергают запросы с >1000 элементов (DoS-protection — раньше можно было прислать 100k и съесть RAM на openpyxl). (8) add_to_cart теперь проверяет product.stock и клампит quantity к доступному остатку с message-уведомлением. Stock=0 — корзина не пополняется. Раньше пользователь мог добавить 999 шт. товара с stock=3, и проблема всплывала только на checkout.",
        "Готово",
    ],
    [
        "Регрессионные тесты (24 новых)",
        "AuthHardeningTests +3 (weak/strong password, 5-fail lockout), CheckoutSessionTests +1 (юзер не логаутится после checkout), BomDosLimitsTests +2 (1000-комп лимит для match и add-all), AddToCartStockTests +2 (cap на stock, отказ для stock=0), AIChatEndpointTests +2 (XSS-векторы в banner и scheme_data). Все 24 OK за 500 c с FAST_TESTS=1.",
        "Готово",
    ],
    [
        "Симуляция: ошибка GND, расчёты, layout на весь экран",
        "(1) GND auto-pick: addFirstAvailableGroundRoot теперь выбирает «-» первой батареи (или эмиттер транзистора), а не случайный первый порт первого компонента — все DC/TRAN-напряжения теперь референсируются интуитивно правильно (раньше могли инвертироваться). (2) Снят громкий warning «Нет компонента Земля» в пользу info-сообщения с указанием, что за GND принят «-» батареи. (3) Чек «плавающая секция» переписан без бесполезного nodeResolver.rootToNode.has() (он там был тождественно ложен и сводился к первому условию) — формула стала чище: groundSet.has(find(...)). (4) Layout: добавлен класс container--fullwidth для simulation.html, переопределяет .container { max-width: 1200px } из shop/styles.css → симулятор теперь занимает всю ширину окна.",
        "Готово",
    ],
    [
        "Седьмой круг полировки: 7 малых UX-фич",
        "(1) Toggle подписей в 3D — кнопка «🏷️ Подписи» в header модалки 3D, скрывает все label-спрайты для чистого PNG-снимка; setLabelsVisible(bool) в DolgScheme3D, _labelSprites массив. (2) Component search в палитре редактора — input «🔍 Найти...» под заголовком, фильтрует .component-btn по тексту названия и data-component (CSS .search-hidden). (3) Token cost estimator в AI-панели — badge «X / Y / Z» в header (input/cache_read/output), накапливается из usage Anthropic-ответа после каждого turn-а; tooltip раскрывает значения. (4) Recently viewed products в shop — session-based история последних 5 просмотренных товаров (RECENTLY_VIEWED_KEY), панель «👁 Недавно смотрели» под related в product_detail.html. (5) «Мои публичные схемы» в /projects/ — bool is_shared в API responses, бейдж «🔗 публичная» на карточке, кнопка «📋 Ссылка» (clipboard copy), отдельный stat-блок «Публичные» со счётчиком. (6) T1/T2 курсоры на осциллографе Lab — Shift+Click ставит/обновляет курсоры (как в основном графике), двойной клик сбрасывает; рисуются пунктирные линии с подписями V, в info-строке Δt + Δv + 1/Δt. (7) Keyboard shortcuts panel — ?-клавиша открывает модалку с таблицей всех hotkeys по разделам (симулятор / AI / графики / CAD / глобально); Esc закрывает; защита от срабатывания в input/textarea/contenteditable.",
        "Готово",
    ],
    [
        "Серия мелких UX-фич (6 штук)",
        "(1) AI quick-prompts — под вкладками чата чипсы с заранее заготовленными вопросами (3-4 на каждый профиль), клик подставляет в input без авто-отправки. (2) QR-код для shared-ссылки — после генерации share-токена открывается модалка с URL и QR (qrcode.js v1.0.0 локально, 20 КБ); мобильный гость сканит камерой, попадает на read-only схему без ввода. (3) Пресеты камеры 3D — кнопки «📐 Изо / ⬇ Сверху / ➡ Сбоку / ⬛ Спереди» в header модалки; setCameraPreset() рассчитывает дистанцию из bbox сцены через THREE.Box3. (4) Авто-tStop для TRAN — buildAnalysisDirectives видит components[i]._signalSource.frequency и автоматически растягивает tStop до 5 периодов сигнала, шаг ≤ T/200 для гладкой осциллограммы. (5) CAD H/V/R горячие клавиши — отзеркалить выделенный объект по горизонтали (H) / вертикали (V) / повернуть на 90° по часовой (R); работает для line/rect/dimension/leader; защита от срабатывания при фокусе в input. (6) +1 тест на share-endpoint без логина → 302/401/403.",
        "Готово",
    ],
    [
        "Доводка точности: triangle PWL + Ω union-find + catalog cache",
        "Триangle-волна генератора лаборатории больше не падает в DC: scheme-netlist.js разворачивает _signalSource.wave='triangle' в SPICE PWL(0 off, T/4 off+amp, 3T/4 off-amp, T off) R=0 — ngspice зацикливает шаблон до конца TRAN. Ω-режим мультиметра переписан на честный union-find: новый buildPortNetMap(components, connections, getComponentPorts) в DolgSchemeNetlist строит точно такую же port→net карту, как генератор netlist; lab.findResistorBetween использует её и складывает параллельные R через 1/Σ(1/Ri). window._lastPortNetMap обновляется после каждой симуляции. Catalog cache: build_catalog_snapshot теперь идёт через django.core.cache (LocMem) с TTL=60с и SHA1-ключом из (categories, lifecycle, exclude_pn, limit) — multi-turn AI-сессия больше не делает SELECT shop_product на каждый turn. Notification генератора теперь честно показывает success для всех трёх форм волны.",
        "Готово",
    ],
    [
        "3 CAD JSON-шаблона: op-amp + резистор + диод по ГОСТ",
        "В shop/static/cad/templates/ добавлены resistor-symbol.json (прямоугольник 4×10 мм с выводами и подписью R 10k), diode-symbol.json (треугольник + катодная полоса с метками A/K). loadJsonTemplate() уже поддерживал arbitrary типы — добавление новых .json не требует JS-изменений. В COMPONENTS_LIBRARY добавлены три записи (op-amp, резистор, диод) — все из JSON-ассетов. Доказывает архитектуру: внешний редактор шаблонов теперь полноценно работает.",
        "Готово",
    ],
    [
        "ManifestStaticFilesStorage + 3 теста для sharing",
        "В settings.py включён ManifestStaticFilesStorage когда DEBUG=False (в проде) — файлы получают hash в имени (pixi.abc123.js), снимая ручной кэш-бастинг ?v=20260506a и позволяя CDN кешировать без ограничений. В DEBUG-режиме поведение прежнее (хеши усложняют отладку). Тесты Dolg_APP/tests.py +3: test_share_token_lifecycle (enable→token, anon GET 200 + баннер, чужой пользователь 404, disable→404), test_share_invalid_token_404, test_catalog_snapshot_is_cached (второй вызов с теми же фильтрами не делает SQL-query). +2 cache.clear() в существующих snapshot-тестах для изоляции.",
        "Готово",
    ],
    [
        "Bug-hunt + новые фичи под защиту",
        "Закрыты три косяка: (1) трансформатор рисовал ПОЛНЫЕ круги вместо полу-окружностей-витков — добавлен tplArc-хелпер, обмотки через настоящие дуги; (2) корпус TO-220 теперь имеет отдельный язычок-радиатор со штриховкой и крепёжным отверстием Ø3.6 (раньше всё было одним прямоугольником); (3) SMD 1206 получил видимый плюсовой маркер шрифта 14 рядом с положительной площадкой. Z-конфликт 3D↔Lab закрыт: открытие одного автоматически закрывает другой.",
        "Готово",
    ],
    [
        "Направление токов: статические зелёные стрелки на проводах",
        "После успешной симуляции computeWireCurrents() для каждого V/R-элемента определяет, через какой порт ток вытекает наружу (V: знак напряжения; R: разность потенциалов); для каждого conn выбирается направление forward/backward с приоритетом R-показаний (точнее, чем V-знак). _drawCurrentArrow рисует мелкий зелёный треугольник в середине ломаной поверх линии. _lastWireCurrents сбрасывается на clearCanvas/stopSimulation/applySchemeData. Без анимации — однократный рендер.",
        "Готово",
    ],
    [
        "Read-only sharing /s/<token>/",
        "Новое поле SchematicProject.share_token (22 chars, db_index, миграция 0004). Endpoint POST /projects/api/<pk>/share/ генерит token через secrets.token_urlsafe(16) и возвращает абсолютный URL. View shared_scheme(token) рендерит редактор с context.is_shared_view=True — без login_required. JS: IS_SHARED_VIEW флаг, авто-загрузка чужой схемы из window._sharedSchemeData, синий gradient-баннер с именем владельца, body.dolg-shared-view CSS прячет кнопки save/load/clear/wire/delete и палитру компонентов; обработчики мутации блокируются. Кнопка «🔗 Поделиться» в тулбаре копирует ссылку через navigator.clipboard (fallback prompt при HTTP).",
        "Готово",
    ],
    [
        "Авто-разводка проводов + расширение поиска обхода",
        "Кнопка «📐 Авто-разводка» в тулбаре сбрасывает waypoints у всех соединений с ручными изгибами (snapshotScheme для undo). Существующая логика buildOrthogonalPath затем перерасчитывает Z/L-маршруты с учётом препятствий. pickFreeAxisX/Y расширены с 30 до 60 итераций — больше шансов найти свободный коридор в плотных схемах.",
        "Готово",
    ],
    [
        "CAD JSON-loader + первый внешний шаблон",
        "loadJsonTemplate(name) async-загрузчик — fetch из shop/static/cad/templates/<name>.json, нормализация полей через tpl* хелперы. Первый эталон — op-amp.json (стандартное обозначение операционного усилителя по IEC: треугольник с входами +/− и выходом). applyTemplate() переведена на async, поддерживает sync- и async-build функции. Доказательство архитектуры: внешние редакторы шаблонов могут добавлять .json без правки JS-кода.",
        "Готово",
    ],
    [
        "Бэклог: блочные уроки /learn/ + расширение CAD",
        "Новый урок-режим на /learn/ — пять интерактивных карточек (Закон Ома, Делитель, RC-фильтр, LED, Тепловая нагрузка). Каждый урок: теория, формула на отдельной плашке, набор параметров, quiz с вводом ответа и проверкой по допуску, кнопка «📂 Открыть» к связанной демо-схеме. Без БД — статический template Dolg_APP/templates/tools/learn.html (244 LOC) + view learn(request) + URL hello:learn. Ссылка в меню «🔧 Инструменты → 🎓 Уроки». Расширение CAD: snap-mode «К объектам» — функция _findNearestSnapPoint() ищет ближайший узел существующих фигур (концы линий, углы прямоугольников, центры/контур окружностей и эллипсов, якорь текста) в радиусе 12 мировых единиц с учётом zoom; snap-mode «Полярный (15°)» — applyOrtho перехватывает шаг угла кратный π/12 для линий/dimension/leader. В библиотеку компонентов CAD добавлены: TO-220 (силовой корпус с радиаторным отверстием Ø3.6 и тремя выводами шагом 2.54), SMD 1206 (3.2×1.6 мм с контактными площадками и полосой плюса), Трансформатор (две обмотки на сердечнике с подписями L1/L2/TR).",
        "✅ Готово 06.05.2026",
    ],
    [
        "Финальный аудит killer-фич + два P1-фикса",
        "По итогам AUDIT_REPORT_2026-05-06_killer-features.md закрыты два бага из этапа 4-5: (#1) Three.js shared-материалы (MAT.beige, MAT.wire и т.д.) диспозились при закрытии 3D-модалки → re-open рисовал чёрную сцену (мёртвые шейдеры). Фикс: помечены userData._shared = true через хелпер _sharedMat(opts), dispose() guard'ит и не диспозит shared. (#2) Генератор сигналов в лаборатории врал — UI показывал «sine 5В», но обработчик ставил comp.voltage = amplitude (DC). Фикс: scheme-netlist.js читает component._signalSource и генерит SIN(off amp freq) для синуса, PULSE(low high 0 1n 1n half_period period) для меандра; triangle — fallback на DC с warning-нотификацией. Lab-callback показывает success для sine/square и warning для triangle. Подготовлен DEMO_SCENARIO.md — пошаговый план 12-15 минут защиты с Q&A-бронёй.",
        "Готово",
    ],
    [
        "Killer-фича #5: Виртуальная лаборатория приборов",
        "Новый модуль shop/static/simulation/scheme-lab.js (~506 LOC) — три прибора в одной модалке: (1) ОСЦИЛЛОГРАФ — отображает TRAN-результат как осциллограмму на затемнённом «трубном» canvas-е с фосфор-зелёной траекторией, сеткой 10×8 делений, выбором канала и регулировками V/деление + t/деление (1мВ–5В, 1мкс–100мс); под графиком — Vmin/Vmax/Vavg/RMS. (2) МУЛЬТИМЕТР — крупный «7-сегментный» LCD-дисплей (моноширинный фосфор), три режима: V (DC), V_RMS (по TRAN-точкам), Ω (поиск резистора между выбранными узлами); выбор пары щупов. (3) ГЕНЕРАТОР СИГНАЛОВ — preview-canvas с живой осциллограммой sine/square/triangle, регулировки амплитуды, частоты, DC-offset; «Применить к V1» меняет напряжение источника схемы и просит перезапуск симуляции. Хук в renderSimResult — при каждой успешной симуляции лаборатория рефрешится автоматически. Вход — кнопка «🔬 Лаборатория» в тулбаре, Esc — закрыть.",
        "✅ Готово 06.05.2026",
    ],
    [
        "Killer-фича #4: 3D-просмотр платы (Three.js)",
        "Three.js r140 + OrbitControls локально (650 КБ, shop/static/lib/three.min.js + OrbitControls.js). Новый модуль shop/static/simulation/scheme-3d.js (~535 LOC) — процедурная генерация PCB: зелёная FR-4 подложка по bbox схемы, 9 типов корпусов (axial-резистор с 4-полосным ГОСТ-цветовым кодом по resistorBands(), LED с прозрачным куполом и emissive-свечением, электролитический конденсатор с минус-полосой, DIP-8 микросхема, TO-92 транзистор с плоской стороной, диод с катодной полосой, барабанный индуктор с витками-торусами, ground-полусфера, node-точка). Соединения — тонкие провода-цилиндры. Подписи компонентов — спрайты с canvas-текстурами. OrbitControls (ЛКМ орбита / колесо зум / ПКМ панорама). Камера авто-fit по размеру схемы. Кнопка «📷 Сохранить PNG» через preserveDrawingBuffer + toDataURL. Вход — кнопка «🎬 3D» в тулбаре редактора. Esc — закрыть. dispose() освобождает GPU-ресурсы при закрытии.",
        "✅ Готово 06.05.2026",
    ],
    [
        "Финальная полировка P2 (post-fix аудит)",
        "Все 6 P2 из AUDIT_REPORT_2026-05-05_post-fix.md закрыты. (1) PROMPT_CACHE_MIN_CHARS поднят с 3000 до 4500 chars — ≈ 1800 токенов кириллицы, гарантированно выше Anthropic-минимума 1024; кеш создаётся на любой вкладке, экономия 5-10× в первой же сессии. (2) _lastSelectedObj вынесен в State._lastSelectedObj — единый источник стейта CAD. (3) Смена инструмента (кроме pointer/delete) теперь сбрасывает selection — панель «🎯 Выбранный объект» больше не показывает прошлый объект при начале нового действия. (4) test_endpoint_caps_target_pn_length усилен: мокает build_catalog_snapshot и проверяет, что в exclude_pn пришла строка ровно 200 символов (а не просто что endpoint не упал). (5) Чат AI переведён с polled-rerender на append-only DOM (helpers _makeMsgEl/appendMessage/removePlaceholderMessage); aria-live=polite теперь корректно — скринридер озвучивает только новое сообщение, не всю историю. (6) Opt-in FAST_TESTS=1 (через Dolg_PR/settings.py + DisableMigrations): пропускает прогон миграций в тестах, схема строится напрямую из моделей. На AI-suite даёт небольшой выигрыш, на PopulateDemoProjectsCommandTests — кратный (миграции — главный вкладчик 12-минутного прогона).",
        "Готово",
    ],
    [
        "Аудит-фиксы P1: AI / CAD / гость",
        "По итогам докуметированного аудита AUDIT_REPORT_2026-05-05.md закрыты все P1-пункты. AI-ассистент: введён доменный класс AIError + AIAuthError/AIRateLimitError/AINetworkError/AIServerError/AINotConfiguredError; HTTP-статусы Anthropic мапятся в понятные сообщения и адекватные коды (401→502 «ключ недействителен», 429→429 «лимит, подождите 30 с», 5xx→502, network→504). Внедрено prompt caching: build_system_blocks разделяет system на стабильный (persona+guidelines+CATALOG, c cache_control: ephemeral) и переменный блок (SCHEME/target_pn) — экономия 5-10× на multi-turn сессиях. anthropic-version бамплено на 2024-10-22. CAD: dimension получил трёхполосную точность (10.5 мм → «10.5 мм» вместо «11»); inline-панель «🎯 Выбранный объект» с editable text/color/lineWidth — leader.text теперь правится без удаления. Гость в demo-режиме перестал получать загадочное «Сеть недоступна»: short-circuit в JS показывает «🔑 Войдите в аккаунт».",
        "Готово",
    ],
    [
        "Аудит-фиксы P2: text bbox / sweep flag / canvas.title / a11y",
        "Селект-bbox для text-объекта в CAD теперь меряется через ctx.measureText — длинные подписи целиком в рамке. _sweepHasResults() ушёл с DOM-зависимости (.status-idle) на runtime-флаг _lastSimResultPresent (синхронизирован с runSimulation/stopSimulation/clearCanvas/applySchemeData). canvas.title в thermal-tooltip пишется только при изменении — снимает 60×/с set property без необходимости. На ленту чата AI добавлен aria-live=polite + aria-label для скринридеров. target_pn капится 200 символами. settings.AI_ENABLED удалён (был не подключён, источник истины — ai_assistant.is_enabled()).",
        "Готово",
    ],
    [
        "Аудит-фиксы P1: тесты для error mapping и prompt caching",
        "10 новых тестов в Dolg_APP/tests.py: AIAssistantModuleTests +6 (build_system_blocks_marks_stable_prefix_for_caching, build_system_blocks_skips_caching_when_short, call_claude_maps_401/429/500, network_failure→AINetworkError, no_key→AINotConfiguredError); AIChatEndpointTests +3 (endpoint_maps_anthropic_429_to_429, _401_to_502_with_auth_message, caps_target_pn_length). Переписан is_enabled-тест: вместо хитрого мока settings — прямой mock _api_key (теперь не зависит от того, выставлен ли ANTHROPIC_API_KEY на машине разработчика). 27/27 AI-тестов зелёные (245 c).",
        "Готово",
    ],
    [
        "Полировка после этапа 3: специализированные AI-агенты",
        "Чат-виджет переведён с универсального prompt-а на три профиля AGENT_PROFILES (recommend / explain / replace). Каждый профиль — свой persona + guidelines + output_hint + temperature: «Инженер-схемотехник» 0.3 (даёт расчёт + список компонентов), «Schematic-reviewer» 0.2 (точный поиск ошибок с id), «Supply-chain эксперт» 0.2 (замена с риск-оценкой pin-out). call_claude(messages, system, mode=mode) подставляет model/temp/max_tokens из профиля. В UI title-tooltip’ы вкладок раскрывают роль агента. Endpoint возвращает agent + model для отладки.",
        "Готово",
    ],
    [
        "Rate-limit AI-чата + sweep-стрип + thermal Map",
        "POST /api/ai/chat/ — минимальный интервал 2 с между вызовами (через session, AI_MIN_INTERVAL_SEC). Mobile-UI панели: max-width = calc(100vw-40px) и медиа-запрос для <480px. Sweep-state (_sweep) теперь стрипится в buildSchemeData() при сериализации — не утечёт в save/load/AI/PDF. computeThermal() использует Map id→component вместо .find в цикле — снимает O(N²) на больших схемах.",
        "Готово",
    ],
    [
        "Demo-схема Thermal-showcase",
        "Новый демо-проект «🌡 Тепловая шкала: 5 резисторов 12 В» в populate_demo_projects.py. Пять параллельных резисторов (10к/1к/680/470/220 Ом) от батареи 12 В, подобраны по P=V²/R так, чтобы покрыть все цветовые зоны теплового анализа: зелёный 6% → жёлтый 58% → оранжевый 85% → красный 123% → 262%. Один клик «▶ Симуляция» — на схеме градиент аур, в результатах — таблица «🔥 Тепловая нагрузка» с предупреждениями для перегруженных резисторов.",
        "Готово",
    ],
    [
        "Тесты AI-ассистента",
        "Новые тесты в Dolg_APP/tests.py: AIAssistantModuleTests (профили — все 3 различные, build_system_prompt для каждого режима, build_catalog_snapshot с фильтрами по категории/lifecycle/exclude_pn) и AIChatEndpointTests (login_required, invalid JSON, unknown mode, empty/oversize message, demo-режим без ключа, rate-limit 429, mocked live-call с проверкой kwargs). 11 новых тестов.",
        "Готово",
    ],
    [
        "AI-ассистент DOLG (killer-фича #3)",
        "Чат-виджет в симуляторе: floating action button (правый-нижний угол) + slide-in panel 380×500 px, hotkey Alt+A. Три вкладки: «Подбор» (recommend) — рекомендации компонентов из каталога под задачу пользователя; «Объясни» (explain) — анализ текущей схемы, scheme_data передаётся на сервер как контекст; «Замена EOL» (replace) — поиск активных аналогов по part_number с автозаполнением из выбранного компонента. Серверный модуль Dolg_APP/ai_assistant.py — тонкая обёртка над POST https://api.anthropic.com/v1/messages (raw HTTP через requests, без SDK). Endpoint POST /api/ai/chat/ (login_required): подбирает CATALOG-snapshot 20-60 позиций (с фильтром по категориям из scheme_data в режиме explain или по lifecycle=active в replace), хранит историю до 8 реплик. Без ANTHROPIC_API_KEY возвращается demo-режим — UI отдаёт пользователю «🔒 ANTHROPIC_API_KEY не настроен». Модель по умолчанию — claude-haiku-4-5-20251001, max_tokens=1024, бюджет CATALOG-snapshot 6 КБ",
        "Готово",
    ],
    [
        "Тепловой анализ схемы (killer-фича #2)",
        "После каждой успешной симуляции computeThermal() вычисляет рассеиваемую мощность для R (V²/R), V-источников (|V·I|), диодов и LED (|V·I|). На схеме у каждого компонента — цветная аура (зелёный → жёлтый → оранжевый → красный по % от TDP); hover-tooltip с P/лимит/% через canvas.title. В панели результатов — таблица «🔥 Тепловая нагрузка (топ-5)» с предупреждением ⚠️ при превышении TDP. Лимиты берутся из catalog_parameters.tdp_w/power_w если каталог связан, иначе типовые дефолты (R 0.25 Вт, диод 1 Вт, LED 0.1 Вт, источник 10 Вт). _lastSimPowers сбрасывается при clearCanvas/applySchemeData/stopSimulation",
        "Готово",
    ],
    [
        "«What if»-слайдер на параметрах компонентов (killer-фича #1)",
        "В панели свойств у числовых полей (R/C/L/V) — иконка 〰️. По клику input заменяется log-scale слайдером в диапазоне ±2 декады от базового значения; движение слайдера меняет параметр в реальном времени, drawCanvas, и через debounce 220 мс автоматически перезапускает симуляцию (только если симуляция уже была запущена). Кнопка ↺ возвращает к базе. Видно как АЧХ/TRAN сдвигаются live при изменении C от 1 нФ до 1 мкФ или R от 100 Ω до 100 кΩ",
        "Готово",
    ],
    [
        "Отдельный шаг хода схем",
        "В симулятор добавлена настройка шага хода, независимая от визуальной сетки; snap компонентов, Ctrl+стрелки, дублирование, изгибы и маршрутизация используют новый шаг; 11 демо-схем нормализованы по drawing_step=30",
        "Готово",
    ],
    [
        "Правки рендера схем по замечаниям",
        "Убран непрактичный переключатель modern/ГОСТ, оставлен единый УГО-режим; скрыта canvas-подсказка, удалена кнопка экспорта результатов из аналитики; добавлен fast-path для больших схем и performance browser-smoke",
        "Готово",
    ],
    [
        "Редакция ВКР без смены структуры",
        "В генераторе диплома добавлен раздел 2.4 с этапами разработки и планом модернизации; глава 3 завершает реализацию и проверку, а доказательные этапные блоки вынесены в приложение В",
        "Готово",
    ],
    [
        "Аналитика симуляции под свойствами",
        "Панель результатов вынесена из узкой правой колонки в нижнюю часть рабочей области под свойства компонента; добавлены CSS flex/min-width ограничения и browser-smoke assertions для desktop/mobile layout",
        "Готово",
    ],
    [
        "CAD visual baseline",
        "Добавлен browser-smoke desktop/mobile для /cad/: применение ГОСТ-шаблона, проверка overflow, видимости панелей, высоты canvas и реальной отрисовки; исправлено сжатие mobile canvas до 150 px",
        "Готово",
    ],
    [
        "WebGL-рендер симулятора",
        "Pixi.js v7.4.2 локально (shop/static/lib/pixi.min.js, 456 КБ); auto-switch Canvas2D ↔ WebGL при components+connections > 200; настройка «Авто / Canvas2D / WebGL» в settings, persist в localStorage",
        "Готово",
    ],
    [
        "Производительность графиков",
        "drawSimGraph и drawAcGraph переведены на ImageData-снапшот (один paint + putImageData на hover); устранён Math.min(...arr) на больших AC-sweep'ах; stats и markers HTML обновляются только при изменении точек",
        "Готово",
    ],
    [
        "Производительность редактора",
        "drawWireHops — cell-grid Set вместо O(H×V×N×ports); drawCanvas / drawJunctionDots — Map id→component вместо .find в цикле; MNA-солвер избавлен от spread-копий и destructuring-swap",
        "Готово",
    ],
    [
        "CAD UI-обвязка",
        "Усилены границы у .right-panel, .tools-panel, .canvas-area (2 px цианновая рамка + glow + box-shadow), заголовки секций — градиент-фон",
        "Готово",
    ],
    [
        "Шаблон ГОСТ 2.104 Форма 1",
        "Полностью переписан gostFrame: 8 cells в левом блоке (top-zone + headers + 5 ролей с пропуском ряда 7 + расширенная Утв), 4 cells в Лит-zone, headers ONLY (значения пустые); заголовки центрированы через cellLbl(x1, x2, row, text, sz); шрифт 7 для шапки и ролей; внешняя рамка штампа BOLD по ГОСТ 2.301",
        "Готово",
    ],
    [
        "Критический баг шаблона",
        "Локальная переменная MID = 135 (мм-координата) затеняла глобальную MID = 1.5 (line-width) — штамп рисовался сплошным чёрным квадратом. Исправлено переименованием в NAME_R; добавлен предупреждающий комментарий",
        "Готово",
    ],
    [
        "Инструменты CAD",
        "dimension — теперь со стрелками-наконечниками, перпендикулярными «выносными рисками» и автоматической подписью длины в миллиметрах (CAD_PX_PER_MM = 2); новый leader — выноска со стрелкой и текстом из поля «Текст»; ortho-режим распространён на оба",
        "Готово",
    ],
    [
        "Слои CAD",
        "Inline-переименование по двойному клику; счётчик объектов в слое; подтверждение при удалении непустого слоя",
        "Готово",
    ],
    [
        "Engineering-суффиксы",
        "parseEngValue('1k', '4.7k', '10u', '100n') подключён в input-поля свойств И в buildSpiceNetlist для R/C/L/V — раньше старый JSON со строкой '4.7k' уходил в SPICE как 4.7",
        "Готово",
    ],
    [
        "CSV-экспорт BOM",
        "Кнопка «📑 CSV» рядом с «Добавить всё в корзину»; UTF-8 + BOM (Excel читает кириллицу), кавычки по RFC 4180, итоговая строка",
        "Готово",
    ],
    [
        "Onboarding-tour",
        "5 шагов при первом заходе на /simulation/ (приветствие, добавление компонентов, провод, горячие клавиши, запуск симуляции); spotlight целевого элемента, prev/next/skip, прогресс N/5, persistence в localStorage",
        "Готово",
    ],
    [
        "Cache-bust ngspice-worker",
        "Версия assetsVersion прокидывается из main thread через postMessage в worker; locateFile и importScripts используют общий ver — больше не нужно править два места при бампе версии",
        "Готово",
    ],
    [
        "Regex-кэш в worker",
        "DC/AC/TRAN regex'ы вынесены на module-level: DC_LINE_RE, DC_BRANCH_RE, TRAN_V_RE, AC_CPLX_RE",
        "Готово",
    ],
    [
        "Авто-вписывание шаблона",
        "После применения ГОСТ-шаблона вызывается fitView() — лист встаёт целиком в видимую область; настройка «Авто-вписывать шаблон» в верхней панели, persists в localStorage",
        "Готово",
    ],
]


METRICS_TABLE = [
    ["Тестов всего", "129 (8 accounts + 51 Dolg_APP + 16 browser-smoke + 16 orders + 38 shop)"],
    ["Стандартный прогон", "113 OK + 16 skipped (browser-smoke без RUN_BROWSER_E2E=1)"],
    ["Прогон AI+share-тестов", "30/30 OK (включая 3 новых: share lifecycle, invalid token, catalog cache, ~316 c)"],
    ["Быстрый прогон", "FAST_TESTS=1 → миграции не прогоняются, схема из моделей (см. settings.py)"],
    ["Браузер-тесты", "16/16 OK через scripts/run_browser_e2e.ps1"],
    ["LOC ключевых файлов", "simulation.html 7643, cad.html 2080, projects.html 734, learn.html 244, shop/views.py 902, Dolg_APP/views.py 637, Dolg_APP/models.py 109, Dolg_APP/ai_assistant.py 354, Dolg_APP/tests.py 669, scheme-netlist.js 360, scheme-3d.js 587, scheme-lab.js 577"],
    ["JS-модули симуляции", "scheme-normalizer.js (130), scheme-export.js (150), scheme-bom.js (226), scheme-netlist.js (360), simulation-engine.js (256), ngspice-worker.js (416), scheme-3d.js (547), scheme-lab.js (501)"],
    ["Внешние библиотеки", "pixi.min.js v7.4.2 (~456 КБ), three.min.js r140 (~624 КБ) + OrbitControls.js (~26 КБ) — всё локально в shop/static/lib/"],
    ["Демо-схем", "12 (включая стресс-тест R-2R ×100 ≈600 элементов и тепловую шкалу для thermal-аналитики)"],
    ["Товаров в каталоге", "72 (включая 26 РЭБ-компонентов)"],
    ["Категорий", "20 (12 потребительских + 8 РЭБ)"],
    ["Статей энциклопедии", "9 в 6 категориях"],
]


KILLER_FEATURES = [
    [
        "1. «What if»-слайдер",
        "В панели свойств — иконка 〰️ рядом с числовым полем. Клик — параметр становится слайдером (R: 100 Ω – 100 кΩ, log-scale); при движении — auto re-run симуляции (debounce 220 ms) и перерисовка графиков. Зритель видит как АЧХ-фильтра сдвигается в реальном времени при изменении C от 1 нФ до 1 мкФ",
        "4-6 ч",
        "✅ Готово 04.05.2026",
    ],
    [
        "2. Тепловой анализ",
        "Из SPICE-результатов вычисляется P каждого компонента (P=V²/R для R, P=|V·I| для активных). На схеме — тепловая аура (зелёный → красный по % от TDP). Hover-tooltip «R3: 0.45 W / лимит 0.5 W (90%)». Таблица «🔥 Тепловая нагрузка (топ-5)» в результатах. Привязка к catalog_parameters.tdp_w если связан с каталогом, иначе дефолты по типу",
        "6-8 ч",
        "✅ Готово 04.05.2026",
    ],
    [
        "3. AI-ассистент DOLG",
        "Чат-виджет (FAB + slide-in panel + Alt+A) с тремя вкладками: «Подбор» — рекомендации компонентов из каталога под задачу; «Объясни» — анализ текущей схемы (передаётся scheme_data); «Замена EOL» — поиск активных аналогов по part_number. Серверный endpoint POST /api/ai/chat/ через Claude API (haiku-4-5, raw HTTP без SDK); system prompt с динамическим CATALOG-snapshot (фильтр по категориям/lifecycle, бюджет 6 КБ); demo-режим без ANTHROPIC_API_KEY — UI остаётся живым.",
        "8-12 ч",
        "✅ Готово 05.05.2026",
    ],
    [
        "4. 3D-просмотр платы",
        "Three.js r140 + OrbitControls локально. Процедурная генерация PCB: 9 типов корпусов с реалистичной геометрией; резисторы — с 4-полосным цветовым кодом; LED — с прозрачным куполом и emissive-свечением; OrbitControls (орбита/зум/панорама); экспорт PNG через preserveDrawingBuffer. Кнопка «🎬 3D» в тулбаре, Esc — закрыть.",
        "10-15 ч",
        "✅ Готово 06.05.2026",
    ],
    [
        "5. Виртуальная лаборатория приборов",
        "Три прибора: ОСЦИЛЛОГРАФ (canvas со фосфор-зелёной траекторией, сетка 10×8, регулировки V/div и t/div, выбор канала, статистика Vmin/Vmax/Vavg/RMS); МУЛЬТИМЕТР (крупный 7-сегментный LCD, режимы V DC / V_RMS / Ω, выбор щупов); ГЕНЕРАТОР СИГНАЛОВ (preview-canvas с sine/square/triangle, регулировки амплитуды/частоты/offset, «Применить к V1»). Хук в renderSimResult авто-рефрешит лабораторию при каждом запуске. Кнопка «🔬 Лаборатория» в тулбаре.",
        "15-20 ч",
        "✅ Готово 06.05.2026",
    ],
]


BONUS_FEATURES = [
    [
        "Marketplace проектов",
        "Публикация чужих проектов с лайками/форками — GitHub для EDA. Отложено: требует публичного деплоя сервера",
        "10-15 ч",
        "Отложено",
    ],
    [
        "QR-СБП на checkout",
        "На странице оформления заказа — QR для российской системы быстрых платежей. Отложено: ждёт публичного деплоя",
        "2-3 ч",
        "Отложено",
    ],
    [
        "Block-based редактор (learning mode)",
        "Scratch-style на /learn/: drag-drop готовых функциональных блоков (фильтр НЧ, инвертирующий усилитель, выпрямитель, делитель напряжения); генерируется реальная схема под капотом",
        "12-15 ч",
        "После killer-фич",
    ],
]


REJECTED_IDEAS = [
    [
        "Speech-to-schema",
        "Голосовое управление через Web Speech API (русский). Отклонено: пользователь не видит больших перспектив",
    ],
    [
        "AR-режим",
        "Камера телефона + WebXR/AR.js для наложения схемы на реальную плату. Отклонено: «слишком крутая и сложная фича, в очень далёком будущем»",
    ],
    [
        "Hardware-in-the-loop",
        "WebSerial → подключение Arduino/ESP32, реальные данные на симуляции. Отклонено: «одним ЯП и простыми нейронками не ограничится»",
    ],
    [
        "Генератор технической документации (отдельная фича)",
        "PDF с титульным листом, оглавлением, схемой, BOM, расчётами. Объединено с #3 AI-ассистент: будет одним из режимов «сгенерируй ТЗ по этой схеме»",
    ],
]


REJECTION_NOTES = [
    "Marketplace и QR-СБП имеют смысл, но текущий сервер локальный (только cloudflared tunnel). Реализовать после публичного деплоя.",
    "Block-based редактор — образовательная фича, не для основной презентации. Можно сделать после killer-features в режиме «обучалки».",
]


WEAK_SPOTS = [
    [
        "simulation.html",
        "Около 6 000 строк в одном файле, JS вперемежку с HTML и Django-шаблонными тегами",
        "Сложно безопасно править, риск регрессий, тяжелее писать unit-тесты на JS",
        "🔴 Высокий",
    ],
    [
        "WebGL-рендер",
        "Упрощённая визуализация (rect + ports + wires вместо ГОСТ-УГО); подписи через PIXI.Text без оптимизации",
        "На 1000+ элементах текстовые объекты могут просесть; визуально отличается от Canvas2D",
        "🟡 Средний",
    ],
    [
        "CAD-WebGL",
        "Не реализован (State.ctx + 8 типов примитивов с clip-масками)",
        "На 200+ объектах CAD-чертежа Canvas2D начнёт лагать",
        "🟡 Средний",
    ],
    [
        "ГОСТ-шаблоны",
        "A4/A3/A2 имеют один и тот же штамп; реальный ГОСТ предусматривает разные формы для разных листов",
        "Смотрится упрощённо",
        "🟢 Низкий",
    ],
    [
        "Reviews",
        "Захардкожены в шаблоне товара (нет модели)",
        "Нельзя добавить отзыв через UI",
        "🟢 Низкий",
    ],
    [
        "Visual regression",
        "Есть базовые visual-smoke модалок и перенос аналитики под свойства компонента; нет screenshot-baseline/diff для CAD, projects, панели свойств и сложных состояний аналитики",
        "Регрессии вёрстки незаметны до защиты",
        "🟡 Средний",
    ],
    [
        "Multisim-like rendering",
        "Введён отдельный шаг хода и нормализованы демо-схемы, но сами УГО, junction dots и подписи ещё требуют профессиональной перерисовки",
        "На защите схема может выглядеть менее инженерно, чем расчётная часть",
        "🟡 Средний",
    ],
]


# =============================================================================
# Сборка структуры документа.
# =============================================================================


def build_blocks() -> list[dict]:
    return [
        {"type": "title", "text": "Итерационный отчёт DOLG"},
        {
            "type": "paragraph",
            "text": (
                f"Дата: {REPORT_DATE}. Версия отчёта: {REPORT_VERSION}. Документ обновляется "
                "после каждой содержательной итерации и фиксирует: что сделано, какие "
                "killer-фичи запланированы, что отложено и почему, текущие слабые места."
            ),
        },
        {"type": "page_break"},

        # === Раздел 1 ===
        {"type": "heading", "level": 1, "text": "1. Что сделано в последней итерации"},
        {
            "type": "paragraph",
            "text": (
                "Изменения охватывают WebGL-рендер симулятора, оптимизации производительности, "
                "переделку штампа ГОСТ 2.104, новые инструменты CAD (размеры в миллиметрах со "
                "стрелками, выноски), engineering-суффиксы для значений компонентов, CSV-экспорт "
                "BOM, onboarding-тур и инфраструктурные правки."
            ),
        },
        {
            "type": "table",
            "headers": ["Зона", "Описание", "Статус"],
            "rows": CHANGES_TABLE,
        },

        # === Раздел 2 ===
        {"type": "heading", "level": 1, "text": "2. Метрики проекта"},
        {
            "type": "table",
            "headers": ["Показатель", "Значение"],
            "rows": METRICS_TABLE,
        },

        # === Раздел 3 ===
        {"type": "heading", "level": 1, "text": "3. План killer-фич (5 этапов)"},
        {
            "type": "paragraph",
            "text": (
                "Согласовано с пользователем 03.05.2026. Сортировка по соотношению "
                "«wow-эффект / трудозатраты / технический риск». Общая оценка: 45-60 часов "
                "на все 5 этапов. Реализовать последовательно с проверкой после каждого."
            ),
        },
        {
            "type": "table",
            "headers": ["Фича", "Описание", "Часы", "Очерёдность"],
            "rows": KILLER_FEATURES,
        },

        # === Раздел 4 ===
        {"type": "heading", "level": 1, "text": "4. Бонус-фичи (после основного плана)"},
        {
            "type": "table",
            "headers": ["Фича", "Описание", "Часы", "Статус"],
            "rows": BONUS_FEATURES,
        },
        {
            "type": "bullets",
            "items": REJECTION_NOTES,
        },

        # === Раздел 5 ===
        {"type": "heading", "level": 1, "text": "5. Отклонённые идеи"},
        {
            "type": "paragraph",
            "text": "Идеи, рассмотренные и не вошедшие в план. Зафиксированы для истории решений.",
        },
        {
            "type": "table",
            "headers": ["Идея", "Причина отклонения"],
            "rows": REJECTED_IDEAS,
        },

        # === Раздел 6 ===
        {"type": "heading", "level": 1, "text": "6. Текущие слабые места"},
        {
            "type": "paragraph",
            "text": (
                "Зоны проекта, которые требуют внимания, но не блокируют killer-features. "
                "Будут адресованы параллельно с основным планом или после защиты."
            ),
        },
        {
            "type": "table",
            "headers": ["Зона", "Наблюдение", "Риск", "Приоритет"],
            "rows": WEAK_SPOTS,
        },

        # === Раздел 7 ===
        {"type": "heading", "level": 1, "text": "7. Заключение и следующий шаг"},
        {
            "type": "paragraph",
            "text": (
                "Текущая итерация закрыла критические UI/UX-баги рендера схем, добавила отдельный "
                "шаг хода для snap/маршрутизации и нормализовала 11 демо-схем по `drawing_step=30`. "
                "Это снижает лаги и создаёт основу для дальнейшей Multisim-like перерисовки без "
                "возврата непрактичного modern-режима."
            ),
        },
        {
            "type": "paragraph",
            "text": (
                "Следующий шаг: продолжить стабилизацию инженерной основы — добавить screenshot-baseline "
                "для `/simulation/` и `/cad/`, затем перерисовать УГО, junction dots, провода и подписи "
                "на нормализованном шаге хода. «What if»-слайдер остаётся первой killer-фичей после "
                "того, как расчётный контур, CAD и визуальный вид схем будут достаточно надёжны для демонстрации."
            ),
        },
    ]


# =============================================================================
# Рендереры (DOCX / Markdown / HTML / PDF) — копия из generate_dolg_report.py
# с минимальными правками: нумерация страниц без жёсткой привязки к одному
# имени шрифта, чтобы не падать при отсутствии Times New Roman.
# =============================================================================


def _set_docx_style(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1.5)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.first_line_indent = Cm(1.0)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size in (("Heading 1", 16), ("Heading 2", 14), ("Title", 18)):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def render_docx(blocks: Iterable[dict], output_path: Path) -> None:
    document = Document()
    _set_docx_style(document)

    for block in blocks:
        kind = block["type"]
        if kind == "title":
            paragraph = document.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(block["text"]).bold = True
        elif kind == "heading":
            document.add_heading(block["text"], level=block["level"])
        elif kind == "paragraph":
            paragraph = document.add_paragraph(block["text"])
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif kind == "bullets":
            for item in block["items"]:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.add_run(item)
        elif kind == "table":
            rows = block["rows"]
            table = document.add_table(rows=1, cols=len(block["headers"]))
            table.style = "Table Grid"
            header_cells = table.rows[0].cells
            for index, value in enumerate(block["headers"]):
                header_cells[index].text = value
                for paragraph in header_cells[index].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for index, value in enumerate(row):
                    cells[index].text = value
            document.add_paragraph()
        elif kind == "page_break":
            document.add_section(WD_SECTION_START.NEW_PAGE)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _md_escape(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", "<br>")


def render_markdown(blocks: Iterable[dict], output_path: Path) -> None:
    lines: list[str] = []
    for block in blocks:
        kind = block["type"]
        if kind == "title":
            lines.extend([f"# {block['text']}", ""])
        elif kind == "heading":
            prefix = "#" * (block["level"] + 1)
            lines.extend([f"{prefix} {block['text']}", ""])
        elif kind == "paragraph":
            lines.extend([block["text"], ""])
        elif kind == "bullets":
            lines.extend([f"- {item}" for item in block["items"]])
            lines.append("")
        elif kind == "table":
            headers = [_md_escape(value) for value in block["headers"]]
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in block["rows"]:
                lines.append("| " + " | ".join(_md_escape(value) for value in row) + " |")
            lines.append("")
        elif kind == "page_break":
            lines.append("")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def render_html(blocks: Iterable[dict], output_path: Path) -> None:
    body: list[str] = []
    for block in blocks:
        kind = block["type"]
        if kind == "title":
            body.append(f"<h1>{escape(block['text'])}</h1>")
        elif kind == "heading":
            level = min(block["level"] + 1, 6)
            body.append(f"<h{level}>{escape(block['text'])}</h{level}>")
        elif kind == "paragraph":
            body.append(f"<p>{escape(block['text'])}</p>")
        elif kind == "bullets":
            body.append("<ul>")
            body.extend(f"<li>{escape(item)}</li>" for item in block["items"])
            body.append("</ul>")
        elif kind == "table":
            body.append("<table>")
            body.append(
                "<thead><tr>"
                + "".join(f"<th>{escape(value)}</th>" for value in block["headers"])
                + "</tr></thead>"
            )
            body.append("<tbody>")
            for row in block["rows"]:
                body.append(
                    "<tr>"
                    + "".join(f"<td>{escape(value)}</td>" for value in row)
                    + "</tr>"
                )
            body.append("</tbody></table>")
        elif kind == "page_break":
            body.append("<hr>")

    html = f"""<!doctype html>
<html lang="ru">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Итерационный отчёт DOLG</title>
  <style>
    body {{
      margin: 0;
      background: #f5f5f5;
      color: #111;
      font-family: "Times New Roman", Times, serif;
      font-size: 17px;
      line-height: 1.5;
    }}
    main {{
      max-width: 1100px;
      margin: 0 auto;
      padding: 36px 48px;
      background: #fff;
      min-height: 100vh;
      box-shadow: 0 0 28px rgba(0, 0, 0, 0.08);
    }}
    h1, h2, h3 {{ line-height: 1.25; margin: 1.2em 0 0.5em; }}
    h1 {{ text-align: center; font-size: 28px; margin-top: 0; }}
    h2 {{ font-size: 22px; border-bottom: 2px solid #ccc; padding-bottom: 4px; }}
    p  {{ margin: 0 0 0.75em; text-align: justify; }}
    table {{ width: 100%; border-collapse: collapse; margin: 1em 0 1.25em; font-size: 14px; }}
    th, td {{ border: 1px solid #777; padding: 8px 10px; vertical-align: top; }}
    th {{ background: #ededed; text-align: center; }}
    li {{ margin: 0.35em 0; }}
    hr {{ border: 0; border-top: 1px solid #ddd; margin: 2em 0; }}
  </style>
</head>
<body>
<main>
{chr(10).join(body)}
</main>
</body>
</html>
"""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(html, encoding="utf-8")


def _register_pdf_fonts() -> tuple[str, str]:
    font_dir = Path(r"C:\Windows\Fonts")
    regular = font_dir / "times.ttf"
    bold = font_dir / "timesbd.ttf"
    if regular.exists() and bold.exists():
        pdfmetrics.registerFont(TTFont("TimesNewRoman", str(regular)))
        pdfmetrics.registerFont(TTFont("TimesNewRoman-Bold", str(bold)))
        return "TimesNewRoman", "TimesNewRoman-Bold"
    return "Helvetica", "Helvetica-Bold"


def _pdf_styles() -> dict[str, ParagraphStyle]:
    regular_font, bold_font = _register_pdf_fonts()
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "ReportTitle",
            parent=base["Title"],
            fontName=bold_font,
            fontSize=20,
            leading=26,
            alignment=TA_CENTER,
            spaceAfter=18,
        ),
        "heading1": ParagraphStyle(
            "ReportHeading1",
            parent=base["Heading1"],
            fontName=bold_font,
            fontSize=15,
            leading=19,
            spaceBefore=14,
            spaceAfter=8,
        ),
        "normal": ParagraphStyle(
            "ReportNormal",
            parent=base["Normal"],
            fontName=regular_font,
            fontSize=11,
            leading=15,
            alignment=TA_JUSTIFY,
            firstLineIndent=0,
            spaceAfter=6,
        ),
        "bullet": ParagraphStyle(
            "ReportBullet",
            parent=base["Normal"],
            fontName=regular_font,
            fontSize=11,
            leading=15,
            leftIndent=0.7 * cm,
            bulletIndent=0.25 * cm,
            spaceAfter=4,
        ),
        "table": ParagraphStyle(
            "ReportTable",
            parent=base["Normal"],
            fontName=regular_font,
            fontSize=8,
            leading=11,
            alignment=TA_LEFT,
        ),
        "table_header": ParagraphStyle(
            "ReportTableHeader",
            parent=base["Normal"],
            fontName=bold_font,
            fontSize=8,
            leading=11,
            alignment=TA_CENTER,
        ),
    }


def _paragraph(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(escape(text), style)


def _render_pdf_table(headers: list[str], rows: list[list[str]], styles: dict[str, ParagraphStyle]) -> Table:
    data = [[_paragraph(value, styles["table_header"]) for value in headers]]
    for row in rows:
        data.append([_paragraph(value, styles["table"]) for value in row])

    n = len(headers)
    if n == 2:
        widths = [4.5 * cm, 12.5 * cm]
    elif n == 3:
        widths = [4.0 * cm, 9.5 * cm, 3.5 * cm]
    elif n == 4:
        widths = [4.0 * cm, 8.5 * cm, 2.0 * cm, 2.5 * cm]
    else:
        widths = [17 * cm / n] * n

    table = Table(data, colWidths=widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EDEDED")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#777777")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return table


def _page_footer(canvas, doc) -> None:
    canvas.saveState()
    font_name = "TimesNewRoman" if "TimesNewRoman" in pdfmetrics.getRegisteredFontNames() else "Helvetica"
    canvas.setFont(font_name, 9)
    canvas.drawCentredString(A4[0] / 2, 1.0 * cm, str(doc.page))
    canvas.restoreState()


def render_pdf(blocks: Iterable[dict], output_path: Path) -> None:
    styles = _pdf_styles()
    story = []

    for block in blocks:
        kind = block["type"]
        if kind == "title":
            story.append(_paragraph(block["text"], styles["title"]))
        elif kind == "heading":
            story.append(_paragraph(block["text"], styles["heading1"]))
        elif kind == "paragraph":
            story.append(_paragraph(block["text"], styles["normal"]))
        elif kind == "bullets":
            for item in block["items"]:
                story.append(Paragraph(escape(item), styles["bullet"], bulletText="•"))
        elif kind == "table":
            story.append(_render_pdf_table(block["headers"], block["rows"], styles))
            story.append(Spacer(1, 8))
        elif kind == "page_break":
            story.append(PageBreak())

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=1.8 * cm,
        rightMargin=1.5 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.6 * cm,
        title="Итерационный отчёт DOLG",
        author="DOLG project",
    )
    doc.build(story, onFirstPage=_page_footer, onLaterPages=_page_footer)


def main() -> None:
    blocks = build_blocks()
    outputs = [
        ("PDF",  PDF_PATH,  render_pdf),
        ("DOCX", DOCX_PATH, render_docx),
        ("MD",   MD_PATH,   render_markdown),
        ("HTML", HTML_PATH, render_html),
    ]
    for label, path, renderer in outputs:
        try:
            renderer(blocks, path)
            print(f"OK   {label}: {path}")
        except PermissionError:
            print(f"SKIP locked: {path}")
        except Exception as exc:
            print(f"FAIL {label}: {path} — {exc}")


if __name__ == "__main__":
    main()

"""Добавляет в Главу 2 диплома листинги кода и скриншоты, НЕ трогая оформление.

- Существующие подписи рисунков (Рис 1..13 в приложениях) сдвигаются на +2,
  чтобы новые рисунки Главы 2 заняли сквозные номера 1 и 2 (по порядку появления).
- Листинги нумеруются отдельной последовательностью (Листинг 1..3).
- Все вставки — новые абзацы со своим форматированием; существующие абзацы,
  стили и таблицы не меняются.
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

DL = Path.home() / 'Downloads'
SRC = DL / 'Дипломная работа (домашняя версия 6.1.1) (1).docx'
# Имя по Положению о проверке (п.37): <год защиты>_<код направленности (профиля)>_<ФИО>.
# Код 09.03.03 — направление; точный «код направленности (профиля)» уточняется у научрука.
DST = DL / '2026_09.03.03_Буряко Дмитрий Сергеевич.docx'
SHOTS = Path(__file__).resolve().parents[1] / 'docs' / 'diploma_assets' / 'screenshots'
EDITOR_IMG = SHOTS / 'ch2_editor_canvas.png'
CAD_IMG = SHOTS / 'ch2_cad_gost.png'

EN = '–'  # en-dash, как в существующих подписях "Таблица 2.1 – ..."

LISTING1 = """class SchematicProject(models.Model):
    user = models.ForeignKey(settings.AUTH_USER_MODEL,
                             on_delete=models.CASCADE,
                             related_name='schematic_projects')
    name = models.CharField(max_length=200)
    category = models.CharField(max_length=20, choices=CATEGORY_CHOICES)
    is_demo = models.BooleanField(default=False)
    # Схема целиком: компоненты, соединения и метаданные в одном JSON-поле
    scheme_data = models.JSONField(default=dict, blank=True)
    # Токен для просмотра по ссылке /s/<token>/ (read-only)
    share_token = models.CharField(max_length=22, blank=True, db_index=True)
    # Мягкое удаление: вместо стирания ставим отметку времени
    deleted_at = models.DateTimeField(null=True, blank=True, db_index=True)
    visibility = models.CharField(max_length=10, choices=VISIBILITY_CHOICES,
                                  default='private')"""

LISTING2 = """class Algorithm(NamedTuple):
    key: str
    title: str
    intents: frozenset[str]
    run: Callable[[dict | None, str | None], list[str]]

# Реестр движков: число берётся из проверяемого кода, а не из текста модели
ALGORITHMS: list[Algorithm] = [
    Algorithm('dc_voltages', 'Расчёт DC напряжений (MNA)',
              frozenset({'measurement', 'overview', 'scheme_overview'}),
              lambda scheme, topology: ai_toolkit.dc_voltage_lines(scheme)),
    Algorithm('rf_filter', 'РЧ-анализ фильтра (scikit-rf)',
              frozenset({'formula', 'measurement'}),
              lambda scheme, topology: ai_toolkit.rf_filter_lines(scheme)),
    Algorithm('tolerance', 'Разброс номиналов (worst-case)',
              frozenset({'thermal', 'why_failed'}),
              lambda scheme, topology: ai_toolkit.tolerance_lines(scheme)),
    # также: power (MNA), formula (формулы), tiny_ai (нейроподсказка)
]

def sections_for_intent(intent, scheme_data, topology=None):
    \"\"\"Вызывает применимые к интенту движки и собирает их результат.\"\"\"
    sections = []
    for algo in ALGORITHMS:
        if intent not in algo.intents:
            continue
        lines = algo.run(scheme_data, topology)
        if lines:
            sections.append((algo.title, lines))
    return sections"""

LISTING3 = """def solve_dc(circuit: dict) -> dict:
    \"\"\"Решает DC методом MNA: формирует G и решает систему G·V = I.\"\"\"
    node_count = circuit['n_nodes'] - 1            # без узла земли
    v_sources = [e for e in circuit['elements'] if e['type'] == 'V']
    size = node_count + len(v_sources)
    A = np.zeros((size, size)); b = np.zeros(size)

    def stamp_g(n1, n2, g):                        # штамповка проводимости в G
        if n1 > 0: A[n1 - 1, n1 - 1] += g
        if n2 > 0: A[n2 - 1, n2 - 1] += g
        if n1 > 0 and n2 > 0:
            A[n1 - 1, n2 - 1] -= g; A[n2 - 1, n1 - 1] -= g

    for e in circuit['elements']:                  # резисторы -> проводимости
        if e['type'] == 'R' and e['value'] > 0:
            stamp_g(e['nodes'][0], e['nodes'][1], 1.0 / e['value'])
    for k, e in enumerate(v_sources):              # источники напряжения
        row = node_count + k
        n_p, n_n = e['nodes']
        if n_p > 0: A[n_p - 1, row] = A[row, n_p - 1] = 1
        if n_n > 0: A[n_n - 1, row] = A[row, n_n - 1] = -1
        b[row] = e['value']

    x = np.linalg.solve(A, b)                      # решение G·V = I
    return {'voltages': {i + 1: float(x[i]) for i in range(node_count)}}"""


def find_anchor(doc, needle):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise SystemExit(f'anchor not found: {needle!r}')


def new_para_after(anchor):
    p = OxmlElement('w:p')
    anchor._p.addnext(p)
    return Paragraph(p, anchor._parent)


def insert_blocks_after(anchor, builders):
    """builders: список функций build(par) -> None. Вставляются по порядку после anchor."""
    cur = anchor
    for build in builders:
        par = new_para_after(cur)
        build(par)
        cur = par
    return cur


def make_listing_caption(text):
    def build(par):
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        par.paragraph_format.space_before = Pt(6)
        par.paragraph_format.space_after = Pt(2)
        run = par.add_run(text)
        run.bold = True

    return build


def make_code_block(code):
    """Один абзац: строки кода через мягкий перенос, моноширинный шрифт."""
    lines = code.split('\n')

    def build(par):
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = par.paragraph_format
        pf.left_indent = Cm(0.5)
        pf.first_line_indent = Cm(0)
        pf.line_spacing = 1.0
        pf.space_after = Pt(8)
        for i, line in enumerate(lines):
            run = par.add_run(line if line else ' ')
            rpr_font(run, 'Consolas', 9)
            if i < len(lines) - 1:
                br = OxmlElement('w:br')
                run._r.append(br)

    return build


def rpr_font(run, name, size_pt):
    run.font.size = Pt(size_pt)
    run.font.name = name
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), name)


def make_image(path, width_cm, caption):
    def build_img(par):
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.space_before = Pt(6)
        par.paragraph_format.space_after = Pt(2)
        run = par.add_run()
        run.add_picture(str(path), width=Cm(width_cm))

    def build_cap(par):
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.space_after = Pt(8)
        par.add_run(caption)

    return [build_img, build_cap]


def renumber_figures(doc, delta=2):
    """Сдвигает номер в подписях 'Рисунок N ...' на delta. Возвращает список изменений."""
    changed = []
    pat = re.compile(r'^(\s*Рисунок\s+)(\d+)(\b)')
    for p in doc.paragraphs:
        m = pat.match(p.text)
        if not m:
            continue
        old_n = int(m.group(2))
        new_n = old_n + delta
        # меняем только в первом run, где лежит номер (как делали раньше)
        full = p.text
        new_full = pat.sub(lambda mm, n=new_n: f'{mm.group(1)}{n}{mm.group(3)}', full, count=1)
        if p.runs:
            p.runs[0].text = new_full
            for r in p.runs[1:]:
                r.text = ''
        changed.append((old_n, new_n))
    return changed


def scan_inline_refs(doc):
    hits = []
    pat = re.compile(r'рисун\w*\s+\d+|рис\.\s*\d+', re.IGNORECASE)
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith('Рисунок'):
            continue  # это подпись, не ссылка
        for m in pat.finditer(p.text):
            hits.append((i, m.group(0)))
    return hits


def main():
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    # 0. Диагностика ссылок на рисунки в тексте (до изменений).
    refs = scan_inline_refs(doc)
    print('INLINE FIGURE REFS (review):', refs if refs else 'none')

    # 1. Сдвигаем номера существующих подписей рисунков +2.
    changed = renumber_figures(doc, delta=2)
    print(f'renumbered captions: {len(changed)} ->', changed)

    # 2. Листинг 1 (2.1) после абзаца про scheme_data.
    a1 = find_anchor(doc, 'Поле scheme_data содержит список компонентов')
    insert_blocks_after(
        a1,
        [
            make_listing_caption(
                f'Листинг 2.1 {EN} Модель проекта схемы (фрагмент): хранение схемы в JSONField'
            ),
            make_code_block(LISTING1),
        ],
    )

    # 3. Листинг 2 (2.2) после абзаца про порядок «сначала объяснимое».
    a2 = find_anchor(doc, 'сначала то, что можно объяснить')
    insert_blocks_after(
        a2,
        [
            make_listing_caption(
                f'Листинг 2.2 {EN} Реестр расчётных движков ассистента и диспетчер по интенту'
            ),
            make_code_block(LISTING2),
        ],
    )

    # 4. Листинг 3 (2.3) после абзаца про формулы (2.1)-(2.6).
    a3 = find_anchor(doc, 'заложены в расчётные движки')
    insert_blocks_after(
        a3,
        [
            make_listing_caption(
                f'Листинг 2.3 {EN} Ядро DC-решателя MNA (фрагмент): формирование и решение G·V = I'
            ),
            make_code_block(LISTING3),
        ],
    )

    # 5. Рисунок 1 (2.3) после абзаца про редактор и BOM.
    f1 = find_anchor(doc, 'переключением между САПР и магазином')
    insert_blocks_after(
        f1,
        make_image(
            EDITOR_IMG,
            12.0,
            f'Рисунок 1 {EN} Редактор принципиальных схем с демонстрационной цепью (УГО-режим)',
        ),
    )

    # 6. Рисунок 2 (2.3) после абзаца про 2D-CAD и ГОСТ.
    f2 = find_anchor(doc, 'Вкладка энциклопедии убрана из CAD-панели')
    insert_blocks_after(
        f2,
        make_image(
            CAD_IMG, 15.5, f'Рисунок 2 {EN} Лист 2D-CAD с рамкой и основной надписью по ГОСТ (формат A3)'
        ),
    )

    doc.save(str(DST))
    print('SAVED:', DST)


if __name__ == '__main__':
    main()

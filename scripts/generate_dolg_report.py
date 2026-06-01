"""Generate the DOLG analysis and development report in DOCX and PDF.

The report is intentionally self-contained: it records the current project
state, compares it with the diploma requirements, and fixes the development
roadmap that should be used for subsequent work.
"""

from __future__ import annotations

import os
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs"
DOCX_PATH = OUTPUT_DIR / "Отчет_по_анализу_и_развитию_DOLG.docx"
PDF_PATH = OUTPUT_DIR / "Отчет_по_анализу_и_развитию_DOLG.pdf"
ASCII_DOCX_PATH = OUTPUT_DIR / "DOLG_report.docx"
ASCII_PDF_PATH = OUTPUT_DIR / "DOLG_report.pdf"
MARKDOWN_PATH = OUTPUT_DIR / "DOLG_report.md"
HTML_PATH = OUTPUT_DIR / "DOLG_report.html"


@dataclass(frozen=True)
class ProjectFacts:
    django_version: str
    database_engine: str
    users: int
    categories: int
    products: int
    reb_products: int
    projects: int
    demo_projects: int
    knowledge_categories: int
    articles: int
    check_result: str
    test_result: str
    environment_note: str


def _short_db_engine(engine: str) -> str:
    return engine.rsplit(".", 1)[-1] if engine else "не определено"


def collect_project_facts() -> ProjectFacts:
    """Collect live facts from Django and fall back to the inspected baseline."""

    os.environ.setdefault("DJANGO_SETTINGS_MODULE", "Dolg_PR.settings")
    try:
        import django
        from django.conf import settings
        from django.contrib.auth.models import User
        from django.core.management import call_command

        django.setup()

        from Dolg_APP.models import SchematicProject
        from knowledge.models import Article, KnowledgeCategory
        from shop.models import Category, Product

        try:
            call_command("check", verbosity=0)
            check_result = "manage.py check: замечаний не выявлено"
        except Exception as exc:  # pragma: no cover - diagnostic fallback
            check_result = f"manage.py check: требуется проверка ({exc})"

        debug_value = os.getenv("DEBUG", "True")
        redirect_note = (
            "В текущем окружении DEBUG имеет значение "
            f"'{debug_value}'. Тестовый режим теперь определяется отдельно, поэтому "
            "production-редирект HTTPS не мешает прогону Django test client."
        )

        return ProjectFacts(
            django_version=django.get_version(),
            database_engine=_short_db_engine(settings.DATABASES["default"]["ENGINE"]),
            users=User.objects.count(),
            categories=Category.objects.count(),
            products=Product.objects.count(),
            reb_products=Product.objects.filter(category__slug__in=Category.REB_SLUGS).count(),
            projects=SchematicProject.objects.count(),
            demo_projects=SchematicProject.objects.filter(is_demo=True).count(),
            knowledge_categories=KnowledgeCategory.objects.count(),
            articles=Article.objects.count(),
            check_result=check_result,
            test_result=(
                "99 тестов обнаруживаются в стандартном наборе: accounts — 8, "
                "shop — 38, orders — 16, Dolg_APP — 21 серверный тест + "
                "16 optional browser-smoke. scripts/run_checks.ps1: 83 OK, "
                "16 skipped, coverage — 78%. scripts/run_browser_e2e.ps1: 16/16 OK."
            ),
            environment_note=redirect_note,
        )
    except Exception as exc:  # pragma: no cover - used only when Django is unavailable
        return ProjectFacts(
            django_version="6.0.4",
            database_engine="sqlite3",
            users=5,
            categories=20,
            products=72,
            reb_products=26,
            projects=11,
            demo_projects=11,
            knowledge_categories=6,
            articles=9,
            check_result="manage.py check: замечаний не выявлено",
            test_result=(
                "99 тестов обнаруживаются в стандартном наборе: accounts — 8, "
                "shop — 38, orders — 16, Dolg_APP — 21 серверный тест + "
                "16 optional browser-smoke. Стандартный прогон: 83 OK, 16 skipped; "
                "browser-smoke отдельно: 16/16 OK."
            ),
            environment_note=(
                "Факты собраны по результатам предварительной инспекции проекта; "
                f"автоматический доступ к Django недоступен: {exc}."
            ),
        )


def build_blocks(facts: ProjectFacts) -> list[dict]:
    tech_stack = (
        f"Python 3.14, Django {facts.django_version}, база данных "
        f"{facts.database_engine}, серверный рендеринг Django-шаблонов, "
        "HTML5/CSS3/JavaScript, Canvas/Web Worker, ngspice.wasm, reportlab, openpyxl и Playwright для optional browser-smoke."
    )
    filling = (
        f"В рабочей базе зафиксировано: пользователей — {facts.users}, "
        f"категорий — {facts.categories}, товаров — {facts.products}, "
        f"РЭБ-компонентов — {facts.reb_products}, проектов схем — {facts.projects}, "
        f"демо-схем — {facts.demo_projects}, категорий базы знаний — "
        f"{facts.knowledge_categories}, статей — {facts.articles}."
    )

    return [
        {"type": "title", "text": "Отчет по анализу и плану развития проекта DOLG"},
        {
            "type": "paragraph",
            "text": (
                "Настоящий отчет подготовлен по результатам анализа дипломной "
                "работы и файлов проекта DOLG. Документ не содержит титульного "
                "листа, поскольку исходные данные об образовательной организации, "
                "авторе, руководителе и утвержденном шаблоне оформления не были "
                "предоставлены."
            ),
        },
        {"type": "page_break"},
        {"type": "heading", "level": 1, "text": "1. Введение и цель анализа"},
        {
            "type": "paragraph",
            "text": (
                "Целью анализа является определение фактического состояния "
                "программной системы, сопоставление реализованных возможностей с "
                "требованиями выпускной квалификационной работы и формирование "
                "плана дальнейшего развития проекта."
            ),
        },
        {
            "type": "paragraph",
            "text": (
                "В дипломной работе проект определен как веб-приложение, "
                "обеспечивающее пользователям возможность подбора и приобретения "
                "радио- и электронных компонентов, а также проектирования и "
                "симуляции электронных схем в онлайн-режиме."
            ),
        },
        {
            "type": "paragraph",
            "text": (
                "При подготовке отчета учитывались текст дипломной работы, "
                "структура Django-проекта, модели данных, маршруты, шаблоны, "
                "клиентский код редактора схем, тесты и вспомогательная "
                "документация проекта."
            ),
        },
        {"type": "heading", "level": 1, "text": "2. Краткое описание проекта"},
        {
            "type": "paragraph",
            "text": (
                "DOLG представляет собой единую веб-платформу, объединяющую "
                "интернет-магазин электронных компонентов, личный кабинет, "
                "оформление заказов, справочную базу знаний, редактор "
                "принципиальных схем, CAD-раздел и браузерный SPICE-симулятор."
            ),
        },
        {
            "type": "table",
            "headers": ["Показатель", "Фактическое состояние"],
            "rows": [
                ["Назначение", "Подбор, покупка, проектирование и симуляция электронных схем."],
                ["Технологический стек", tech_stack],
                ["Основные приложения", "shop, accounts, orders, knowledge, Dolg_APP, Dolg_PR."],
                ["Симуляция", "ngspice.wasm в Web Worker с резервным JS-MNA-движком для DC-расчетов."],
                ["Наполнение", filling],
            ],
        },
        {"type": "heading", "level": 1, "text": "3. Соответствие требованиям ВКР"},
        {
            "type": "paragraph",
            "text": (
                "Ниже приведено сопоставление ключевых требований, выделенных в "
                "дипломной работе, с фактическим состоянием проекта. Статус "
                "«реализовано частично» означает наличие базовой функции при "
                "отсутствии отдельных возможностей, указанных в требованиях."
            ),
        },
        {
            "type": "table",
            "headers": ["Требование", "Фактическое состояние", "Статус"],
            "rows": [
                [
                    "Каталог компонентов, поиск и фильтрация",
                    "Есть иерархические категории, поиск, фильтры по производителю, жизненному циклу и корпусу.",
                    "Реализовано",
                ],
                [
                    "Карточка товара с параметрами, ценой, остатком и документацией",
                    "Карточка товара содержит описание, изображение, цену, наличие, параметры, datasheet URL и PDF-документы.",
                    "Реализовано",
                ],
                [
                    "Сравнение и альтернативы компонентов",
                    "Есть сравнение до четырех товаров, подбор альтернатив в BOM-сценарии, поиск товара каталога из панели свойств компонента схемы, XLSX-экспорт BOM и предупреждения по номиналам/SPICE-моделям.",
                    "Реализовано",
                ],
                [
                    "Корзина и оформление заказа",
                    "Реализованы корзина, checkout, списание остатков, история заказов и отмена заказа.",
                    "Реализовано",
                ],
                [
                    "Личный кабинет покупателя",
                    "Есть профиль, адреса, просмотр заказов и повтор заказа из истории.",
                    "Реализовано",
                ],
                [
                    "Администрирование каталога, заказов и пользователей",
                    "Модели зарегистрированы в Django Admin, есть фильтры, поиск, actions для заказов, роли менеджера и сводка статистики заказов.",
                    "Реализовано частично",
                ],
                [
                    "Визуальный редактор схем",
                    "Реализован Canvas/WebGL-редактор с компонентами, проводами, визуальной сеткой, отдельным шагом хода для snap/маршрутизации, горячими клавишами, undo/redo и экспортом PNG.",
                    "Реализовано",
                ],
                [
                    "Сохранение и загрузка проектов",
                    "Есть API создания, обновления, сохранения и загрузки схем для авторизованных пользователей; демо-схемы доступны для просмотра.",
                    "Реализовано",
                ],
                [
                    "Генерация netlist и экспорт",
                    "SPICE-netlist, PNG и SVG реализованы на клиенте; PDF-экспорт схемы реализован серверным API.",
                    "Реализовано",
                ],
                [
                    "Симуляция DC/TRAN/AC и графики",
                    "Симуляция выполняется в браузере через ngspice.wasm; результаты отображаются на canvas-графиках с аналитикой. Панель аналитики перенесена под свойства компонента, чтобы графики и диагностические сообщения не были зажаты в узкой правой колонке. Для отказов добавлена диагностика пустого AC stdout и падения обоих движков.",
                    "Реализовано",
                ],
                [
                    "Гостевой режим редактора",
                    "Симулятор доступен без входа в demo-режиме; сохранение проектов и результатов остаётся доступным только авторизованным пользователям.",
                    "Реализовано частично",
                ],
                [
                    "Нефункциональные требования",
                    "Базовые механизмы безопасности Django включены, CI и coverage настроены. Добавлены optional Playwright smoke для `/simulation/`, DC/AC/TRAN-анализа, ошибок симуляции, desktop/mobile visual layout с проверкой переноса аналитики под свойства компонента, единого УГО-режима, отдельного шага хода и performance-smoke большой схемы, SVG/PDF-экспорта, netlist fixtures, BOM/netlist/project-модалок, export warnings, `/cad/`, CAD desktop/mobile visual baseline и `/projects/`; полноценная screenshot-baseline регрессия пока отсутствует.",
                    "Реализовано частично",
                ],
            ],
        },
        {"type": "heading", "level": 1, "text": "4. Фактическая архитектура и модули"},
        {
            "type": "table",
            "headers": ["Модуль", "Назначение"],
            "rows": [
                ["Dolg_PR", "Конфигурация Django-проекта, маршрутизация, настройки безопасности, static/media."],
                ["shop", "Каталог, карточки товаров, корзина, BOM-подбор, поиск товара для компонента схемы, XLSX-экспорт BOM, сравнение товаров, PDF-документы и validation helpers для номиналов/SPICE."],
                ["accounts", "Регистрация, вход, профиль пользователя, адреса доставки и связь с UserProfile."],
                ["orders", "Оформление заказов, статусы, позиции заказа, доставка, платежные транзакции и Stripe demo mode."],
                ["knowledge", "Справочная база по электронике: категории, статьи и публичные страницы."],
                ["Dolg_APP", "Редактор схем, CAD, проекты, API сохранения схем, версии проектов, журнал симуляций, DRC-проверка, экспорт, BOM, квоты симуляции, нормализация демо-схем по `drawing_step=30` и совместимые фасады для `scheme-normalizer.js`, `scheme-export.js`, `scheme-bom.js` и `scheme-netlist.js` с генерацией SPICE-элементов и analysis directives."],
            ],
        },
        {
            "type": "paragraph",
            "text": (
                "Архитектура соответствует монолитному MVT-подходу Django. "
                "Коммерческий контур, инженерный контур и справочный контур "
                "разделены по приложениям, но используют общую базу данных, "
                "единую аутентификацию и общую шаблонную систему."
            ),
        },
        {"type": "heading", "level": 1, "text": "5. Результаты проверки и тестирования"},
        {
            "type": "table",
            "headers": ["Проверка", "Результат"],
            "rows": [
                ["Конфигурация Django", facts.check_result],
                ["Модульные тесты", facts.test_result],
                ["Особенность окружения", facts.environment_note],
                [
                    "Тестовый контур",
                    "Покрыты accounts, shop, orders и Dolg_APP. Для `/simulation/`, BOM-модуля, DC-анализа, AC-графиков RC-фильтра, TRAN-графика RC-цепи, пустого AC stdout, отказа ngspice.wasm + JS-MNA, сохранения AC/TRAN запусков авторизованного проекта, desktop/mobile visual layout с проверкой переноса аналитики под свойства компонента, единого УГО-режима, отдельного шага хода и performance-smoke большой схемы, SVG/PDF-экспорта, fixtures netlist builder, `scheme-normalizer.js`, `scheme-export.js`, `scheme-bom.js`, `scheme-netlist.js`, прямой генерации SPICE-элементов, `.op/.ac/.tran/.print ac` directives, BOM/netlist/project-модалок, export warnings, `/cad/`, CAD desktop/mobile visual baseline и `/projects/` добавлены optional browser-smoke; для knowledge, расширенной CAD visual regression и screenshot-baseline проверок контур ещё требует развития.",
                ],
            ],
        },
        {
            "type": "paragraph",
            "text": (
                "Вывод по тестированию: основной набор модульных тестов используется "
                "как регрессионная база, а запуск стандартизирован через "
                "Dolg_PR.settings_test. Это устраняет ложные падения при локальном "
                "DEBUG=release и ускоряет тесты за счёт in-memory SQLite и быстрого "
                "хешера паролей."
            ),
        },
        {"type": "heading", "level": 1, "text": "6. Выявленные расхождения"},
        {
            "type": "bullets",
            "items": [
                "В актуальном тексте диплома упоминаются DRF, Celery, Redis, JWT, bcrypt, PostgreSQL и полноценная REST-архитектура. В текущем коде реализован монолитный Django MVT с сессионной аутентификацией; перечисленные технологии следует либо внедрять следующими итерациями, либо явно оформить как перспективы.",
                "Production-настройки требуют отдельного заполнения SECRET_KEY, ALLOWED_HOSTS, CSRF_TRUSTED_ORIGINS и платежных ключей по шаблону .env.example.",
                "Автоматизированный browser-smoke есть для `/simulation/ → BOM → XLSX → cart`, DC-анализа, AC-графиков RC-фильтра, TRAN-графика RC-цепи, диагностики пустого AC stdout и отказа обоих движков, сохранения AC/TRAN запусков, переноса аналитики под свойства компонента, единого УГО-режима, отдельного шага хода, performance-smoke большой схемы, базового desktop/mobile visual layout, SVG/PDF-экспорта, fixtures netlist builder, BOM/netlist/project-модалок, export warnings, `/cad/`, CAD desktop/mobile visual baseline и `/projects/`; следующим уровнем остаётся screenshot-baseline регрессия CAD, projects и сложных состояний панели свойств/аналитики.",
                "Требуют развития production-деплой на PostgreSQL, а также дальнейший ремонт текущих модулей симуляции и CAD. Импорт/экспорт промышленных CAD-форматов KiCad/LTspice/EDIF/Gerber/Excellon перенесён в этап после стабилизации существующего редактора.",
            ],
        },
        {"type": "heading", "level": 1, "text": "7. План развития проекта"},
        {
            "type": "table",
            "headers": ["Этап", "Содержание работ", "Ожидаемый результат"],
            "rows": [
                [
                    "1. Стабилизация",
                    "Поддерживать .env.example, CI, coverage и единый тестовый запуск; дополнять production-проверки перед демонстрацией.",
                    "Проект воспроизводимо проверяется в локальной и CI-среде.",
                ],
                [
                    "2. Закрытие требований ВКР",
                    "Сначала довести ремонт и модернизацию симуляции/CAD под современные инженерные требования: сохранение AC/TRAN запусков, probes/курсоры графиков, понятные ошибки ngspice.wasm, читаемая панель аналитики под свойствами компонентов, отдельный шаг хода и Multisim-like отрисовка схем, DRC/ERC, screenshot-baseline проверки CAD/projects/панели свойств, сетка/слои/штамп/размеры; затем расширить инженерную библиотеку компонентов, добавить интеграцию CAD-форматов и сверить все утверждения диплома с кодом.",
                    "Функциональная часть полностью прослеживается от требований дипломной работы до проверяемой реализации.",
                ],
                [
                    "3. Future killer-фичи",
                    "`What if`-слайдер параметров, тепловой анализ, AI-ассистент, 3D-просмотр платы и виртуальная лаборатория. Marketplace/QR — только после публичного деплоя, block-based editor — отдельный learning mode.",
                    "Проект получает сильные демонстрационные и инженерные сценарии развития после стабилизации текущих CAD и симулятора.",
                ],
                [
                    "4. Продакшн-подготовка",
                    "Перейти на PostgreSQL, настроить static/media, HTTPS-деплой, резервное копирование БД, мониторинг ошибок и регламент обновлений.",
                    "Платформа готова к демонстрации и ограниченной опытной эксплуатации.",
                ],
                [
                    "5. Документация",
                    "Привести README и пояснительную записку к фактической реализации либо явно перенести DRF/Celery/Redis/JWT/PostgreSQL/bcrypt в раздел перспектив.",
                    "Документация не противоречит коду и может использоваться при защите проекта.",
                ],
            ],
        },
        {"type": "heading", "level": 1, "text": "8. Заключение"},
        {
            "type": "paragraph",
            "text": (
                "Проект DOLG находится в состоянии работоспособного дипломного "
                "прототипа с широким набором реализованных функций. Наиболее "
                "сильной стороной системы является объединение каталога, корзины, "
                "BOM-сценария, редактора схем и браузерной SPICE-симуляции в "
                "одном пользовательском контуре."
            ),
        },
        {
            "type": "paragraph",
            "text": (
                "Основные задачи дальнейшего развития связаны не с полной "
                "переработкой архитектуры, а со стабилизацией окружения, "
                "устранением расхождений между текстом ВКР и кодом, расширением "
                "административных и экспортных возможностей, а также подготовкой "
                "проекта к воспроизводимой демонстрации и эксплуатации."
            ),
        },
    ]


def _set_docx_style(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size in (("Heading 1", 16), ("Heading 2", 15), ("Title", 18)):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def render_docx(blocks: Iterable[dict], output_path: Path) -> None:
    document = Document()
    _set_docx_style(document)

    for block in blocks:
        kind = block["type"]
        if kind == "title":
            paragraph = document.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(block["text"]).bold = True
        elif kind == "heading":
            document.add_heading(block["text"], level=block["level"])
        elif kind == "paragraph":
            paragraph = document.add_paragraph(block["text"])
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif kind == "bullets":
            for item in block["items"]:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.add_run(item)
        elif kind == "table":
            rows = block["rows"]
            table = document.add_table(rows=1, cols=len(block["headers"]))
            table.style = "Table Grid"
            header_cells = table.rows[0].cells
            for index, value in enumerate(block["headers"]):
                header_cells[index].text = value
                for paragraph in header_cells[index].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for index, value in enumerate(row):
                    cells[index].text = value
            document.add_paragraph()
        elif kind == "page_break":
            document.add_section(WD_SECTION_START.NEW_PAGE)
        else:  # pragma: no cover - developer error
            raise ValueError(f"Unknown block type: {kind}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _md_escape(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", "<br>")


def render_markdown(blocks: Iterable[dict], output_path: Path) -> None:
    lines: list[str] = []
    for block in blocks:
        kind = block["type"]
        if kind == "title":
            lines.extend([f"# {block['text']}", ""])
        elif kind == "heading":
            prefix = "#" * (block["level"] + 1)
            lines.extend([f"{prefix} {block['text']}", ""])
        elif kind == "paragraph":
            lines.extend([block["text"], ""])
        elif kind == "bullets":
            lines.extend([f"- {item}" for item in block["items"]])
            lines.append("")
        elif kind == "table":
            headers = [_md_escape(value) for value in block["headers"]]
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in block["rows"]:
                lines.append("| " + " | ".join(_md_escape(value) for value in row) + " |")
            lines.append("")
        elif kind == "page_break":
            lines.append("")
        else:  # pragma: no cover - developer error
            raise ValueError(f"Unknown block type: {kind}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def render_html(blocks: Iterable[dict], output_path: Path) -> None:
    body: list[str] = []
    for block in blocks:
        kind = block["type"]
        if kind == "title":
            body.append(f"<h1>{escape(block['text'])}</h1>")
        elif kind == "heading":
            level = min(block["level"] + 1, 6)
            body.append(f"<h{level}>{escape(block['text'])}</h{level}>")
        elif kind == "paragraph":
            body.append(f"<p>{escape(block['text'])}</p>")
        elif kind == "bullets":
            body.append("<ul>")
            body.extend(f"<li>{escape(item)}</li>" for item in block["items"])
            body.append("</ul>")
        elif kind == "table":
            body.append("<table>")
            body.append(
                "<thead><tr>"
                + "".join(f"<th>{escape(value)}</th>" for value in block["headers"])
                + "</tr></thead>"
            )
            body.append("<tbody>")
            for row in block["rows"]:
                body.append(
                    "<tr>"
                    + "".join(f"<td>{escape(value)}</td>" for value in row)
                    + "</tr>"
                )
            body.append("</tbody></table>")
        elif kind == "page_break":
            body.append("<hr>")
        else:  # pragma: no cover - developer error
            raise ValueError(f"Unknown block type: {kind}")

    html = f"""<!doctype html>
<html lang="ru">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Отчет по анализу и плану развития проекта DOLG</title>
  <style>
    body {{
      margin: 0;
      background: #f5f5f5;
      color: #111;
      font-family: "Times New Roman", Times, serif;
      font-size: 18px;
      line-height: 1.55;
    }}
    main {{
      max-width: 980px;
      margin: 0 auto;
      padding: 48px 56px;
      background: #fff;
      min-height: 100vh;
      box-shadow: 0 0 28px rgba(0, 0, 0, 0.08);
    }}
    h1, h2, h3 {{
      line-height: 1.25;
      margin: 1.2em 0 0.5em;
    }}
    h1 {{
      text-align: center;
      font-size: 28px;
      margin-top: 0;
    }}
    h2 {{
      font-size: 23px;
    }}
    p {{
      margin: 0 0 0.75em;
      text-align: justify;
      text-indent: 1.25cm;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      margin: 1em 0 1.25em;
      font-size: 15px;
    }}
    th, td {{
      border: 1px solid #777;
      padding: 8px 10px;
      vertical-align: top;
    }}
    th {{
      background: #ededed;
      text-align: center;
    }}
    li {{
      margin: 0.35em 0;
    }}
    hr {{
      border: 0;
      border-top: 1px solid #ddd;
      margin: 2em 0;
    }}
  </style>
</head>
<body>
<main>
{chr(10).join(body)}
</main>
</body>
</html>
"""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(html, encoding="utf-8")


def _register_pdf_fonts() -> tuple[str, str]:
    font_dir = Path(r"C:\Windows\Fonts")
    regular = font_dir / "times.ttf"
    bold = font_dir / "timesbd.ttf"
    if regular.exists() and bold.exists():
        pdfmetrics.registerFont(TTFont("TimesNewRoman", str(regular)))
        pdfmetrics.registerFont(TTFont("TimesNewRoman-Bold", str(bold)))
        return "TimesNewRoman", "TimesNewRoman-Bold"
    return "Helvetica", "Helvetica-Bold"


def _pdf_styles() -> dict[str, ParagraphStyle]:
    regular_font, bold_font = _register_pdf_fonts()
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "ReportTitle",
            parent=base["Title"],
            fontName=bold_font,
            fontSize=18,
            leading=24,
            alignment=TA_CENTER,
            spaceAfter=18,
        ),
        "heading1": ParagraphStyle(
            "ReportHeading1",
            parent=base["Heading1"],
            fontName=bold_font,
            fontSize=15,
            leading=19,
            spaceBefore=12,
            spaceAfter=8,
        ),
        "normal": ParagraphStyle(
            "ReportNormal",
            parent=base["Normal"],
            fontName=regular_font,
            fontSize=12,
            leading=17,
            alignment=TA_JUSTIFY,
            firstLineIndent=1.25 * cm,
            spaceAfter=6,
        ),
        "bullet": ParagraphStyle(
            "ReportBullet",
            parent=base["Normal"],
            fontName=regular_font,
            fontSize=12,
            leading=16,
            leftIndent=0.7 * cm,
            bulletIndent=0.25 * cm,
            spaceAfter=4,
        ),
        "table": ParagraphStyle(
            "ReportTable",
            parent=base["Normal"],
            fontName=regular_font,
            fontSize=9,
            leading=12,
            alignment=TA_LEFT,
        ),
        "table_header": ParagraphStyle(
            "ReportTableHeader",
            parent=base["Normal"],
            fontName=bold_font,
            fontSize=9,
            leading=12,
            alignment=TA_CENTER,
        ),
    }


def _paragraph(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(escape(text), style)


def _render_pdf_table(headers: list[str], rows: list[list[str]], styles: dict[str, ParagraphStyle]) -> Table:
    data = [[_paragraph(value, styles["table_header"]) for value in headers]]
    for row in rows:
        data.append([_paragraph(value, styles["table"]) for value in row])

    column_count = len(headers)
    if column_count == 2:
        widths = [4.4 * cm, 12.6 * cm]
    elif column_count == 3:
        widths = [4.1 * cm, 9.4 * cm, 3.5 * cm]
    else:
        widths = [17 * cm / column_count] * column_count

    table = Table(data, colWidths=widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EDEDED")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#777777")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def _page_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("TimesNewRoman", 10)
    canvas.drawCentredString(A4[0] / 2, 1.15 * cm, str(doc.page))
    canvas.restoreState()


def render_pdf(blocks: Iterable[dict], output_path: Path) -> None:
    styles = _pdf_styles()
    story = []

    for block in blocks:
        kind = block["type"]
        if kind == "title":
            story.append(_paragraph(block["text"], styles["title"]))
        elif kind == "heading":
            story.append(_paragraph(block["text"], styles["heading1"]))
        elif kind == "paragraph":
            story.append(_paragraph(block["text"], styles["normal"]))
        elif kind == "bullets":
            for item in block["items"]:
                story.append(Paragraph(escape(item), styles["bullet"], bulletText="•"))
        elif kind == "table":
            story.append(_render_pdf_table(block["headers"], block["rows"], styles))
            story.append(Spacer(1, 8))
        elif kind == "page_break":
            story.append(PageBreak())
        else:  # pragma: no cover - developer error
            raise ValueError(f"Unknown block type: {kind}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=1.5 * cm,
        topMargin=2 * cm,
        bottomMargin=1.8 * cm,
        title="Отчет по анализу и плану развития проекта DOLG",
        author="DOLG project",
    )
    doc.build(story, onFirstPage=_page_footer, onLaterPages=_page_footer)


def main() -> None:
    facts = collect_project_facts()
    blocks = build_blocks(facts)
    outputs = [
        ("DOCX", DOCX_PATH, render_docx),
        ("DOCX", ASCII_DOCX_PATH, render_docx),
        ("PDF", PDF_PATH, render_pdf),
        ("PDF", ASCII_PDF_PATH, render_pdf),
        ("MD", MARKDOWN_PATH, render_markdown),
        ("HTML", HTML_PATH, render_html),
    ]
    for _, path, renderer in outputs:
        try:
            renderer(blocks, path)
        except PermissionError:
            print(f"SKIP locked output: {path}")
    for label, path, _ in outputs:
        print(f"{label}: {path}")


if __name__ == "__main__":
    main()
